#!/usr/bin/env python3
"""
PPT Skill — 口播稿 docx 生成模板
对齐 PM-AI-SOP-script-v3.docx 视觉规范：
- PingFang SC · slate 调色板（#1E293B 标题 / #475569 正文 / #E5E7EB 分隔线 / #2D81FF eyebrow / #94A3B8 过渡）
- 流动段落（非 bullet 列表）· 行距 1.6 · 内联加粗强调
- 每页：eyebrow chip + 大标题 + 薄横线 + 正文段 + 斜体过渡句

用法：复制到项目 scripts/ 目录，填入 NOTES 数据，运行 python3 gen_notes_v1.py
产物：ppt-{主题}-notes-v{N}.docx（微信发手机当提词器）

NOTES 数据结构：每页一个 dict
- eyebrow: str         eyebrow 标签（如「开场 · 02」「核心机制 · 03」）
- title: str           页标题
- paragraphs: list     正文段落列表，每项可以是：
    - str            纯正文段
    - list[(str, bool)]  含内联加粗的段，bool=True 表示该段加粗 + slate-800 深色
- transition: str      过渡句（斜体浅灰，自然衔接下一页）
"""

import os

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ================================================================
# 视觉常量（不要随便改，对齐 PM-AI-SOP-script-v3.docx）
# ================================================================
FONT_CN = 'PingFang SC'
FONT_EN = 'PingFang SC'

CLR_TITLE = '1E293B'       # slate-800 · 标题 / 内联强调
CLR_BODY = '475569'        # slate-600 · 正文
CLR_EYEBROW = '2D81FF'     # blue · eyebrow 标签
CLR_DIVIDER = 'E5E7EB'     # gray-200 · 标题下薄横线
CLR_TRANSITION = '94A3B8'  # slate-400 · 过渡句

# ================================================================
# 数据 — 每次使用时替换为实际口播内容
# ================================================================
NOTES = [
    {
        'eyebrow': '开场 · 01',
        'title': 'Tab 标题',
        'paragraphs': [
            # 纯正文段
            '这是一段正文，描述本页核心论点。演讲时开门见山用这一段。',
            # 含内联加粗的段（强调词加粗 + 深色）
            [
                ('讲解要点 1 · ', False),
                ('关键词', True),
                (' 后面跟描述。', False),
            ],
            [
                ('讲解要点 2 · 用', False),
                ('内联加粗', True),
                ('代替 bullet 列表，自然朗读不卡顿。', False),
            ],
        ],
        'transition': '→ 下一页讲 XXX，承接关系是 YYY',
    },
    # ... 每页一个对象
]

OUTPUT_PATH = '../deliverables/ppt-{主题}-notes-v1.docx'  # 修改为实际路径

# ================================================================
# 生成逻辑（不需要改）
# ================================================================

def set_run_style(run, *, size_pt=12, bold=False, color_hex=CLR_BODY, italic=False):
    """统一字体 / 字号 / 颜色 / 加粗"""
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color_hex)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), FONT_EN)
    rFonts.set(qn('w:hAnsi'), FONT_EN)
    rFonts.set(qn('w:eastAsia'), FONT_CN)
    rFonts.set(qn('w:cs'), FONT_EN)
    rPr.insert(0, rFonts)


def add_horizontal_rule(paragraph, color_hex=CLR_DIVIDER):
    """在段落底部加一条横线（用 pBdr/bottom 实现）"""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def set_line_spacing(paragraph, line_pt=20, space_after_pt=9, space_before_pt=0):
    """正文 1.6 倍行距 + 段后 9pt"""
    pf = paragraph.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.6
    pf.space_after = Pt(space_after_pt)
    pf.space_before = Pt(space_before_pt)


def render_paragraph_runs(paragraph, content):
    """content 可以是 str 或 [(text, bold?), ...] 数组"""
    if isinstance(content, str):
        run = paragraph.add_run(content)
        set_run_style(run, size_pt=12, color_hex=CLR_BODY)
    else:
        for text, bold in content:
            run = paragraph.add_run(text)
            set_run_style(run, size_pt=12, bold=bold,
                          color_hex=(CLR_TITLE if bold else CLR_BODY))


def build_notes(notes_data, output_path):
    doc = Document()

    section = doc.sections[0]
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)

    # 全局默认字体
    style = doc.styles['Normal']
    style.font.name = FONT_EN
    style.element.rPr.rFonts.set(qn('w:eastAsia'), FONT_CN)

    for idx, page in enumerate(notes_data):
        # ── eyebrow（蓝色小字标签）──
        eb_para = doc.add_paragraph()
        eb_pf = eb_para.paragraph_format
        eb_pf.space_after = Pt(4)
        eb_pf.space_before = Pt(0)
        eb_run = eb_para.add_run(page['eyebrow'])
        set_run_style(eb_run, size_pt=9, bold=True, color_hex=CLR_EYEBROW)

        # ── title（大字深色）──
        title_para = doc.add_paragraph()
        title_pf = title_para.paragraph_format
        title_pf.space_before = Pt(2)
        title_pf.space_after = Pt(6)
        t_run = title_para.add_run(page['title'])
        set_run_style(t_run, size_pt=18, bold=True, color_hex=CLR_TITLE)

        # ── 分隔线（薄灰色 horizontal rule）──
        rule_para = doc.add_paragraph()
        rule_para.paragraph_format.space_after = Pt(12)
        rule_para.paragraph_format.space_before = Pt(0)
        add_horizontal_rule(rule_para)

        # ── 正文段落（流动段，非 bullet）──
        for content in page['paragraphs']:
            p = doc.add_paragraph()
            set_line_spacing(p)
            render_paragraph_runs(p, content)

        # ── 过渡句（斜体浅灰，向下指向下页）──
        if page.get('transition'):
            t_para = doc.add_paragraph()
            t_pf = t_para.paragraph_format
            t_pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            t_pf.line_spacing = 1.4
            t_pf.space_before = Pt(14)
            t_pf.space_after = Pt(6)
            t_run = t_para.add_run(page['transition'])
            set_run_style(t_run, size_pt=11, italic=True, color_hex=CLR_TRANSITION)

        # 每页分页符（最后一页不加）
        if idx < len(notes_data) - 1:
            add_page_break(doc)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)) or '.', exist_ok=True)
    doc.save(output_path)
    print(f'✅ 已生成: {output_path}')
    print(f'   页数: {len(notes_data)}')


if __name__ == '__main__':
    build_notes(NOTES, OUTPUT_PATH)

#!/usr/bin/env python3
# PM-Workspace | (c) 2026 CaufieldZ | Apache 2.0 + AI Training Restriction
"""
PPT Skill — 口播稿 docx 生成模板
用法：复制到项目 scripts/ 目录，填入 NOTES 数据，运行 python3 gen_notes_v1.py
产物：ppt-{主题}-notes-v{N}.docx（微信发手机当提词器）
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ================================================================
# 数据 — 每次使用时替换为实际口播内容
# ================================================================
NOTES = [
    {
        'id': 'page-id',
        'title': 'Tab 标题',
        'core': '一句话核心论点，演讲时开门见山用这一句',
        'points': [
            '讲解要点 1',
            '讲解要点 2',
            '讲解要点 3',
        ],
        'transition': '→ 下一页讲 XXX，承接关系是 YYY',
    },
    # ... 每页一个对象
]

OUTPUT_PATH = '../deliverables/ppt-{主题}-notes-v1.docx'  # 修改为实际路径

# ================================================================
# 生成逻辑
# ================================================================

def set_font_cn(run, size_pt=16):
    """设置中文字体（微软雅黑）+ 字号"""
    run.font.size = Pt(size_pt)
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:eastAsia'), '微软雅黑')
    rFonts.set(qn('w:ascii'), 'Calibri')
    rPr.insert(0, rFonts)


def add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def build_notes(notes_data, output_path):
    doc = Document()

    # 全局行距 + 页边距（手机阅读友好）
    section = doc.sections[0]
    section.left_margin = Pt(48)
    section.right_margin = Pt(48)
    section.top_margin = Pt(36)
    section.bottom_margin = Pt(36)

    for idx, page in enumerate(notes_data):
        # ── 标题 ──
        h = doc.add_heading(page['title'], level=1)
        h.runs[0].font.size = Pt(20)
        h.runs[0].font.bold = True
        for run in h.runs:
            rPr = run._r.get_or_add_rPr()
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:eastAsia'), '微软雅黑')
            rFonts.set(qn('w:ascii'), 'Calibri')
            rPr.insert(0, rFonts)

        # ── 核心论点（加粗） ──
        core_para = doc.add_paragraph()
        core_para.paragraph_format.space_before = Pt(6)
        core_run = core_para.add_run(page['core'])
        core_run.bold = True
        set_font_cn(core_run, 16)

        # ── 讲解要点（列表） ──
        for point in page.get('points', []):
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Pt(12)
            p.paragraph_format.space_before = Pt(2)
            run = p.add_run(point)
            set_font_cn(run, 15)

        # ── 过渡句（斜体，灰色） ──
        if page.get('transition'):
            trans_para = doc.add_paragraph()
            trans_para.paragraph_format.space_before = Pt(8)
            trans_run = trans_para.add_run(page['transition'])
            trans_run.italic = True
            trans_run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            set_font_cn(trans_run, 13)

        # 每页分页符（最后一页不加）
        if idx < len(notes_data) - 1:
            add_page_break(doc)

    os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    doc.save(output_path)
    print(f'✅ 已生成: {output_path}')
    print(f'   页数: {len(notes_data)}')


if __name__ == '__main__':
    build_notes(NOTES, OUTPUT_PATH)

#!/usr/bin/env python3
# PM-Workspace | (c) 2026 CaufieldZ | Apache 2.0 + AI Training Restriction
# pm-ws-canary-236a5364
"""
PRD 通用框架函数库
各项目的 gen_prd_vN.py 通过 import 使用，只需写内容区。

用法：
    import sys, os
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), '../../../.claude/skills/prd/references'))
    from gen_prd_base import *
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os, re

# ── 颜色常量 ──────────────────────────────────────────────────────────────
C = {
    "textPrimary":   "333333",
    "textSecondary": "666666",
    "textMuted":     "888888",
    "textHeading":   "1A1A2E",
    "textLink":      "2E75B6",
    "tableHeaderBg": "2D81FF",
    "tableHeaderText": "FFFFFF",
    "tableAltRow":   "F8FAFB",
    "tableBorder":   "CCCCCC",
    "tagNew":        "15803D",
    "tagChange":     "D97706",
    "tagGreen":      "15803D",
    "accentBlue":    "2E75B6",
}

# ── 底层工具 ──────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    """设置单元格边框。用法：set_cell_border(cell, bottom={'val':'single','sz':4,'color':'2E75B6'})"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            tag.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), kwargs[edge].get('color', 'CCCCCC'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def para_run(para, text, font="Arial", size_pt=10, bold=False, color=None, italic=False):
    run = para.add_run(text)
    run.font.name = font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:eastAsia'), 'Arial')
    rPr.insert(0, rFonts)
    return run


_MD_BOLD_RE = __import__("re").compile(r"\*\*(.+?)\*\*")

def para_run_md(para, text, size_pt=10, color=None):
    """渲染内联 **bold** 片段。非 bold 段落继承默认颜色，bold 段落用 textHeading。"""
    pos = 0
    for m in _MD_BOLD_RE.finditer(text):
        if m.start() > pos:
            para_run(para, text[pos:m.start()], size_pt=size_pt, color=color)
        para_run(para, m.group(1), size_pt=size_pt, bold=True,
                 color=C["textHeading"] if color is None else color)
        pos = m.end()
    if pos < len(text):
        para_run(para, text[pos:], size_pt=size_pt, color=color)

# ── 段落级 ────────────────────────────────────────────────────────────────

def add_p(doc, text="", size_pt=10, bold=False, color=None, italic=False,
          align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        para_run(p, text, size_pt=size_pt, bold=bold, color=color, italic=italic)
    return p

def h1(doc, text):
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    para_run(p, text, size_pt=16, bold=True, color=C["textHeading"])
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '2E75B6')
    pBdr.append(bottom)
    pPr.append(pBdr)

def h2(doc, text):
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    para_run(p, text, size_pt=13, bold=True, color=C["accentBlue"])

def h3(doc, text):
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    para_run(p, text, size_pt=11, bold=True, color=C["textHeading"])

def bullet(doc, text, size_pt=10):
    lines = text.split('\n')
    first = True
    for line in lines:
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0 if first else 0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        para_run(p, f"• {line}", size_pt=size_pt, color=C["textPrimary"])
        first = False

# ── 表格级 ────────────────────────────────────────────────────────────────

def cell_text(cell, text, size_pt=9, bold=False, color=None, italic=False,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    para_run(p, text, size_pt=size_pt, bold=bold, color=color, italic=italic)
    return p

def make_table(doc, headers, rows_data, col_widths_cm, row_height_cm=None):
    """通用表格生成：headers=列名列表, rows_data=二维列表, col_widths_cm=每列宽度"""
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, (hdr, w) in enumerate(zip(headers, col_widths_cm)):
        cell = tbl.rows[0].cells[j]
        cell.width = Cm(w)
        set_cell_bg(cell, C["tableHeaderBg"])
        cell_text(cell, hdr, size_pt=9, bold=True, color=C["tableHeaderText"],
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows_data):
        for j, (val, w) in enumerate(zip(row, col_widths_cm)):
            cell = tbl.rows[i+1].cells[j]
            cell.width = Cm(w)
            if i % 2 == 1:
                set_cell_bg(cell, C["tableAltRow"])
            cell_text(cell, str(val), size_pt=9, color=C["textPrimary"])
        if row_height_cm:
            tbl.rows[i+1].height = Cm(row_height_cm)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

def scene_table(doc, scene_id, scene_name, right_blocks):
    """
    两列 Scene 表格
    right_blocks: list of (title, lines) -- title 可含「（变更）」或「（新增）」
    """
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    left_cell = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]
    left_cell.width = Cm(8.5)
    right_cell.width = Cm(13.0)
    set_cell_bg(left_cell, C["tableAltRow"])
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = left_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    para_run(p, f"\U0001f4f1 {scene_id} \u00b7 {scene_name}", size_pt=9, color=C["textMuted"], italic=True)
    p2 = left_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_run(p2, "\u2190 \u6b64\u5904\u7c98\u8d34\u539f\u578b\u622a\u56fe", size_pt=8, color=C["textMuted"], italic=True)

    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    fill_cell_blocks(right_cell, right_blocks)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


_NUMBERED_PREFIX_RE = re.compile(r"^\d+\.\s")

def fill_cell_blocks(cell, blocks, numbered=True):
    """
    在已存在的 cell 内填充结构化 blocks（title 粗体 + 子条目缩进）。
    假设 cell 的首段已为空或将被覆写。不清空 cell，不会处理旧内容。
    被 scene_table 和 update_prd_base.set_cell_blocks 共用。

    blocks: list[tuple[str, list[str]]]
    numbered: True 时自动给每个 block 内的主级 line 加 "1./2./3." 前缀
              （已含 "N. " 前缀的 line 保持原样，二级缩进行 "  - " 不编号）。
              prd-template.md 规定"每个模块下用 1./2./3. 数字编号"，默认开启。
    """
    first = True
    for (title, lines) in blocks:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        if "\uff08\u53d8\u66f4\uff09" in title or "\uff08\u65b0\u589e\uff09" in title:
            tag = "\uff08\u53d8\u66f4\uff09" if "\uff08\u53d8\u66f4\uff09" in title else "\uff08\u65b0\u589e\uff09"
            base = title.replace(tag, "")
            para_run(p, base, size_pt=10, bold=True, color=C["textHeading"])
            para_run(p, tag, size_pt=10, bold=True, color=C["tagChange"])
        else:
            para_run(p, title, size_pt=10, bold=True, color=C["textHeading"])
        counter = 0
        for line in lines:
            pl = cell.add_paragraph()
            pl.paragraph_format.space_before = Pt(0)
            pl.paragraph_format.space_after = Pt(1)
            stripped = line.lstrip()
            is_sublevel = stripped.startswith("- ") and line != stripped
            if is_sublevel:
                pl.paragraph_format.left_indent = Cm(0.9)
                body = stripped
            else:
                pl.paragraph_format.left_indent = Cm(0.3)
                if numbered and not _NUMBERED_PREFIX_RE.match(line):
                    counter += 1
                    body = f"{counter}. {line}"
                else:
                    body = line
            para_run_md(pl, body, size_pt=9, color=C["textPrimary"])

# ── 文档初始化 ────────────────────────────────────────────────────────────

def init_doc(landscape=True):
    """创建并返回一个配置好页面的 Document 对象"""
    doc = Document()
    section = doc.sections[0]
    if landscape:
        section.page_width  = Cm(27.94)
        section.page_height = Cm(21.59)
        section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(1.905)
    section.bottom_margin = Cm(1.905)
    return doc

def cover_page(doc, title, subtitle="产品需求文档（PRD）", scope="", meta_rows=None):
    """
    生成封面页
    meta_rows: list of [label, value] 用于文档属性表
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(4)
    para_run(p, title, size_pt=24, bold=True, color=C["textHeading"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    para_run(p, subtitle, size_pt=14, color=C["textSecondary"])

    if scope:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(30)
        para_run(p, scope, size_pt=10, color=C["textMuted"], italic=True)

    if meta_rows:
        meta_tbl = doc.add_table(rows=len(meta_rows), cols=2)
        meta_tbl.style = 'Table Grid'
        meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, value) in enumerate(meta_rows):
            cell_text(meta_tbl.rows[i].cells[0], label, size_pt=9, bold=True,
                      color=C["textSecondary"], align=WD_ALIGN_PARAGRAPH.RIGHT)
            meta_tbl.rows[i].cells[0].width = Cm(4)
            cell_text(meta_tbl.rows[i].cells[1], value, size_pt=9,
                      color=C["textPrimary"])
            meta_tbl.rows[i].cells[1].width = Cm(14)

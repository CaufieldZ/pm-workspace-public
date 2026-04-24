#!/usr/bin/env python3
# PM-Workspace | (c) 2026 CaufieldZ | Apache 2.0 + AI Training Restriction
"""
push_to_confluence_base.py — 推送 PRD docx 到 Confluence Server

两种模式:
    A. 按 space + title upsert(找不到则建新页):
        python3 push_to_confluence_base.py <docx> --title <标题> --space <key> [--parent-id <id>]
    B. 按 pageId 直接覆盖(已知页面,跳过 title 查找):
        python3 push_to_confluence_base.py <docx> --page-id <id> [--title <新标题>]

说明:
    - 凭证自动从 .mcp.json / .mcp-disabled.json 读取(MCP 关闭也能跑,脚本走 REST)
    - 图片提取后作为附件上传,再用 ac:image 宏引用
    - push 前 pre-flight 检查 docx 的 Heading 1/2 计数,为 0 时 warning
      (memory #1/#3 踩坑:gen 脚本 BASE 路径错 → h1/h2 函数失效 → docx 全 Normal
       → Confluence 大纲树失效,不先检查事后很难发现)
    - 依赖: pip install mammoth requests

示例:
    # 模式 A:新建/upsert
    python3 .claude/skills/prd/references/push_to_confluence_base.py \\
        projects/htx-activity-center/deliverables/prd-htx-v1.docx \\
        --title "HTX 活动中心 PRD v1" --space HTX --parent-id 12345678

    # 模式 B:按 pageId 覆盖
    python3 .claude/skills/prd/references/push_to_confluence_base.py \\
        projects/htx-community/deliverables/prd-htx-community-v3.4.docx \\
        --page-id 155652375
"""

import sys, re, json, argparse, requests
from pathlib import Path

try:
    import mammoth
except ImportError:
    sys.exit("缺依赖: pip install mammoth requests")

# ── 从 .mcp.json / .mcp-disabled.json 读取 Confluence 配置 ──────────────────
# toggle-mcp.sh off 会把 server 搬到 .mcp-disabled.json,env 字段保留;
# 脚本走 REST API 不需要 MCP server 在线,凭据在哪都能用。
_ROOT = Path(__file__).resolve().parents[4]  # pm-workspace 根目录

def _load_conf_creds():
    for src in [_ROOT / ".mcp.json", _ROOT / ".mcp-disabled.json"]:
        if not src.exists():
            continue
        env = json.loads(src.read_text()).get("mcpServers", {}).get("confluence", {}).get("env")
        if env and env.get("CONF_BASE_URL") and env.get("CONF_TOKEN"):
            return env["CONF_BASE_URL"].rstrip("/"), env["CONF_TOKEN"]
    sys.exit("找不到 confluence 凭据(.mcp.json 和 .mcp-disabled.json 都没有 env)")

BASE_URL, TOKEN = _load_conf_creds()
_HEADERS = {"Authorization": f"Bearer {TOKEN}"}


# ── REST API 封装 ────────────────────────────────────────────────────────────

def _rest(method: str, path: str, **kwargs):
    url = f"{BASE_URL}/rest/api/{path.lstrip('/')}"
    headers = {**_HEADERS, **kwargs.pop("extra_headers", {})}
    r = requests.request(method, url, headers=headers, **kwargs)
    try:
        r.raise_for_status()
    except requests.HTTPError as e:
        sys.exit(f"Confluence API 错误 [{r.status_code}]: {r.text[:300]}\n{e}")
    return r.json() if r.content else {}


def get_or_create_page(space_key: str, title: str, parent_id: str = None) -> tuple:
    """返回 (page_id, current_version_number)。页面不存在则新建占位页。"""
    resp = _rest("GET", "content", params={
        "spaceKey": space_key, "title": title, "expand": "version"
    })
    results = resp.get("results", [])
    if results:
        p = results[0]
        return p["id"], p["version"]["number"]

    # 新建占位页（先建后填内容，这样附件上传有 pageId）
    payload = {
        "type": "page",
        "title": title,
        "space": {"key": space_key},
        "body": {"storage": {"value": "", "representation": "storage"}},
    }
    if parent_id:
        payload["ancestors"] = [{"id": str(parent_id)}]
    p = _rest("POST", "content",
              json=payload,
              extra_headers={"Content-Type": "application/json"})
    return p["id"], 1


def upload_attachment(page_id: str, filename: str, data: bytes, mime: str = "image/png"):
    """上传图片附件，已存在则更新（upsert）。"""
    base = f"{BASE_URL}/rest/api/content/{page_id}/child/attachment"
    headers = {**_HEADERS, "X-Atlassian-Token": "no-check"}

    # 查是否已有同名附件
    existing = requests.get(base, headers=_HEADERS,
                            params={"filename": filename}).json().get("results", [])
    if existing:
        att_id = existing[0]["id"]
        url = f"{base}/{att_id}/data"
    else:
        url = base

    requests.post(url, headers=headers,
                  files={"file": (filename, data, mime)}).raise_for_status()


# ── docx → Confluence storage format ────────────────────────────────────────

def convert_docx(docx_path: str, page_id: str, img_width: int = 600) -> str:
    """
    1. mammoth 把 docx 转 HTML，图片提取到内存
    2. 图片上传为页面附件
    3. <img src="..."> 替换为 <ac:image> 宏
    返回 Confluence storage format HTML 字符串
    """
    _images = {}   # filename → (bytes, mime)
    _idx = [0]

    def handle_image(image):
        with image.open() as f:
            data = f.read()
        mime = getattr(image, "content_type", None) or "image/png"
        ext  = mime.split("/")[-1].split("+")[0]   # image/svg+xml → svg
        _idx[0] += 1
        name = f"prd-img-{_idx[0]}.{ext}"
        _images[name] = (data, mime)
        return {"src": f"__ATTACH__{name}"}

    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(
            f,
            convert_image=mammoth.images.img_element(handle_image)
        )

    if result.messages:
        for msg in result.messages:
            print(f"  ⚠ mammoth: {msg}")

    html = result.value

    # 上传附件
    for name, (data, mime) in _images.items():
        upload_attachment(page_id, name, data, mime)
        print(f"  ↑ 附件: {name} ({len(data)//1024}KB)")

    # <img src="__ATTACH__xxx"> → <ac:image> 宏
    def to_ac(m):
        src = m.group(1)
        if src.startswith("__ATTACH__"):
            fname = src[len("__ATTACH__"):]
            return f'<ac:image ac:width="{img_width}"><ri:attachment ri:filename="{fname}"/></ac:image>'
        return m.group(0)   # 外链图片保留原样

    html = re.sub(r'<img[^>]+src="([^"]+)"[^>]*/?>', to_ac, html)

    # 恢复段落缩进(mammoth 默认丢失)。
    # fill_cell_blocks 产出两级缩进:一级 "N. xxx"(Cm 0.3),二级 "  - xxx"(Cm 0.9)。
    # 映射到 HTML 的 padding-left,让 Confluence 渲染保留层次。
    def indent_p(m):
        inner = m.group(1)
        # 去掉段落内前置 <strong>/<em>/<br/> 标签,看真实文本首字符
        txt = re.sub(r'<[^>]+>', '', inner).lstrip()
        if re.match(r'^\s*-\s', txt) or inner.startswith(('  -', '\t-')):
            return f'<p style="padding-left:40px;">{inner}</p>'
        if re.match(r'^\d+\.\s', txt):
            return f'<p style="padding-left:20px;">{inner}</p>'
        return m.group(0)
    html = re.sub(r'<p>(.*?)</p>', indent_p, html, flags=re.DOTALL)

    return html


def update_page(page_id: str, title: str, body: str, current_version: int):
    payload = {
        "version": {"number": current_version + 1},
        "title": title,
        "type": "page",
        "body": {"storage": {"value": body, "representation": "storage"}},
    }
    _rest("PUT", f"content/{page_id}",
          json=payload,
          extra_headers={"Content-Type": "application/json"})


def get_page_info(page_id: str) -> dict:
    """按 pageId 直接取当前页面 title + version + spaceKey,跳过 title 查找。
    模式 B 用;当 page 已知时比 get_or_create_page 稳(无同名页歧义)。"""
    info = _rest("GET", f"content/{page_id}", params={"expand": "version,space"})
    return {
        "id": info["id"],
        "title": info["title"],
        "version": info["version"]["number"],
        "space_key": info["space"]["key"],
    }


# ── Pre-flight:docx 结构体检 ─────────────────────────────────────────────────

def preflight_check_headings(docx_path: str) -> tuple:
    """扫 docx Heading 1/2 计数。为 0 时 stderr 打印 warning 但不中止。
    memory #1/#3 踩坑:gen 脚本 BASE 路径错时 h1/h2 函数无声失败,
    所有章节都是 Normal style,推 Confluence 后大纲树失效。
    事前 5 秒检查省事后几小时排查。
    返回 (h1_count, h2_count)。"""
    try:
        from docx import Document
    except ImportError:
        print("  ⚠ 跳过 Heading 体检(python-docx 未安装)", file=sys.stderr)
        return (-1, -1)
    doc = Document(docx_path)
    h1 = sum(1 for p in doc.paragraphs if (p.style and p.style.name == "Heading 1"))
    h2 = sum(1 for p in doc.paragraphs if (p.style and p.style.name == "Heading 2"))
    print(f"🔍 pre-flight: Heading 1={h1}, Heading 2={h2}")
    if h1 == 0 and h2 == 0:
        print(
            "⚠️  docx 里没有任何 Heading 1/2!推到 Confluence 后左侧大纲树会失效。\n"
            "   常见原因:① gen 脚本 sys.path 路径错,h1/h2 函数无声失败(memory #1)\n"
            "            ② 章节全部用 add_paragraph 直接写,没调 h1/h2\n"
            "   先跑 normalize_headings(doc) 归一化或修 gen 脚本,再 push。",
            file=sys.stderr,
        )
    elif h1 == 0:
        print("⚠️  docx 里没有 Heading 1,只有 H2 — Confluence 大纲树会缺顶层。", file=sys.stderr)
    return (h1, h2)


# ── 主流程 ───────────────────────────────────────────────────────────────────

def _build_parser():
    p = argparse.ArgumentParser(
        description="推送 PRD docx 到 Confluence Server(两种模式)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    p.add_argument("docx", help="PRD docx 文件路径")
    p.add_argument("--page-id", dest="page_id", default=None,
                   help="已知 pageId,直接按 ID 覆盖(模式 B);与 --space/--parent-id 互斥")
    p.add_argument("--title", default=None,
                   help="页面标题。模式 A 必填;模式 B 可选(不传则保留 Confluence 现有 title)")
    p.add_argument("--space", dest="space_key", default=None,
                   help="Confluence spaceKey(模式 A 必填)")
    p.add_argument("--parent-id", dest="parent_id", default=None,
                   help="父页面 pageId(模式 A 可选)")
    p.add_argument("--img-width", type=int, default=600,
                   help="图片附件展示宽度(像素),默认 600")
    p.add_argument("--skip-preflight", action="store_true",
                   help="跳过 Heading 体检(不推荐,除非明知 docx 故意无 H1/H2)")
    return p


def _resolve_mode(args):
    """返回 (page_id, title, version, space_key) —— 无论模式 A/B 都归一化为这 4 个字段。"""
    if args.page_id:
        info = get_page_info(args.page_id)
        title = args.title or info["title"]
        return info["id"], title, info["version"], info["space_key"]
    # 模式 A:按 title upsert
    if not args.title or not args.space_key:
        sys.exit("❌ 模式 A 必须传 --title 和 --space(或用 --page-id 走模式 B)")
    page_id, version = get_or_create_page(args.space_key, args.title, args.parent_id)
    return page_id, args.title, version, args.space_key


def main(argv=None):
    # 向后兼容:旧位置参数 `docx title space [parent]` 仍能工作
    if argv is None:
        argv = sys.argv[1:]
    if argv and not argv[0].startswith("-") and len(argv) >= 3 and not any(
        a.startswith("--") for a in argv[:4]
    ):
        # 位置参数模式:[docx, title, space, parent_id?]
        legacy = ["--title", argv[1], "--space", argv[2]]
        if len(argv) >= 4:
            legacy += ["--parent-id", argv[3]]
        argv = [argv[0]] + legacy

    args = _build_parser().parse_args(argv)

    if not Path(args.docx).exists():
        sys.exit(f"❌ 文件不存在: {args.docx}")

    if not args.skip_preflight:
        preflight_check_headings(args.docx)

    print(f"📄 转换: {args.docx}")
    page_id, title, version, space_key = _resolve_mode(args)
    print(f"📑 页面 ID: {page_id}  标题: {title}  space: {space_key}  当前版本: {version}")

    body = convert_docx(args.docx, page_id, img_width=args.img_width)

    print("📤 更新页面内容...")
    update_page(page_id, title, body, version)

    page_url = f"{BASE_URL}/pages/viewpage.action?pageId={page_id}"
    print(f"✅ 完成: {page_url}")


if __name__ == "__main__":
    main()

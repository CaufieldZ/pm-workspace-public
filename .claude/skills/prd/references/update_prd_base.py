#!/usr/bin/env python3
# PM-Workspace | (c) 2026 CaufieldZ | Apache 2.0 + AI Training Restriction
"""
update_prd_base.py — PRD 升版/更新的通用辅助函数

与 gen_prd_base.py（新建 PRD）互补，本文件处理「修改已有 docx」场景。

用法（从 projects/{项目}/scripts/ 执行）：
    import sys, os
    _ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '../../..'))
    sys.path.insert(0, os.path.join(_ROOT, '.claude/skills/prd/references'))
    from update_prd_base import *
"""

from docx.shared import Cm, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── fix_dpi ────────────────────────────────────────────────────────────────────

def fix_dpi(png_path: str, dpi: int = 144) -> str:
    """
    修正 @2x 截图 DPI 元数据（Playwright deviceScaleFactor=2 输出 72dpi，
    python-docx 按 DPI 推算尺寸会导致截图虚化）。

    原地修改，返回原路径。
    """
    from PIL import Image
    img = Image.open(png_path)
    img.save(png_path, dpi=(dpi, dpi))
    return png_path


# ── 段落级操作 ─────────────────────────────────────────────────────────────────

def replace_para_text(para, new_text: str):
    """
    整体替换段落文字，保留第一个 run 的格式（字号/粗体/颜色）。
    后续 run 清空。
    """
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.add_run(new_text)


def search_replace_para(para, old: str, new: str) -> bool:
    """
    段落内搜索替换（跨 run 拼接后匹配）。
    替换成功返回 True，未命中返回 False。保留第一个 run 的格式。
    """
    full = ''.join(r.text for r in para.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ''
    return True


# ── 单元格级操作 ──────────────────────────────────────────────────────────────

def set_cell_text(cell, text: str):
    """
    清空单元格所有内容，在首段写入纯文本。
    保留首段首 run 的格式；无 run 则新建。
    """
    # 清空所有段落文字
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    # 移除多余段落（只留第一个）
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)
    # 写入
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def set_cell_blocks(cell, blocks, numbered=True):
    """
    清空单元格，按结构化 blocks 重建（title 粗体 + 子条目缩进 + 「（变更）」「（新增）」染色）。
    视觉和 gen_prd_base.scene_table 的右列一致。

    用途：升版/更新 PRD 时，给 Scene 表格右列或后台 table 的复杂内容做结构化替换，
          替代 set_cell_text —— 后者只能单 run 纯文本，没有层次。

    参数：
        cell   — python-docx Table cell
        blocks — list[tuple[str, list[str]]]
                 [(title, [line1, line2, ...]), ...]
                 title 含「（变更）」或「（新增）」会自动拆 run 染成强调色。

    示例：
        set_cell_blocks(cell, [
            ("合格投资者校验（R1/R2）", [
                "进入认购流程前后端双重校验：QI 认证 + 风险匹配",
                "校验通过：展示绿色 Badge；失败：拦截 + 引导认证",
            ]),
            ("基金信息展示（变更）", [
                "产品全称 + 风险等级标签",
                "单位净值 + 成立以来收益率",
            ]),
        ])
    """
    # 清空 cell 所有现有内容（文字 + 图片）
    for para in cell.paragraphs:
        para.clear()
    # 仅保留第一个段落壳
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)
    # 调用 gen_prd_base 的公共填充逻辑（避免 DRY 违反）
    import os, sys
    _DIR = os.path.dirname(os.path.abspath(__file__))
    sys.path.insert(0, _DIR)
    from gen_prd_base import fill_cell_blocks
    fill_cell_blocks(cell, blocks, numbered=numbered)


def replace_cell_image(cell, img_path: str, width_cm: float = 7.0):
    """
    清空单元格（文字+旧图），插入新截图。
    内部强制调用 fix_dpi()，确保 docx 里不虚化。

    参数：
        cell       — python-docx Table cell
        img_path   — PNG 文件路径（str 或 Path）
        width_cm   — 插入宽度，前端 Scene 建议 7.0，App 建议 5.0
    """
    img_path = str(img_path)
    fix_dpi(img_path)

    # 清空所有段落内容（文字 + drawing）
    for para in cell.paragraphs:
        para.clear()
    # 移除多余段落，只留一个
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)
    # 插入图片
    p = cell.paragraphs[0]
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))


# ── 标题样式归一化 ─────────────────────────────────────────────────────────────

def normalize_headings(doc):
    """
    修旧 docx 的 H1/H2 样式:补 run 级字色/字号/粗体 + H1 补下划线 pBdr。
    旧脚本常用 add_paragraph(style='Heading 1') 直接产段落,run 级属性缺失,
    渲染效果取决于 Word 默认 style(黑无线),和 gen_prd_base.h1()/h2() 产出的标准
    视觉不一致。本函数幂等:已有 run 属性的跳过不覆盖。

    H1: fg=#1A1A2E + 16pt + bold + 段落下边框 #2E75B6
    H2: fg=#2E75B6 + 13pt + bold

    返回修复数 (h1_count, h2_count)。
    """
    import sys, os
    _DIR = os.path.dirname(os.path.abspath(__file__))
    sys.path.insert(0, _DIR)
    from gen_prd_base import C

    h1_fixed = h2_fixed = 0
    for p in doc.paragraphs:
        s = p.style.name if p.style else ""
        if s == "Heading 1":
            _apply_heading_style(p, C["textHeading"], 16, add_border=C["accentBlue"])
            h1_fixed += 1
        elif s == "Heading 2":
            _apply_heading_style(p, C["accentBlue"], 13, add_border=None)
            h2_fixed += 1
    return h1_fixed, h2_fixed


def _apply_heading_style(p, hex_color, size_pt, add_border=None):
    """给 heading 段落强制刷 run 样式,可选追加段落下边框。
    强制覆盖颜色/字号/粗体:老 docx 里即使 H1 已设错颜色(如黑色)也统一刷成标准色。
    """
    for r in p.runs:
        r.font.color.rgb = RGBColor.from_string(hex_color)
        r.font.size = Pt(size_pt)
        r.bold = True
    # 无 run(段落只有 text 没 run)时新建一个
    if not p.runs and p.text:
        text = p.text
        for child in list(p._p):
            if child.tag.endswith('}r') or child.tag.endswith('}t'):
                p._p.remove(child)
        run = p.add_run(text)
        run.font.color.rgb = RGBColor.from_string(hex_color)
        run.font.size = Pt(size_pt)
        run.bold = True
    # 段落下边框(H1 专用)
    if add_border:
        pPr = p._p.get_or_add_pPr()
        if pPr.find(qn('w:pBdr')) is None:
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), add_border)
            pBdr.append(bottom)
            pPr.append(pBdr)


def normalize_punctuation(doc):
    """
    中文标点规范化:中文相邻的半角 , : ( ) 改为全角 ，：（）。
    遵循 soul.md「中文里禁止混半角逗号冒号括号」。
    幂等;保留 run 级样式(bold / color 不变)。

    段落级判定:拼接段落所有 run 的文字后做上下文判断,避免 run 边界把
    「中文<run1 end>:<run2 start>英文」这种场景看成孤立冒号漏改。

    判定规则:
      - 半角 , : ( ) 只要左右任一是中文字符 → 替换全角
      - 排除 URL 场景:前 6 字符含 http / ftp
    代码/URL/纯英文上下文不触发(两侧都没中文)。
    """
    import re
    CJK_RE = re.compile(r'[\u4e00-\u9fff]')
    MAP = {',': '\uff0c', ':': '\uff1a', '(': '\uff08', ')': '\uff09'}

    def process_paragraph(p):
        runs = p.runs
        if not runs: return 0
        # 拼接 + 记录 run 边界
        full = ''.join(r.text or '' for r in runs)
        if not any(c in MAP for c in full): return 0

        out = list(full)
        for i, c in enumerate(full):
            if c not in MAP: continue
            prev = full[i-1] if i > 0 else ''
            nxt  = full[i+1] if i+1 < len(full) else ''
            if not (CJK_RE.match(prev) or CJK_RE.match(nxt)):
                continue
            if c == ':':
                window = full[max(0, i-6):i].lower()
                if 'http' in window or 'ftp' in window:
                    continue
            out[i] = MAP[c]

        new_full = ''.join(out)
        if new_full == full: return 0

        # 按原 run 边界切回写入;替换都是 1:1 字符,偏移不变
        n = 0
        pos = 0
        for r in runs:
            L = len(r.text or '')
            new_slice = new_full[pos:pos+L]
            if new_slice != (r.text or ''):
                r.text = new_slice
                n += sum(1 for a, b in zip(r.text, full[pos:pos+L]) if a != b)
            pos += L
        return n

    total = 0
    for p in doc.paragraphs:
        total += process_paragraph(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total += process_paragraph(p)
    return total

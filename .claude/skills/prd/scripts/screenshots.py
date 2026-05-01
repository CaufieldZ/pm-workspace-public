#!/usr/bin/env python3
"""PRD 截图回填 + 升版断言：Playwright 截图 → 圆角 → 插入 docx。

合并自原 prd_screenshots.py（截图主体）+ update_prd_base.assert_screenshots_fresh
（升版断言）。原 prd_screenshots.py 已退化为本文件的 alias。

依赖：
- core.images → fix_dpi, replace_cell_image, replace_cell_image_keep_title, round_phone_corners
- playwright（按需）

用法：
    python3 screenshots.py --project community/leaderboard
    python3 screenshots.py --project community/base --imap deliverables/imap-foo-v2.html
    python3 screenshots.py --project foo --docx deliverables/PRD_v2.docx
    python3 screenshots.py --project foo --scenes A-1,B-1,C-1
    python3 screenshots.py --project foo --shot-only    # 只截图不插入
    python3 screenshots.py --project foo --insert-only  # 只插入已有截图
"""
import argparse
import os
import re
import sys
from pathlib import Path

from core.images import (
    fix_dpi,
    replace_cell_image, replace_cell_image_keep_title,
    round_phone_corners,
)


ROOT = Path(__file__).resolve().parents[4]  # pm-workspace 根


# ── 升版断言（防止改完 docx 文字忘重拍截图）─────────────────────────────

def assert_screenshots_fresh(docx_path, prototype_html, shot_dir,
                             excluded_subdir_patterns=('archive', '_v', 'deprecated', 'old')):
    """升版硬规则：改完 docx 文字后调用，断言截图比原型新；过期则 raise。

    诊断三角：
      - prototype html 的 mtime（A）= 原型最后改动时间
      - 截图目录下所有 .png 的最旧 mtime（B）= 截图集合的"最弱一环"
      - 如果 A > B → 至少有一张截图比原型旧 → 漏重拍 → raise

    用法（升版脚本末尾，doc.save() 之后）：
        from screenshots import assert_screenshots_fresh
        doc.save(str(TARGET_DOCX))
        assert_screenshots_fresh(
            TARGET_DOCX,
            PROJ / 'deliverables/活动中心_可交互原型_v5.1.html',
            PROJ / 'screenshots/prd',
        )

    历史：连续多次踩坑「docx 文字改完忘重拍截图」→ 加此断言强制开发者补流程。
    """
    docx_path = Path(docx_path)
    prototype_html = Path(prototype_html)
    shot_dir = Path(shot_dir)

    if not prototype_html.exists():
        return  # 原型不存在不强制（可能纯文字升版无原型依赖）
    if not shot_dir.exists():
        return

    proto_t = prototype_html.stat().st_mtime

    pngs = []
    for p in shot_dir.rglob('*.png'):
        if any(pat in str(p.relative_to(shot_dir)) for pat in excluded_subdir_patterns):
            continue
        pngs.append(p)

    if not pngs:
        return

    stale = [p for p in pngs if p.stat().st_mtime < proto_t]
    if not stale:
        return

    msg = [
        f'\n❌ [assert_screenshots_fresh] 截图比原型旧，漏重拍：',
        f'   原型: {prototype_html.name} (mtime={proto_t:.0f})',
        f'   过期截图 ({len(stale)}/{len(pngs)} 张):',
    ]
    for p in stale[:10]:
        delta_h = (proto_t - p.stat().st_mtime) / 3600
        msg.append(f'     - {p.relative_to(shot_dir)} (旧 {delta_h:.1f}h)')
    if len(stale) > 10:
        msg.append(f'     ... 共 {len(stale)} 张')
    msg.append(f'   修复：跑 screenshot_*.py 重拍 + insert_*.py 回填到 docx，再调本断言')
    raise AssertionError('\n'.join(msg))


# ── 工具函数 ──────────────────────────────────────────────────────────────

def find_latest(directory: Path, glob: str):
    files = sorted(directory.glob(glob), key=lambda p: p.stat().st_mtime, reverse=True)
    return files[0] if files else None


def discover_scenes_from_html(html_path: Path):
    """从 imap/proto HTML 解析 scene 列表，返回 [(scene_id, device_type), ...]。
    device_type: 'phone' | 'web'
    通过检查 scene 容器内是否含 .phone 或 .webframe 判断。"""
    text = html_path.read_text(encoding="utf-8")
    scene_ids = re.findall(r'<div[^>]+class="[^"]*fade-section[^"]*"[^>]+id="(scene-[^"]+)"', text)
    if not scene_ids:
        scene_ids = re.findall(r'id="(scene-[a-z0-9-]+)"', text)

    results = []
    for sid in scene_ids:
        pattern = rf'id="{re.escape(sid)}".*?(?=id="scene-|</body>)'
        block = re.search(pattern, text, re.DOTALL)
        device = "phone"
        if block:
            block_text = block.group(0)
            if 'class="webframe' in block_text or "class='webframe" in block_text:
                device = "web"
            elif 'class="phone' not in block_text and "class='phone" not in block_text:
                device = "web"
        results.append((sid, device))
    return results


# ── 截图 ──────────────────────────────────────────────────────────────────

def take_screenshots(html_path: Path, out_dir: Path, scenes=None):
    """打开 HTML，对每个 scene 截设备框，返回 {scene_id: png_path}。"""
    from playwright.sync_api import sync_playwright

    out_dir.mkdir(parents=True, exist_ok=True)
    url = html_path.as_uri()

    if scenes is None:
        scenes = discover_scenes_from_html(html_path)
    if not scenes:
        sys.exit(f"错误：在 {html_path.name} 里找不到 scene-* 容器")

    shots = {}

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True)
        page = browser.new_page(
            viewport={"width": 1440, "height": 900},
            device_scale_factor=2,
        )
        page.goto(url)
        page.wait_for_load_state("networkidle")

        page.evaluate("""
            document.querySelectorAll('.fade-section').forEach(el => el.classList.add('visible'));
            document.querySelectorAll(
              '.anno, .anno-n, .ann-card, .ann-tag, .aw, .flow-note, .side-nav, .scene-nav'
            ).forEach(el => { el.style.display = 'none'; });
        """)
        page.wait_for_timeout(600)

        for scene_id, device_type in scenes:
            scene_loc = page.locator(f"#{scene_id}").first
            if not scene_loc.is_visible():
                print(f"  ⊘ {scene_id}: 容器不可见，跳过")
                continue
            scene_loc.scroll_into_view_if_needed()
            page.wait_for_timeout(200)

            selector = ".phone" if device_type == "phone" else ".webframe"
            device_loc = scene_loc.locator(selector).first
            if not device_loc.is_visible():
                alt = ".webframe" if device_type == "phone" else ".phone"
                device_loc = scene_loc.locator(alt).first
                if device_loc.is_visible():
                    device_type = "web" if device_type == "phone" else "phone"
                else:
                    print(f"  ⊘ {scene_id}: 找不到 .phone / .webframe")
                    continue

            fname = scene_id.removeprefix("scene-") + ".png"
            out_path = out_dir / fname
            device_loc.screenshot(path=str(out_path), omit_background=True)

            if device_type == "phone":
                round_phone_corners(str(out_path))

            shots[scene_id] = out_path
            print(f"  截图 {scene_id} ({device_type}) → {fname}")

        browser.close()

    return shots


# ── 插入 docx ─────────────────────────────────────────────────────────────

_SCENE_ID_RE = re.compile(r'\b([A-Z])-(\d+[a-z]?)\b')
_SCENE_TITLE_RE = re.compile(r'(📱\s*[A-Z]-\d+[a-z]?\s*·\s*[^\n]+)')


def _normalize_scene_id(raw: str) -> str:
    """把 scene-a-1 / A-1 / a-1 统一成 A-1 大写形式。"""
    raw = raw.removeprefix("scene-")
    parts = raw.split("-", 1)
    if len(parts) == 2:
        return f"{parts[0].upper()}-{parts[1]}"
    return raw.upper()


def _cell_scene_ids(cell):
    text = cell.text.strip()
    return {f"{m.group(1)}-{m.group(2)}" for m in _SCENE_ID_RE.finditer(text)}


def insert_into_docx(docx_path: Path, shots,
                     width_phone_cm: float = 5.0, width_web_cm: float = 7.0):
    """把截图插入 PRD docx，按场景编号匹配表格首行首列。

    自动检测 cell 首段是否含 scene_table 标识（📱 + 编号），有则保留场景标题段
    （memory feedback_prd_iter_screenshot_pitfalls.md 坑 2：默认 replace_cell_image
    会清掉 scene_table 写在 cell 首段的「📱 X-N · 名称」标识）。
    """
    from docx import Document
    doc = Document(str(docx_path))

    shots_norm = {_normalize_scene_id(k): v for k, v in shots.items()}
    inserted = set()

    for table in doc.tables:
        if not table.rows:
            continue
        cell = table.rows[0].cells[0]
        cell_ids = _cell_scene_ids(cell)
        for sid in cell_ids:
            if sid in shots_norm and sid not in inserted:
                img_path = shots_norm[sid]
                is_phone = "phone" in str(img_path) or _is_portrait(img_path)
                width = width_phone_cm if is_phone else width_web_cm
                first_para_text = cell.paragraphs[0].text.strip() if cell.paragraphs else ""
                title_match = _SCENE_TITLE_RE.match(first_para_text)
                if title_match:
                    replace_cell_image_keep_title(cell, title_match.group(1), str(img_path), width_cm=width)
                else:
                    replace_cell_image(cell, str(img_path), width_cm=width)
                print(f"  插入 {sid} → Table ({table.rows[0].cells[0].text[:20].strip()!r})")
                inserted.add(sid)
                break

    doc.save(str(docx_path))
    print(f"\n  已保存: {docx_path.name}")
    missing = set(shots_norm) - inserted
    if missing:
        print(f"  未匹配场景（docx 里没找到对应表格）: {', '.join(sorted(missing))}")
    return inserted


def _is_portrait(png_path: Path) -> bool:
    """粗判断是否竖屏（phone），高 > 宽时认为是。"""
    try:
        from PIL import Image
        w, h = Image.open(str(png_path)).size
        return h > w
    except Exception:
        return False


# ── 主流程 ────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="PRD 截图回填通用脚本")
    parser.add_argument("--project", "-p", required=True, help="项目名（projects/ 下目录名）")
    parser.add_argument("--imap", help="imap/proto HTML 路径（默认：自动找 deliverables/imap-*.html）")
    parser.add_argument("--docx", help="PRD docx 路径（默认：自动找 deliverables/*PRD*.docx）")
    parser.add_argument("--scenes", help="限定场景列表，逗号分隔，如 A-1,B-2（默认：全部）")
    parser.add_argument("--shot-only", action="store_true", help="只截图，不插入 docx")
    parser.add_argument("--insert-only", action="store_true", help="只插入已有截图，跳过截图步骤")
    parser.add_argument("--width-phone", type=float, default=5.0, help="phone 截图宽度 cm（默认 5.0）")
    parser.add_argument("--width-web", type=float, default=7.0, help="web 截图宽度 cm（默认 7.0）")
    args = parser.parse_args()

    proj_dir = ROOT / "projects" / args.project
    if not proj_dir.is_dir():
        sys.exit(f"错误：项目目录不存在 {proj_dir}")

    deliverables = proj_dir / "deliverables"
    shot_dir = proj_dir / "screenshots" / "prd"

    html_path = None
    if not args.insert_only:
        if args.imap:
            html_path = Path(args.imap) if Path(args.imap).is_absolute() else proj_dir / args.imap
        else:
            html_path = find_latest(deliverables, "imap-*.html")
            if not html_path:
                html_path = find_latest(deliverables, "proto-*.html")
            if not html_path:
                html_path = find_latest(deliverables, "*.html")
        if not html_path or not html_path.exists():
            sys.exit(f"错误：找不到 imap/proto HTML，请用 --imap 指定路径")
        print(f"HTML 源: {html_path.relative_to(ROOT)}")

    docx_path = None
    if not args.shot_only:
        if args.docx:
            docx_path = Path(args.docx) if Path(args.docx).is_absolute() else proj_dir / args.docx
        else:
            docx_path = find_latest(deliverables, "*PRD*.docx")
            if not docx_path:
                docx_path = find_latest(deliverables, "*prd*.docx")
        if not docx_path or not docx_path.exists():
            sys.exit(f"错误：找不到 PRD docx，请用 --docx 指定路径")
        print(f"PRD:   {docx_path.relative_to(ROOT)}")

    scene_filter = None
    if args.scenes:
        scene_filter = {_normalize_scene_id(s.strip()) for s in args.scenes.split(",")}

    shots = {}
    if not args.insert_only:
        all_scenes = discover_scenes_from_html(html_path)
        if scene_filter:
            all_scenes = [(sid, dt) for sid, dt in all_scenes
                          if _normalize_scene_id(sid) in scene_filter]
        print(f"\n截图 {len(all_scenes)} 个 Scene...")
        shots = take_screenshots(html_path, shot_dir, scenes=all_scenes)
    else:
        if not shot_dir.exists():
            sys.exit(f"错误：截图目录不存在 {shot_dir}，请先运行截图步骤")
        for png in shot_dir.glob("*.png"):
            sid = "scene-" + png.stem
            shots[sid] = png
        if scene_filter:
            shots = {k: v for k, v in shots.items()
                     if _normalize_scene_id(k) in scene_filter}
        print(f"加载已有截图 {len(shots)} 张...")

    if not shots:
        sys.exit("没有可用截图，退出")

    if args.shot_only:
        print(f"\n截图完成，共 {len(shots)} 张，存于 {shot_dir.relative_to(ROOT)}")
        return

    print(f"\n插入 docx...")
    inserted = insert_into_docx(docx_path, shots,
                                width_phone_cm=args.width_phone,
                                width_web_cm=args.width_web)
    print(f"\n完成：{len(inserted)}/{len(shots)} 个场景已插入")


if __name__ == "__main__":
    main()

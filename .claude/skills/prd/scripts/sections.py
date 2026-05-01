#!/usr/bin/env python3
"""PRD 产品思维章节 helper（4 个高 ROI 要素）。

按 Plan Layer 1 沉淀，把顶尖大厂 PRD 的 4 个高 ROI 章节标准化：

1. north_star_section  ——  1.2 北极星 + 配套 + 反向（Guardrail）+ 非目标（全档位必填）
2. non_goals_section   ——  1.5 Non-goals（标准/完整必填，极简可省）
3. open_questions_table —  6.x Open Questions（全档位必填，无未决项写「无」）
4. risks_table         ——  7.x Risks & Mitigations（完整必填，标准选填）

使用：

    from sections import north_star_section, non_goals_section, \\
        open_questions_table, risks_table

    north_star_section(
        doc,
        north_star="活动中心 DAU 12K → 25K（3 个月内）",
        supporting=["首页 CTR ≥ 8%", "活动报名率 ≥ 15%"],
        guardrail=["首页 Crash < 0.3‰", "客诉率 ≤ 现状 +10%"],
        non_objectives=["不优化加载性能", "不改 KYC 流程"],
    )

    non_goals_section(doc, items=[
        ("不优化首屏加载性能", "现状 1.8s 满足 SLA"),
    ])
    # 空列表 → 自动生成「本期无 Non-goals」占位
    non_goals_section(doc, items=[])

    open_questions_table(doc, rows=[
        ("是否开放白名单 API", "影响外部接入排期", "等技术方案", "@张三", "2026-05-15"),
    ])

    risks_table(doc, rows=[
        ("业务系统数据延迟 > 5min", "中", "高",
         "业务方提供 SLA + 本端兜底缓存", "监控告警 + 切人工配置"),
    ])

Layer 2 拆 core/ 后，本文件依赖路径不变（gen_prd_base.py 退化为 re-export 壳，
para_run / C / make_table / h3 / bullet 等 API 依旧可从 gen_prd_base 拿到）。

章节锚点常量（与 humanize.scan_human_voice tier 体检配套）：

    SECTION_TITLES = {
        'non_goals':       '1.5 Non-goals',
        'alternatives':    '1.6 方案选型说明',
        'open_questions':  'Open Questions',
        'risks':           'Risks & Mitigations',
    }
"""
from docx.shared import Cm, Pt

# 章节锚点常量 —— save_prd 体检用同一份字符串扫描 docx，确保函数生成的标题
# 和体检 grep 的关键词永远一致（避免双轨漂移）
SECTION_TITLES = {
    'non_goals':      '1.5 Non-goals',
    'alternatives':   '1.6 方案选型说明',
    'open_questions': 'Open Questions',
    'risks':          'Risks & Mitigations',
    'assumptions':    '关键假设清单',
}

# Guardrail 关键词（save_prd 体检扫 1.2 是否含反向指标）
GUARDRAIL_KEYWORDS = ('反向指标', 'Guardrail', '不能恶化', '反向')

EMPTY_OPEN_QUESTIONS_PLACEHOLDER = '本期无未决项'
EMPTY_NON_GOALS_PLACEHOLDER = '本期无 Non-goals（本期所有需求都做）'
EMPTY_ASSUMPTIONS_PLACEHOLDER = '本期无关键假设'


def _import_base():
    """惰性 import gen_prd_base 的公共 API。

    放在函数内是为了允许 Layer 2 重构时 gen_prd_base 退化为 re-export 壳，
    不引入循环依赖。
    """
    import os, sys
    _DIR = os.path.dirname(os.path.abspath(__file__))
    if _DIR not in sys.path:
        sys.path.insert(0, _DIR)
    import gen_prd_base
    return gen_prd_base


# ── 1.2 北极星 / 配套 / 反向 / 非目标（全档位必填） ────────────────────────

def north_star_section(doc, *, north_star: str,
                       supporting=None, guardrail=None, non_objectives=None):
    """生成 1.2 核心目标 4 段：北极星 + 配套 + 反向（Guardrail）+ 非目标。

    四段强制写齐（除 north_star 外其他可传空列表，但建议都填），逼 PM 同时想清楚
    「要追求什么」「不能毁掉什么」「不追求什么」。

    方法论见 ../references/metrics-framework.md（北极星三段式 / 配套漏斗推导 /
    Guardrail 默认三件套 / 非目标边界声明 + 自检 8 项 + 反例对照）。本 helper
    只做渲染，不做语义校验，方法论靠 PM 写代码前内化。

    Args:
        doc: python-docx Document
        north_star: 单一最关键指标，含现状值 → 目标值 + 时间窗
        supporting: 配套指标列表（2-3 个），支撑北极星的中间指标
        guardrail: 反向指标列表，本期不能恶化的（**必填，至少 1 项**）
        non_objectives: 非目标列表，本期明确不追求的方向

    生成结构（H3 + bullet）：

        1.2.1 北极星指标
            • [north_star]
        1.2.2 配套指标
            • supporting[0]
            • supporting[1]
        1.2.3 反向指标（Guardrail）
            • guardrail[0]
            • guardrail[1]
        1.2.4 非目标
            • non_objectives[0]
    """
    base = _import_base()
    h3, bullet, add_p, C = base.h3, base.bullet, base.add_p, base.C

    # 1.2.1 北极星
    h3(doc, '1.2.1 北极星指标')
    bullet(doc, north_star)

    # 1.2.2 配套指标
    h3(doc, '1.2.2 配套指标')
    if supporting:
        for s in supporting:
            bullet(doc, s)
    else:
        add_p(doc, '（无配套指标）', size_pt=10, italic=True, color=C['textMuted'])

    # 1.2.3 反向指标 —— 用关键词命中体检（含「反向指标」「Guardrail」字样）
    h3(doc, '1.2.3 反向指标（Guardrail · 不能恶化）')
    if guardrail:
        for g in guardrail:
            bullet(doc, g)
    else:
        # 空 guardrail 仍写一行占位文本，体检 warn 但不阻断
        add_p(doc, '（待 PM 补充：本期不能恶化的指标，如 Crash / 客诉 / 留存）',
              size_pt=10, italic=True, color=C['textMuted'])

    # 1.2.4 非目标
    h3(doc, '1.2.4 非目标')
    if non_objectives:
        for n in non_objectives:
            bullet(doc, n)
    else:
        add_p(doc, '（无明确非目标）', size_pt=10, italic=True, color=C['textMuted'])


# ── 1.5 Non-goals（标准/完整必填） ──────────────────────────────────────────

def non_goals_section(doc, items=None):
    """生成 1.5 Non-goals 章节（H2 + 2 列表格 或 占位）。

    写「不做什么」比写「做什么」更难，强迫 PM 跟自己拉锯一遍边界。
    无 Non-goals 时显式生成占位文本，不留空。

    Args:
        doc: python-docx Document
        items: list of (range, reason) tuples，None / 空 → 生成占位

    生成结构：

        ## 1.5 Non-goals
        | 范围 | 原因 |
        |------|------|
        | ... | ... |

    或空列表时：

        ## 1.5 Non-goals
        本期无 Non-goals（本期所有需求都做）
    """
    base = _import_base()
    h2, add_p, make_table, C = base.h2, base.add_p, base.make_table, base.C

    h2(doc, SECTION_TITLES['non_goals'])

    if not items:
        add_p(doc, EMPTY_NON_GOALS_PLACEHOLDER, size_pt=10,
              italic=True, color=C['textSecondary'])
        return

    rows = [[scope, reason] for scope, reason in items]
    make_table(doc, headers=['范围', '原因'],
               rows_data=rows, col_widths_cm=[8, 13])


# ── 1.6 方案选型说明（仅完整档） ─────────────────────────────────────────────

def alternatives_table(doc, rows=None):
    """生成 1.6 方案选型说明（仅完整档触发，2 选 1+ 时必填）。

    回答「为什么选 A 不选 B/C」。让评审看到推理链而非结论。多方案讨论流水
    归 context.md 第 7 章，本章只写最终态决策矩阵。

    Args:
        doc: python-docx Document
        rows: list of (方案, 描述, 优势, 劣势, 选弃) tuples

    生成结构：

        ## 1.6 方案选型说明
        | 方案 | 描述 | 优势 | 劣势 | 选/弃 |
        |------|------|------|------|-------|
        | A 当前 | ... | ... | ... | ✓ |
        | B 替代 | ... | ... | ... | ✗（理由）|
    """
    base = _import_base()
    h2, add_p, make_table, C = base.h2, base.add_p, base.make_table, base.C

    h2(doc, SECTION_TITLES['alternatives'])

    if not rows:
        add_p(doc, '（本期单方案直出，无替代方案需说明）', size_pt=10,
              italic=True, color=C['textSecondary'])
        return

    table_rows = [[plan, desc, pros, cons, choice]
                  for plan, desc, pros, cons, choice in rows]
    make_table(doc,
               headers=['方案', '描述', '优势', '劣势', '选/弃'],
               rows_data=table_rows,
               col_widths_cm=[2.5, 6, 4.5, 4.5, 3.5])


# ── 6.x Open Questions（全档位必填，无未决项写「无」） ─────────────────────

def open_questions_table(doc, rows=None):
    """生成 Open Questions 表（H3 + 表格 或「无未决项」占位）。

    未决项暴露出来 + 责任人 + deadline，评审会有的放矢，不在会上现场惊讶。
    无未决项时显式写「本期无未决项」，本节不能省略（体检会扫）。

    Args:
        doc: python-docx Document
        rows: list of (议题, 影响, 待决, 责任人, deadline) tuples

    生成结构：

        ### 6.X Open Questions
        | 议题 | 影响 | 待决 | 责任人 | Deadline |
        |------|------|------|--------|----------|

    或空列表时：

        ### 6.X Open Questions
        **本期无未决项**
    """
    base = _import_base()
    h3, add_p, make_table, C = base.h3, base.add_p, base.make_table, base.C

    h3(doc, f'6.X {SECTION_TITLES["open_questions"]}')

    if not rows:
        add_p(doc, EMPTY_OPEN_QUESTIONS_PLACEHOLDER, size_pt=11, bold=True,
              color=C['textHeading'])
        return

    table_rows = [list(r) for r in rows]
    make_table(doc,
               headers=['议题', '影响', '待决', '责任人', 'Deadline'],
               rows_data=table_rows,
               col_widths_cm=[5, 4, 5, 3.5, 3.5])


# ── 6.x 关键假设清单（六列格式，承载项目级假设追踪）─────────────────────────

def assumptions_validation_table(doc, rows=None, heading='6.X 关键假设清单'):
    """生成关键假设清单六列表（H3 + 表格 或「无关键假设」占位）。

    项目级假设追踪：「我们赌什么 → 怎么验 → 何时验 → 错了如何回退 → 现状」
    跟 PRD 6.x 主表 if-then 业务规则不同：业务规则是已确定的，假设是未确定的赌注。
    无假设时显式写「本期无关键假设」，本节不能省略。

    源：context.md 第 6 章子章「关键假设」（context-template 模板 6.6 位置）。
    PRD 用本 helper 投影到 docx，方便评审时一并审视项目假设面。

    Args:
        doc: python-docx Document
        rows: list of (议题, 假设内容, 验证方法, 验证时机, 证伪条件, 当前状态) tuples
              当前状态枚举：未验证 / 验证中 / 已验证 / 已证伪
        heading: 章节标题（默认 '6.X 关键假设清单'，PRD 嵌入位置由调用方决定）

    生成结构：

        ### 6.X 关键假设清单
        | 议题 | 假设内容 | 验证方法 | 验证时机 | 证伪条件 | 当前状态 |
        |------|---------|---------|---------|---------|---------|

    或空列表时：

        ### 6.X 关键假设清单
        **本期无关键假设**
    """
    base = _import_base()
    h3, add_p, make_table, C = base.h3, base.add_p, base.make_table, base.C

    h3(doc, heading)

    if not rows:
        add_p(doc, EMPTY_ASSUMPTIONS_PLACEHOLDER, size_pt=11, bold=True,
              color=C['textHeading'])
        return

    table_rows = [list(r) for r in rows]
    make_table(doc,
               headers=['议题', '假设内容', '验证方法', '验证时机', '证伪条件', '当前状态'],
               rows_data=table_rows,
               col_widths_cm=[3, 4.5, 3.5, 2.5, 4, 2.5])


# ── 7.x Risks & Mitigations（完整必填，标准选填） ──────────────────────────

def risks_table(doc, rows=None):
    """生成 Risks & Mitigations 表（H3 + 表格 或占位）。

    技术 + 业务 + 合规风险全列。完整档触发条件（跨系统 / 资金）通常意味着
    风险面更大，必填；标准档选填（无显著风险时本节可省）。

    Args:
        doc: python-docx Document
        rows: list of (风险, 概率, 影响, 缓解, 触发响应) tuples
              概率 / 影响建议取值：低 / 中 / 高

    生成结构：

        ### 7.X Risks & Mitigations
        | 风险 | 概率 | 影响 | 缓解措施 | 触发响应 |
        |------|------|------|---------|---------|
    """
    base = _import_base()
    h3, add_p, make_table, C = base.h3, base.add_p, base.make_table, base.C

    h3(doc, f'7.X {SECTION_TITLES["risks"]}')

    if not rows:
        add_p(doc, '（本期无显著风险）', size_pt=10,
              italic=True, color=C['textSecondary'])
        return

    table_rows = [list(r) for r in rows]
    make_table(doc,
               headers=['风险', '概率', '影响', '缓解措施', '触发响应'],
               rows_data=table_rows,
               col_widths_cm=[5, 1.8, 1.8, 6, 6.4])


# ── 档位感知产品思维完整性体检（save_prd 内部调用，warn 不阻断）─────────────
#
# 目的：写完 docx 自动 stderr 提醒「PRD 是不是漏写了反向指标 / Non-goals /
#       Open Questions / Risks」。warn 级（不阻断），逼 PM 自查。
#
# 档位行为差异：
#   mini      ：不体检（极简档不强制 4 大要素）
#   standard  ：必查 Guardrail + Non-goals + Open Questions
#   full      ：在 standard 之上加 Risks + Alternatives 必查
#
# 与 humanize.scan_human_voice 互补：humanize 扫「写得像不像 AI」，本扫描
# 扫「写得是否完整」。两条扫描在 save_prd 内并列，串行调用。

import sys


def scan_completeness(doc, tier: str = 'standard') -> dict:
    """扫描 PRD docx 的产品思维完整性（4 大要素是否齐备）。

    Args:
        doc: python-docx Document
        tier: 'mini' / 'standard' / 'full'

    Returns:
        {
            'tier':                 str,
            'missing_guardrail':    bool,  # 1.2 缺反向指标关键词
            'missing_non_goals':    bool,  # 1.5 章节不存在（standard / full）
            'missing_open_questions': bool, # 6.x Open Questions 章不存在
            'missing_risks':        bool,  # 7.x Risks 章不存在（仅 full）
            'missing_alternatives': bool,  # 1.6 方案选型不存在（仅 full）
        }
    """
    if tier not in ('mini', 'standard', 'full'):
        tier = 'standard'

    result = {
        'tier': tier,
        'missing_guardrail': False,
        'missing_non_goals': False,
        'missing_open_questions': False,
        'missing_risks': False,
        'missing_alternatives': False,
    }

    if tier == 'mini':
        return result  # 极简档不体检

    # 拼接全文（段落 + 表格 cell），扫关键词
    chunks = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                chunks.append(cell.text)
    full_text = '\n'.join(chunks)

    # Guardrail：扫 1.2 是否含反向指标关键词
    if not any(kw in full_text for kw in GUARDRAIL_KEYWORDS):
        result['missing_guardrail'] = True

    # Non-goals：扫 1.5 章节标题（标准 + 完整必填）
    if SECTION_TITLES['non_goals'] not in full_text and 'Non-goals' not in full_text:
        result['missing_non_goals'] = True

    # Open Questions：扫章节标题或占位文字（全档位必填，含「无」也算）
    if (SECTION_TITLES['open_questions'] not in full_text
            and EMPTY_OPEN_QUESTIONS_PLACEHOLDER not in full_text):
        result['missing_open_questions'] = True

    if tier == 'full':
        # Alternatives：1.6 方案选型（完整档必填）
        if SECTION_TITLES['alternatives'] not in full_text and 'Alternatives' not in full_text:
            result['missing_alternatives'] = True
        # Risks：7.x Risks & Mitigations（完整档必填）
        if SECTION_TITLES['risks'] not in full_text and 'Risks' not in full_text:
            result['missing_risks'] = True

    return result


def report_completeness(result: dict, file=sys.stderr) -> int:
    """打印体检结果到 stderr。warn 级，永远返回 0（不阻断 save_prd）。

    Args:
        result: scan_completeness 返回的字典
        file: 输出流，默认 stderr

    Returns:
        warning 数（0 = 全通过）
    """
    tier = result.get('tier', 'standard')
    if tier == 'mini':
        return 0

    warnings = []
    if result['missing_guardrail']:
        warnings.append(
            f"1.2 缺反向指标（Guardrail）—— PRD 必须同时讲「不能恶化什么」。"
            f"用 sections.north_star_section(doc, north_star=..., guardrail=[...]) 生成"
        )
    if result['missing_non_goals']:
        warnings.append(
            f"1.5 Non-goals 章节缺失 —— 写「不做什么」比写「做什么」更难，"
            f"逼 PM 想清楚边界。用 sections.non_goals_section(doc, items=[...])，"
            f"无 Non-goals 时传 items=[] 自动占位"
        )
    if result['missing_open_questions']:
        warnings.append(
            f"6.x Open Questions 章节缺失 —— 未决项暴露出来评审会有的放矢。"
            f"用 sections.open_questions_table(doc, rows=[...])，无未决项传 rows=[] 自动占位"
        )
    if result['missing_alternatives']:
        warnings.append(
            f"1.6 方案选型说明缺失（完整档必填）—— "
            f"用 sections.alternatives_table(doc, rows=[...])"
        )
    if result['missing_risks']:
        warnings.append(
            f"7.x Risks & Mitigations 章节缺失（完整档必填）—— "
            f"用 sections.risks_table(doc, rows=[...])"
        )

    if warnings:
        print(f"\n⚠️  PRD 产品思维完整性体检（档位：{tier}）", file=file)
        for w in warnings:
            print(f"   - {w}", file=file)
        print(f"   规范：.claude/skills/prd/references/prd-template.md", file=file)
    return len(warnings)


# ── 第 2 章三节：用户旅程主线 / 端角色拆分 / 跨端时序 ────────────────────────
# 详细叙事投影规则见 references/prd-storytelling.md 第 2 章。
# 渲染走 mermaid_screenshots.render_mermaid（PRD docx 嵌图专用，flowchart skill
# 不服务 PRD docx，参 .claude/skills/flowchart/SKILL.md 边界）。

CROSS_END_PLACEHOLDER = '本项目仅 1 端，无跨端时序'


def _embed_mermaid(doc, mermaid_src, shot_dir, png_name,
                   picture_width_cm=15.5):
    """内部 helper：渲染 mermaid 源码到 PNG 嵌入当前 doc。

    替代原 _embed_flowchart（X6/dagre）。v5.3 视觉反馈：X6 嵌入 docx 字号偏小、
    留白巨大无法阅读 → 改 mermaid LR 横排 + Claude Design 主题。flowchart skill
    退回非 PRD 场景专用（IMAP / arch / PPT 横屏 HTML），PRD 嵌图全走本路径。

    Args:
        doc: python-docx Document
        mermaid_src: mermaid 源码字符串（含 flowchart LR + classDef + 节点 + 边）
        shot_dir: PNG 输出目录
        png_name: PNG 文件名（不含扩展名）。helper 一对一独立 PNG，避免 v5.2/v5.1
                  scope='first' 写死导致 cross_end 嵌错图的 bug
        picture_width_cm: docx 嵌图宽度，默认 15.5 cm

    Returns:
        bool 是否成功嵌入
    """
    if not mermaid_src:
        return False
    from docx.shared import Cm as _Cm
    from pathlib import Path as _Path
    import sys as _sys
    import os as _os
    _DIR = _os.path.dirname(_os.path.abspath(__file__))
    if _DIR not in _sys.path:
        _sys.path.insert(0, _DIR)
    from mermaid_screenshots import render_mermaid

    if shot_dir is None:
        shot_dir = _Path('.')
    shot_dir = _Path(shot_dir)
    out_path = shot_dir / f'{png_name}.png'
    try:
        render_mermaid(mermaid_src, out_path)
    except Exception as e:
        import sys as _s
        print(f'⚠️  mermaid 渲染失败 ({png_name}): {e}', file=_s.stderr)
        return False
    doc.add_picture(str(out_path), width=_Cm(picture_width_cm))
    return True


def journey_main_section(doc, *, narrative: str, mermaid_src=None,
                         shot_dir=None):
    """生成 2.1 用户旅程主线（一段话 + 主线图）。

    narrative 跟 context.md 第 4 章「场景叙事顺序」逐字对齐。涉及 ≥ 2 端项目时
    必须传 mermaid_src（mermaid LR 横排源码字符串，参 mermaid_screenshots.py
    HTML_TEMPLATE 主题约定）。

    v5.3 后从 X6/dagre 切换到 mermaid——flowchart skill 退回非 PRD 场景专用。
    旧参数名 flowchart_html 已停用。

    Args:
        doc: python-docx Document
        narrative: 一段话主线（背景 → 冲突 → 关键路径 → 收束）
        mermaid_src: mermaid 源码字符串（含 flowchart LR + classDef + 节点 + 边）
        shot_dir: PNG 输出目录，默认 projects/{项目}/screenshots/prd/mermaid/

    生成结构：

        ## 2.1 用户旅程主线
        [narrative 段落]
        [主线图 PNG，宽 15.5cm]
    """
    base = _import_base()
    h2, add_p, C = base.h2, base.add_p, base.C

    h2(doc, '2.1 用户旅程主线')
    add_p(doc, narrative)

    embedded = _embed_mermaid(doc, mermaid_src, shot_dir, png_name='journey-main')
    if not embedded and mermaid_src:
        add_p(doc, '⚠️ 主线图渲染失败，详见 stderr',
              size_pt=10, italic=True, color=C['textMuted'])


def end_role_split_section(doc, *, role_summaries, scene_index):
    """生成 2.2 端 / 角色拆分总览（每端职责一句话 + 场景索引表）。

    把 2.1 主线落到第 3/4/5 章模块视图。每端职责强制一句话，避免堆段落。

    Args:
        doc: python-docx Document
        role_summaries: list of (端名, 一句话职责) tuples
                        例：[('View 1 · App 用户端', '看 Feed 推荐 → 进个人主页 → 订阅'),
                             ('View 2 · 主播工作台', '开播 → 连麦管理 → 推策略卡')]
        scene_index: list of (编号, 场景名, 模块, 优先级) tuples
                     跟 prd-template.md 第 2 章索引表同 schema

    生成结构：

        ## 2.2 端 / 角色拆分总览
        ### 每端职责
        • View 1 · App 用户端：[一句话]
        • View 2 · 主播工作台：[一句话]
        ### 场景索引
        | 编号 | 场景 | 模块 | P |
    """
    base = _import_base()
    h2, h3, bullet, make_table, C = (
        base.h2, base.h3, base.bullet, base.make_table, base.C,
    )

    h2(doc, '2.2 端 / 角色拆分总览')

    h3(doc, '每端职责')
    if not role_summaries:
        base.add_p(doc, '（待补充：列出各端 / 角色一句话职责）',
                   size_pt=10, italic=True, color=C['textMuted'])
    else:
        for role, summary in role_summaries:
            bullet(doc, f'{role}：{summary}')

    h3(doc, '场景索引')
    if not scene_index:
        base.add_p(doc, '（待补充：见 scene-list.md）',
                   size_pt=10, italic=True, color=C['textMuted'])
    else:
        rows = [list(r) for r in scene_index]
        make_table(doc, headers=['编号', '场景', '模块', 'P'],
                   rows_data=rows, col_widths_cm=[2, 9, 5, 1.5])


def cross_end_sequence_section(doc, *, sequences=None, shot_dir=None):
    """生成 2.3 跨端时序（多张跨端流程图 + 一句话说明）。

    分模块写第 3/4/5 章时跨端逻辑会被切碎，每端只看到半边，时序图把它们焊起来。
    QA 写测试用例时这一节是黄金素材。

    涉及 ≥ 2 端项目强制至少 1 张跨端流程图（check_prd.sh 自检）。单端项目传
    sequences=None 或 [] 自动占位。

    **图类型按场景选**（mermaid 支持多种，按故事和复杂度选）：
    - 跨端协作含分支判定 / Yes-No 路径 → `flowchart LR + subgraph` 分组流程图
      （目前 PRD docx 嵌图主用，subgraph 模拟角色分组，**不是严格泳道图**）
    - 强调时序消息 / 请求-响应往返 → `sequenceDiagram`
    - 状态机切换 / 多状态流转 → `stateDiagram-v2`
    - 单端简单流程（≤ 5 节点）→ 不出图，文字一段说清
    严格 swimlane 泳道图 mermaid 不直接支持，需要时回归 X6 swimlane（参
    flowchart skill），但 X6 嵌入 docx 视觉差被淘汰，慎用。

    v5.3 后从 X6/dagre 切换到 mermaid——每个 sequences 元素独立 mermaid 源码 →
    独立 PNG（一对一）。修了 v5.1/v5.2 cross_end 嵌错图的 bug（原 _embed_flowchart
    scope='first' 写死取第 0 张图，导致 cross_end 永远嵌主线图）。

    Args:
        doc: python-docx Document
        sequences: list of (标题, mermaid_src, 一句话说明) tuples
                   每个元素独立 mermaid 源码，独立渲染独立 PNG。
                   例：[('观众举手 → 主播同意', '''flowchart LR\\n  ...''',
                         '观众端举手 → 主播工作台审批 → 双端席位状态同步'),
                        ('上麦中断异常', '''flowchart LR\\n  ...''',
                         '网络断连时主播 / 观众如何感知，自动重连 vs 强制下麦')]
        shot_dir: PNG 输出目录，默认 projects/{项目}/screenshots/prd/mermaid/

    生成结构：

        ## 2.3 跨端时序
        ### 观众举手 → 主播同意
        [一句话说明]
        [跨端流程图 PNG]
        ### 上麦中断异常
        [一句话说明]
        [跨端流程图 PNG]

    单端 / 空 sequences 时：

        ## 2.3 跨端时序
        本项目仅 1 端，无跨端时序
    """
    base = _import_base()
    h2, h3, add_p, C = base.h2, base.h3, base.add_p, base.C
    import re as _re

    h2(doc, '2.3 跨端时序')

    if not sequences:
        add_p(doc, CROSS_END_PLACEHOLDER, italic=True, color=C['textMuted'])
        return

    for idx, item in enumerate(sequences):
        if len(item) == 3:
            title, mermaid_src, note = item
        elif len(item) == 2:
            title, mermaid_src = item
            note = ''
        else:
            raise ValueError(
                f'cross_end_sequence_section: sequences 元素需 '
                f'(title, mermaid_src, note?) 2 或 3 元组，收到 {item}'
            )

        h3(doc, title)
        if note:
            add_p(doc, note)
        # 一对一独立 PNG：用 idx + 标题 slug 命名，避免 v5.1/v5.2 嵌错图 bug
        slug = _re.sub(r'[^\w一-鿿]+', '-', title).strip('-')[:30]
        png_name = f'cross-end-{idx:02d}-{slug}'
        embedded = _embed_mermaid(doc, mermaid_src, shot_dir, png_name=png_name)
        if not embedded and mermaid_src:
            add_p(doc, f'⚠️ 跨端时序图渲染失败：{title}',
                  size_pt=10, italic=True, color=C['textMuted'])

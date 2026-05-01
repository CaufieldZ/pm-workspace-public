"""标题与段落级 helper：h1/h2/h3 + 段落 + bullet + chapter_story + pullquote +
normalize_headings（修旧 docx Heading run 级样式）。

依赖 core.styles 的 C / para_run。
"""
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from core.styles import C, FONT, para_run


# ── 段落级（基础）──────────────────────────────────────────────────────────

def add_p(doc, text="", size_pt=10, bold=False, color=None, italic=False,
          align=WD_ALIGN_PARAGRAPH.LEFT, before=0, after=4):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    if text:
        para_run(p, text, size_pt=size_pt, bold=bold, color=color, italic=italic)
    return p


def bullet(doc, text, size_pt=10):
    lines = text.split('\n')
    first = True
    for line in lines:
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0 if first else 0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(0.5)
        para_run(p, f"• {line}", size_pt=size_pt, color=C["textPrimary"])
        first = False


# ── 章节标题（h1/h2/h3）────────────────────────────────────────────────────

def h1(doc, text):
    """章标题。出版物语言：加粗字号 + 加厚下边框 terra cotta 横线，留白拉开。"""
    p = doc.add_paragraph(style='Heading 1')
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(12)
    para_run(p, text, font='title', size_pt=18, bold=True, color=C["textHeading"])
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')           # 1.5pt 粗线（原 6 = 0.75pt）
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), C["accentBlue"])  # terra cotta
    pBdr.append(bottom)
    pPr.append(pBdr)


def h2(doc, text):
    """节标题。加 6pt 左色块 border 作 anchor（出版物锚点用法），文字走 terra cotta。"""
    p = doc.add_paragraph(style='Heading 2')
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(0.3)  # 让左色块和文字有空隙
    para_run(p, text, font='title', size_pt=13, bold=True, color=C["accentBlue"])
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')           # 3pt 短色块
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), C["accentBlue"])
    pBdr.append(left)
    pPr.append(pBdr)


def h3(doc, text):
    p = doc.add_paragraph(style='Heading 3')
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    para_run(p, text, font='title', size_pt=11, bold=True, color=C["textHeading"])


# ── 章节用户故事 / 重点金句 / 关键数字 ─────────────────────────────────────

def chapter_story(doc, text):
    """章节用户故事引言（pm-workflow.md「PART/章节用户故事陈述」强制）。
    紧跟 h1 后调用，斜体浅灰一句话 ≤ 30 字，告诉读者这章在讲谁、做什么、为什么。
    技术骨架章（背景/排期/非功能）可省略。"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    para_run(p, text, size_pt=10, italic=True, color="6F7A85")
    return p


def pullquote(doc, text):
    """重点金句块，替代正文里堆 bold。出版物 pullquote 语言：
    左 3pt terra cotta border + Lora italic + 12pt + textHeading 暗色"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.left_indent = Cm(0.6)
    para_run(p, text, font='title', size_pt=12, italic=True, color=C["textHeading"])
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '24')           # 3pt
    left.set(qn('w:space'), '12')
    left.set(qn('w:color'), C["accentBlue"])
    pBdr.append(left)
    pPr.append(pBdr)
    return p


def metric_run(p, value):
    """在段落内插入 mono terra cotta 数字（如 "47%" "5%+"），
    替代正文全 bold，让数字像 FT/经济学人那样亮起来。"""
    return para_run(p, value, font='mono', size_pt=10, bold=True, color=C["accentBlue"])


# ── 锚点插入 + Heading 修缮（patch / 升版用） ──────────────────────────────

def _apply_heading_border(p, hex_color: str):
    """H1 专用的段落下边框。幂等：已有 pBdr 不覆盖。"""
    pPr = p._p.get_or_add_pPr()
    if pPr.find(qn('w:pBdr')) is not None:
        return
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def insert_heading_before(anchor, text: str, level: int = 2):
    """在 anchor 段落之前插入 Heading {level} 段落，run 级样式和 h1/h2/h3 一致。

    用途：patch 脚本给 PRD 新增章节（如 6.4 归因漏斗 / 8.3 事件列表）。
    不传 level 默认 H2。anchor 不可为 None。

    对比 anchor.insert_paragraph_before(text)：默认产 Normal style，
    Confluence 识别不出层级 → 大纲树失效（memory #2 踩坑）。
    """
    p = anchor.insert_paragraph_before('')
    # 直接改底层 pStyle，绕过 python-docx styles[key] 在重复 styleId 文档下误报
    # KeyError 的 bug（部分 gen_prd_base 产出的 docx 有 Heading1/2 重复定义）
    pPr = p._p.get_or_add_pPr()
    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = OxmlElement('w:pStyle')
        pPr.insert(0, pStyle)
    pStyle.set(qn('w:val'), f'Heading{level}')
    # style 只套模板，run 级样式还得手动刷（和 h1/h2/h3 对齐）
    if level == 1:
        color, size_pt = C["textHeading"], 16
        _apply_heading_border(p, C["accentBlue"])
    elif level == 2:
        color, size_pt = C["accentBlue"], 13
    else:
        color, size_pt = C["textHeading"], 11
    para_run(p, text, font='title', size_pt=size_pt, bold=True, color=color)
    return p


def _apply_heading_style(p, hex_color, size_pt, add_border=None):
    """给 heading 段落强制刷 run 样式，可选追加段落下边框。
    强制覆盖颜色/字号/粗体/字体：老 docx 里即使 H1 已设错颜色（如黑色）也统一刷成标准色，
    字体强制刷成 title role（Lora + Noto Serif SC）对齐 h1/h2。
    """
    cfg = FONT['title']
    for r in p.runs:
        r.font.color.rgb = RGBColor.from_string(hex_color)
        r.font.size = Pt(size_pt)
        r.bold = True
        r.font.name = cfg['ascii']
        rPr = r._r.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.insert(0, rFonts)
        rFonts.set(qn('w:ascii'), cfg['ascii'])
        rFonts.set(qn('w:hAnsi'), cfg['ascii'])
        rFonts.set(qn('w:eastAsia'), cfg['eastAsia'])
    # 无 run（段落只有 text 没 run）时新建一个 - 走 para_run 含字体配置
    if not p.runs and p.text:
        text = p.text
        for child in list(p._p):
            if child.tag.endswith('}r') or child.tag.endswith('}t'):
                p._p.remove(child)
        para_run(p, text, font='title', size_pt=size_pt, bold=True, color=hex_color)
    # 段落下边框（H1 专用）
    if add_border:
        pPr = p._p.get_or_add_pPr()
        if pPr.find(qn('w:pBdr')) is None:
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), add_border)
            pBdr.append(bottom)
            pPr.append(pBdr)


def normalize_headings(doc):
    """修旧 docx 的 H1/H2 样式：补 run 级字色/字号/粗体 + H1 补下划线 pBdr。
    旧脚本常用 add_paragraph(style='Heading 1') 直接产段落，run 级属性缺失，
    渲染效果取决于 Word 默认 style（黑无线），和 h1()/h2() 产出的标准
    视觉不一致。本函数幂等：已有 run 属性的跳过不覆盖。

    H1: fg=#141413 (Anthropic Dark) + 16pt + bold + 段落下边框 #D97757 (terra cotta)
    H2: fg=#D97757 + 13pt + bold

    返回修复数 (h1_count, h2_count)。
    """
    h1_fixed = h2_fixed = 0
    for p in doc.paragraphs:
        s = p.style.name if p.style else ""
        if s == "Heading 1":
            _apply_heading_style(p, C["textHeading"], 16, add_border=C["accentBlue"])
            h1_fixed += 1
        elif s == "Heading 2":
            _apply_heading_style(p, C["accentBlue"], 13, add_border=None)
            h2_fixed += 1
    return h1_fixed, h2_fixed

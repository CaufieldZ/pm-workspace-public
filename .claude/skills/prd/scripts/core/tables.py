"""表格 helper：通用 make_table / scene 双列表格 / cell 结构化填充 + 老 cell 归一。

依赖 core.styles 的 C / FONT / para_run / para_run_md / cell_text / set_cell_bg。
"""
import re

from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL

from core.styles import C, para_run, para_run_md, cell_text, set_cell_bg


# ── 通用表格 ───────────────────────────────────────────────────────────────

def make_table(doc, headers, rows_data, col_widths_cm, row_height_cm=None,
               mono_first_col=False):
    """通用表格生成：headers=列名列表, rows_data=二维列表, col_widths_cm=每列宽度。
    mono_first_col=True 时首列走 mono（适合编号 / 角色 ID / 路径，让节奏感出来）"""
    tbl = doc.add_table(rows=1 + len(rows_data), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for j, (hdr, w) in enumerate(zip(headers, col_widths_cm)):
        cell = tbl.rows[0].cells[j]
        cell.width = Cm(w)
        set_cell_bg(cell, C["tableHeaderBg"])
        cell_text(cell, hdr, size_pt=9, bold=True, color=C["tableHeaderText"],
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, row in enumerate(rows_data):
        for j, (val, w) in enumerate(zip(row, col_widths_cm)):
            cell = tbl.rows[i+1].cells[j]
            cell.width = Cm(w)
            if i % 2 == 1:
                set_cell_bg(cell, C["tableAltRow"])
            # 首列 mono 渲染（编号节奏）
            if mono_first_col and j == 0:
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                para_run(p, str(val), font='mono', size_pt=9, bold=True,
                         color=C["textHeading"])
            else:
                cell_text(cell, str(val), size_pt=9, color=C["textPrimary"])
        if row_height_cm:
            tbl.rows[i+1].height = Cm(row_height_cm)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl


# ── Scene 双列表格（核心：左截图占位 + 右说明）──────────────────────────

def scene_table(doc, scene_id, scene_name, right_blocks):
    """两列 Scene 表格
    right_blocks: list of (title, lines) -- title 可含「（变更）」或「（新增）」
    """
    tbl = doc.add_table(rows=1, cols=2)
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    left_cell = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]
    left_cell.width = Cm(8.5)
    right_cell.width = Cm(13.0)
    set_cell_bg(left_cell, C["tableAltRow"])
    left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = left_cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    para_run(p, f"\U0001f4f1 {scene_id} · {scene_name}", size_pt=9, color=C["textMuted"], italic=True)
    p2 = left_cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_run(p2, "← 此处粘贴原型截图", size_pt=8, color=C["textMuted"], italic=True)

    right_cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    fill_cell_blocks(right_cell, right_blocks)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ── cell 结构化填充（公共逻辑，scene_table + set_cell_blocks 共用）────────

_NUMBERED_PREFIX_RE = re.compile(r"^\d+\.\s")


def fill_cell_blocks(cell, blocks, numbered=True):
    """在已存在的 cell 内填充结构化 blocks（title 粗体 + 子条目缩进）。
    假设 cell 的首段已为空或将被覆写。不清空 cell，不会处理旧内容。
    被 scene_table 和 set_cell_blocks 共用。

    blocks: list[tuple[str, list[str]]]
    numbered: True 时自动给每个 block 内的主级 line 加 "1./2./3." 前缀
              （已含 "N. " 前缀的 line 保持原样，二级缩进行 "  - " 不编号）。
              prd-template.md 规定"每个模块下用 1./2./3. 数字编号"，默认开启。
    """
    first = True
    for (title, lines) in blocks:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        is_first_block = first
        first = False
        # 第一个 block 顶贴 cell 顶部，后续 block 前留 6pt 拉开层级
        p.paragraph_format.space_before = Pt(0) if is_first_block else Pt(6)
        p.paragraph_format.space_after = Pt(2)
        if "（变更）" in title or "（新增）" in title:
            tag = "（变更）" if "（变更）" in title else "（新增）"
            base = title.replace(tag, "")
            # title 11pt + bold + textHeading；tag 走 mono 染色
            para_run(p, base, font='title', size_pt=11, bold=True, color=C["textHeading"])
            para_run(p, tag, font='mono', size_pt=11, bold=True, color=C["tagChange"])
        else:
            para_run(p, title, font='title', size_pt=11, bold=True, color=C["textHeading"])
        counter = 0
        for line in lines:
            pl = cell.add_paragraph()
            pl.paragraph_format.space_before = Pt(0)
            pl.paragraph_format.space_after = Pt(1)
            stripped = line.lstrip()
            is_sublevel = stripped.startswith("- ") and line != stripped
            if is_sublevel:
                pl.paragraph_format.left_indent = Cm(1.2)
                body = stripped
            else:
                pl.paragraph_format.left_indent = Cm(0.6)
                if numbered and not _NUMBERED_PREFIX_RE.match(line):
                    counter += 1
                    body = f"{counter}. {line}"
                else:
                    body = line
            para_run_md(pl, body, size_pt=9, color=C["textPrimary"])


# ── 升版用 cell 操作 ──────────────────────────────────────────────────────

def set_cell_text(cell, text: str):
    """清空单元格所有内容，在首段写入纯文本。
    保留首段首 run 的格式；无 run 则新建。
    """
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ''
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
    else:
        p.add_run(text)


def set_cell_blocks(cell, blocks, numbered=True):
    """清空单元格，按结构化 blocks 重建（title 粗体 + 子条目缩进 + 「（变更）」「（新增）」染色）。
    视觉和 scene_table 的右列一致。

    用途：升版/更新 PRD 时，给 Scene 表格右列或后台 table 的复杂内容做结构化替换，
          替代 set_cell_text —— 后者只能单 run 纯文本，没有层次。

    参数：
        cell   — python-docx Table cell
        blocks — list[tuple[str, list[str]]]
                 [(title, [line1, line2, ...]), ...]
                 title 含「（变更）」或「（新增）」会自动拆 run 染成强调色。
    """
    # 清空 cell 所有现有内容（文字 + 图片）
    for para in cell.paragraphs:
        para.clear()
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)
    fill_cell_blocks(cell, blocks, numbered=numbered)


# ── 老 cell 归一（V2.7-era docx 风格不一致 → 统一 fill_cell_blocks 风格）──

def cell_paragraphs_to_blocks(cell):
    """识别 scene_table 右列 cell 内的 (title, lines) 块结构。

    兼容两种 V2.7-era docx 的 cell 模式：

    模式 A（title=bold + 无空段 + line 自带 1./2./3.）— T9/E-2 风格：
        title (bold) → 1. line → 2. line → next title (bold) → ...
    模式 B（title=普通 + 空段分隔 + line 平铺无编号）— T20/C-3 风格：
        title → line → line → 空段 → next title → ...

    通用规则：空段 OR bold 段 = block 强分隔符；
            分隔符之间第一段是 title，其余是 lines。

    返回 list[tuple[str, list[str]]] 格式，可直接喂 set_cell_blocks(cell, blocks)。
    """
    blocks = []
    cur_title = None
    cur_lines = []

    def flush():
        nonlocal cur_title, cur_lines
        if cur_title is not None:
            blocks.append((cur_title, cur_lines))
        cur_title = None
        cur_lines = []

    for p in cell.paragraphs:
        t = p.text.strip()
        if not t:
            flush()
            continue
        is_bold = bool(p.runs) and bool(p.runs[0].bold)
        if is_bold and (cur_title is not None or cur_lines):
            flush()
        if cur_title is None:
            cur_title = t
        else:
            cur_lines.append(t)
    flush()
    return blocks


def fix_scene_cell_numbering(doc):
    """扫所有 scene_table 右列，≥4 段时用 cell_paragraphs_to_blocks 拆分后
    用 set_cell_blocks 重写，统一为「title 11pt bold + lines 9pt 0.6cm 缩进 + 自动编号」。

    跳过条件：
      - 表第一列非 scene_table anchor（不以 📱/🖥/🔧 开头）
      - cell 段 < 4（描述太短，加 numbered 反而冗余）
      - 单 block 且 lines < 2（无意义层次）

    背景：V2.7-era docx 的 scene_table 右列内容 mixed 风格不一致（有 numbered
    有平铺）。本 helper 统一为 set_cell_blocks 风格，配合 fill_cell_blocks 的
    11pt title + 9pt 缩进编号渲染。
    """
    for t in doc.tables:
        if not t.rows:
            continue
        if len(t.rows[0].cells) < 2:
            continue
        cell0_text = t.rows[0].cells[0].text.strip()
        if not (cell0_text.startswith('📱') or cell0_text.startswith('🖥') or cell0_text.startswith('🔧')):
            continue
        cell = t.rows[0].cells[1]
        if len(cell.paragraphs) < 4:
            continue
        blocks = cell_paragraphs_to_blocks(cell)
        if not blocks:
            continue
        if len(blocks) == 1 and len(blocks[0][1]) < 2:
            continue
        set_cell_blocks(cell, blocks, numbered=True)

"""docx 规范化三件套：标点 + 字体 + theme。

被 save_prd 强制串调用，确保 PRD 渲染：
1. CJK 旁半角 → 全角（soul.md 硬规则）
2. 老字体（Arial / Microsoft YaHei / HarmonyOS）→ Lora / Poppins / Noto SC
3. 注入 word/theme/theme1.xml（让 docDefaults 的 themed 引用不悬空）

老 docx 升版漏跑会残留旧字体风。check_prd.sh 已加自动扫，缺则 fail。
"""
import re

from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def normalize_punctuation(doc):
    """中文标点规范化：中文相邻的半角 , : ( ) 改为全角 ，：（）。
    遵循 soul.md「中文里禁止混半角逗号冒号括号」。
    幂等；保留 run 级样式（bold / color 不变）。

    段落级判定：拼接段落所有 run 的文字后做上下文判断，避免 run 边界把
    「中文<run1 end>:<run2 start>英文」这种场景看成孤立冒号漏改。

    判定规则：
      - 半角 , : ( ) 只要左右任一是中文字符 → 替换全角
      - 排除 URL 场景：前 6 字符含 http / ftp
    代码 / URL / 纯英文上下文不触发（两侧都没中文）。
    """
    CJK_RE = re.compile(r'[一-鿿]')
    MAP = {',': '，', ':': '：', '(': '（', ')': '）'}

    def process_paragraph(p):
        runs = p.runs
        if not runs: return 0
        full = ''.join(r.text or '' for r in runs)
        if not any(c in MAP for c in full): return 0

        out = list(full)
        for i, c in enumerate(full):
            if c not in MAP: continue
            prev = full[i-1] if i > 0 else ''
            nxt  = full[i+1] if i+1 < len(full) else ''
            if not (CJK_RE.match(prev) or CJK_RE.match(nxt)):
                continue
            if c == ':':
                window = full[max(0, i-6):i].lower()
                if 'http' in window or 'ftp' in window:
                    continue
            out[i] = MAP[c]

        new_full = ''.join(out)
        if new_full == full: return 0

        n = 0
        pos = 0
        for r in runs:
            L = len(r.text or '')
            new_slice = new_full[pos:pos+L]
            if new_slice != (r.text or ''):
                r.text = new_slice
                n += sum(1 for a, b in zip(r.text, full[pos:pos+L]) if a != b)
            pos += L
        return n

    total = 0
    for p in doc.paragraphs:
        total += process_paragraph(p)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total += process_paragraph(p)
    return total


def normalize_fonts(doc):
    """老 docx 升版：把硬编码 Arial / Microsoft YaHei 替换为 FONT 字典定义的设计字体。

    覆盖：
    1. styles.xml docDefaults rFonts 写死 → 改为 themed (minorHAnsi/minorEastAsia)。
       让 Word 找不到 run 级字体时 fallback 到主题字体（Mac 走 PingFang+Calibri，
       Win 走 Microsoft YaHei+Calibri），而不是 substitute 成 Arial。
    2. 所有 run rFonts：老 ascii (Arial/Calibri/Times/Source Serif 4/Plus Jakarta Sans) → Poppins；
       老 eastAsia (Arial/Microsoft YaHei/SimSun/HarmonyOS Sans SC 等) → Noto Sans SC。
       已是设计字体的 run 不动。

    背景：framework 美学切到 Anthropic 官方 brand-guidelines 后，PRD 字体规范升到
    Lora / Poppins / Noto Serif SC / Noto Sans SC / JetBrains Mono。本函数把老字体
    （含上一代 Source Serif 4 / Plus Jakarta Sans / HarmonyOS Sans SC）一次刷成新规范。

    幂等。返回 (run_changed, defaults_changed)。
    """
    LEGACY_ASCII = {'Arial', 'Calibri', 'Times New Roman', 'Times',
                    'Helvetica', 'Roboto', 'Open Sans',
                    'Source Serif 4', 'Plus Jakarta Sans', 'Inter'}
    LEGACY_EAST = {'Arial', 'Microsoft YaHei', '微软雅黑', 'SimSun', '宋体',
                   'SimHei', '黑体', 'STSong', 'STHeiti',
                   'HarmonyOS Sans SC'}
    DESIGN = {'Lora', 'Poppins', 'Noto Serif SC',
              'Noto Sans SC', 'JetBrains Mono', 'PingFang SC'}

    # 1. docDefaults → themed
    defaults_changed = False
    styles_el = doc.styles.element
    docDefaults = styles_el.find(qn('w:docDefaults'))
    if docDefaults is not None:
        rPrDefault = docDefaults.find(qn('w:rPrDefault'))
        if rPrDefault is not None:
            rPr = rPrDefault.find(qn('w:rPr'))
            if rPr is not None:
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is not None:
                    has_legacy = any(
                        rFonts.get(qn(f'w:{k}')) in (LEGACY_ASCII | LEGACY_EAST)
                        for k in ('ascii', 'hAnsi', 'eastAsia', 'cs')
                    )
                    if has_legacy:
                        for k in ('ascii', 'hAnsi', 'eastAsia', 'cs'):
                            attr = qn(f'w:{k}')
                            if attr in rFonts.attrib:
                                del rFonts.attrib[attr]
                        rFonts.set(qn('w:asciiTheme'), 'minorHAnsi')
                        rFonts.set(qn('w:hAnsiTheme'), 'minorHAnsi')
                        rFonts.set(qn('w:eastAsiaTheme'), 'minorEastAsia')
                        rFonts.set(qn('w:cstheme'), 'minorBidi')
                        defaults_changed = True

    # 2. 所有 run rFonts：老字体 / 缺失 → 强制设 body 设计字体
    run_changed = 0
    W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
    for r in doc.element.body.iter(f'{{{W}}}r'):
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            r.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rFonts.set(qn('w:ascii'), 'Poppins')
            rFonts.set(qn('w:hAnsi'), 'Poppins')
            rFonts.set(qn('w:eastAsia'), 'Noto Sans SC')
            rPr.insert(0, rFonts)
            run_changed += 1
            continue
        changed = False
        for slot, default_font, legacy_set in (
            ('ascii',    'Poppins',       LEGACY_ASCII),
            ('hAnsi',    'Poppins',       LEGACY_ASCII),
            ('eastAsia', 'Noto Sans SC',  LEGACY_EAST),
        ):
            v = rFonts.get(qn(f'w:{slot}'))
            if v is None or (v in legacy_set and v not in DESIGN):
                rFonts.set(qn(f'w:{slot}'), default_font)
                changed = True
        if changed:
            run_changed += 1

    return run_changed, defaults_changed


def ensure_theme(docx_path):
    """python-docx 生成的 docx 缺 word/theme/theme1.xml；docDefaults 的 themed 引用
    （minorHAnsi/minorEastAsia）会悬空，Word 只能 fallback 到 Arial。

    本函数把 .claude/skills/prd/assets/theme1.xml 标准 Office 主题注入 docx zip：
      - 写 word/theme/theme1.xml
      - [Content_Types].xml 加 theme override
      - word/_rels/document.xml.rels 加 theme relationship

    幂等：已有 theme1.xml 直接返回 False。

    用法：doc.save() **之后**调用，传入 docx 文件路径（不能传 Document 对象）。
    """
    import zipfile
    import shutil
    import os

    # core/normalize.py 在 scripts/core/ 下，向上 2 级到 skill 根
    _SKILL_ROOT = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
    THEME_XML = os.path.join(_SKILL_ROOT, 'assets', 'theme1.xml')
    if not os.path.exists(THEME_XML):
        return False

    docx_path = str(docx_path)
    with zipfile.ZipFile(docx_path) as z:
        if 'word/theme/theme1.xml' in z.namelist():
            return False

    with open(THEME_XML, 'rb') as f:
        theme_bytes = f.read()

    THEME_OVERRIDE = (
        '<Override PartName="/word/theme/theme1.xml" '
        'ContentType="application/vnd.openxmlformats-officedocument.theme+xml"/>'
    )
    THEME_REL = (
        '<Relationship Id="rIdTheme1" '
        'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/theme" '
        'Target="theme/theme1.xml"/>'
    )

    temp_path = docx_path + '.temp'
    with zipfile.ZipFile(docx_path) as zin:
        with zipfile.ZipFile(temp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    s = data.decode('utf-8')
                    if 'theme/theme1.xml' not in s:
                        s = s.replace('</Types>', THEME_OVERRIDE + '</Types>')
                    data = s.encode('utf-8')
                elif item.filename == 'word/_rels/document.xml.rels':
                    s = data.decode('utf-8')
                    if 'theme1.xml' not in s:
                        s = s.replace('</Relationships>', THEME_REL + '</Relationships>')
                    data = s.encode('utf-8')
                zout.writestr(item, data)
            zout.writestr('word/theme/theme1.xml', theme_bytes)
    shutil.move(temp_path, docx_path)
    return True

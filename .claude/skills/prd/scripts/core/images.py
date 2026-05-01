"""图片操作：DPI 修正 + cell 内换图 + phone 圆角透明化。

被 screenshots.py 复用，给 ops-handbook 等下游 skill 截图统一处理。

依赖 core.styles 的 C / para_run。
"""
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from core.styles import C, para_run


# ── DPI 修正 ──────────────────────────────────────────────────────────────

def fix_dpi(png_path: str, dpi: int = 144) -> str:
    """修正 @2x 截图 DPI 元数据（Playwright deviceScaleFactor=2 输出 72dpi，
    python-docx 按 DPI 推算尺寸会导致截图虚化）。

    原地修改，返回原路径。
    """
    from PIL import Image
    img = Image.open(png_path)
    img.save(png_path, dpi=(dpi, dpi))
    return png_path


# ── phone 圆角透明化（截图后处理）─────────────────────────────────────────

def round_phone_corners(png_path: str, radius_px: int = 72):
    """给 phone 截图加透明圆角蒙版。
    CSS border-radius 36px × deviceScaleFactor 2 = 72px。"""
    from PIL import Image, ImageDraw
    img = Image.open(png_path).convert("RGBA")
    mask = Image.new("L", img.size, 0)
    ImageDraw.Draw(mask).rounded_rectangle([(0, 0), img.size], radius=radius_px, fill=255)
    img.putalpha(mask)
    img.save(png_path, "PNG")


# ── cell 内换图 ───────────────────────────────────────────────────────────

def replace_cell_image(cell, img_path: str, width_cm: float = 7.0):
    """清空单元格（文字+旧图），插入新截图。
    内部强制调用 fix_dpi()，确保 docx 里不虚化。

    参数：
        cell       — python-docx Table cell
        img_path   — PNG 文件路径（str 或 Path）
        width_cm   — 插入宽度，前端 Scene 建议 7.0，App 建议 5.0
    """
    img_path = str(img_path)
    fix_dpi(img_path)

    # 清空所有段落内容（文字 + drawing）
    for para in cell.paragraphs:
        para.clear()
    # 移除多余段落，只留一个
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)
    # 插入图片
    p = cell.paragraphs[0]
    run = p.add_run()
    run.add_picture(img_path, width=Cm(width_cm))


def replace_cell_image_keep_title(cell, scene_title: str, img_path: str, width_cm: float = 5.0):
    """清空 cell 占位 → 保留场景标题（首段）→ 插入截图（次段）。

    替代 replace_cell_image 用于 scene_table 左列：后者会把 scene_table 写在 cell 内的
    场景标题段（"📱 B-1 · ..."）也清掉，导致 Confluence 页面看不出图属于哪个场景。

    Felix 反馈（memory `feedback_prd_iter_screenshot_pitfalls.md` 坑 2）：
        「你这些怎么都没写具体是哪个页面？」

    用法：
        from core.images import replace_cell_image_keep_title
        # 从 cell 首段提取场景标题
        title = cell.paragraphs[0].text.split('\\n')[0].strip()
        replace_cell_image_keep_title(cell, title, png_path, width_cm=5.0)

    prd_screenshots.py 的 insert_into_docx 已自动在 cell 首段含 `📱 X-N · ...` 时调用本函数。
    """
    fix_dpi(str(img_path))
    for para in cell.paragraphs:
        para.clear()
    while len(cell.paragraphs) > 1:
        p = cell.paragraphs[-1]
        p._element.getparent().remove(p._element)

    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(4)
    para_run(p1, scene_title, size_pt=9, color=C["textMuted"], italic=True)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run().add_picture(str(img_path), width=Cm(width_cm))

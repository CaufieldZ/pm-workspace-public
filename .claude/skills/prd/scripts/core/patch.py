"""patch / 升版用 helper：锚点定位 + 段落插入 + Scene blocks 插入 + 段落级 search/replace。

依赖 core.styles 的 para_run / para_run_md / C，core.headings 的 insert_heading_before。
"""
import re

from docx.shared import Cm, Pt

from core.styles import C, para_run, para_run_md
from core.headings import insert_heading_before


# ── 锚点定位 ───────────────────────────────────────────────────────────────

def _find_anchor_paragraph(doc, anchor_keyword: str):
    """按 text.strip().startswith(anchor_keyword) 找到第一个匹配段落。
    找不到抛 ValueError，不静默失败（patch 脚本最忌锚点错位）。"""
    for p in doc.paragraphs:
        if p.text.strip().startswith(anchor_keyword):
            return p
    raise ValueError(f"anchor 未命中: {anchor_keyword!r}")


# ── 纯插入（反 #6 踩坑：模糊段落覆盖容易咬错下一级标题）────────────────────

def insert_paragraph_before(anchor, text: str, size_pt: int = 10,
                             bold: bool = False, color: str = None):
    """在 anchor 前插入 Normal 段落。纯插入不覆盖已有段 text。

    设计意图：替代「查下一段落覆盖 text」这种模糊定位（memory #6 踩坑 ——
    patch_prd_v3_3.py 因为跳过 2.x 条件太宽，咬到 3. 大标题把 H1 text 覆盖了）。

    硬规则（配 SKILL.md）：patch 脚本插入新段落一律用本函数，
    禁止 anchor.insert_paragraph_before(...) 裸调用（会产 Normal 无格式），
    禁止遍历 paragraphs[i+1:i+5] 找「下一段」改 text。
    """
    p = anchor.insert_paragraph_before('')
    para_run(p, text, size_pt=size_pt, bold=bold, color=color)
    return p


def insert_description_after(doc, anchor_keyword: str, text: str,
                              size_pt: int = 10):
    """找到 anchor_keyword 段，在其后插入描述段。锚点未命中抛异常，
    不做「跳过若干段找位置」的模糊定位。

    实现：定位到 anchor 的下一段，在下一段之前插入 → 等价于 anchor 后插入。
    anchor 已经是文档最后一段时直接 doc.add_paragraph。
    """
    anchor = _find_anchor_paragraph(doc, anchor_keyword)
    body = anchor._p.getparent()
    idx = list(body).index(anchor._p)
    if idx + 1 < len(body):
        next_elem = body[idx + 1]
        from docx.text.paragraph import Paragraph
        next_para = Paragraph(next_elem, anchor._parent)
        p = next_para.insert_paragraph_before('')
    else:
        p = doc.add_paragraph()
    para_run(p, text, size_pt=size_pt)
    return p


_NUMBERED_PREFIX_RE = re.compile(r"^\d+\.\s")


def insert_scene_blocks(anchor, blocks, heading_level: int = 3):
    """在 anchor 段前批量插入「H{level} 小节标题 + numbered list」结构，
    用于 body 级别的 Scene 详细说明（纯后台 / CMS / 无截图场景降级写法，
    替代 scene_table 2 列表格）。视觉与 fill_cell_blocks 对齐。

    blocks: list[tuple[str, list[str]]]
            [(section_title, [line1, line2, ...]), ...]
            line 支持 `**label**：content` 内联 bold（通过 para_run_md 解析）；
            以 `  - ` 开头的行作为二级缩进，不编号；
            已含 `N. ` 前缀的行保持原样（不重复编号）。
    heading_level: 默认 H3(3)。

    用途场景：CMS 管理后台等「本次无 UI 改动 / 无独立截图」的 Scene 章节，
    用 H2 场景标题 + 本函数展开子章节，避免 scene_table 左列空截图占位符冗余。
    """
    for (title, lines) in blocks:
        insert_heading_before(anchor, title, level=heading_level)
        counter = 0
        for line in lines:
            stripped = line.lstrip()
            is_sublevel = stripped.startswith("- ") and line != stripped
            p = anchor.insert_paragraph_before('')
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(1)
            if is_sublevel:
                p.paragraph_format.left_indent = Cm(0.9)
                body = stripped
            else:
                p.paragraph_format.left_indent = Cm(0.3)
                if not _NUMBERED_PREFIX_RE.match(line):
                    counter += 1
                    body = f"{counter}. {line}"
                else:
                    body = line
            para_run_md(p, body, size_pt=10, color=C["textPrimary"])


def remove_table(table):
    """从 body 中移除整张 table（含所有 row / cell）。
    配合 insert_scene_blocks 做「去 scene_table、改 H3 + numbered list」重构。
    """
    table._element.getparent().remove(table._element)


# ── 段落级 search / replace ───────────────────────────────────────────────

def replace_para_text(para, new_text: str):
    """整体替换段落文字，保留第一个 run 的格式（字号/粗体/颜色）。
    后续 run 清空。
    """
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ''
    else:
        para.add_run(new_text)


def search_replace_para(para, old: str, new: str) -> bool:
    """段落内搜索替换（跨 run 拼接后匹配）。
    替换成功返回 True，未命中返回 False。保留第一个 run 的格式。
    """
    full = ''.join(r.text for r in para.runs)
    if old not in full:
        return False
    new_full = full.replace(old, new)
    if para.runs:
        para.runs[0].text = new_full
        for r in para.runs[1:]:
            r.text = ''
    return True

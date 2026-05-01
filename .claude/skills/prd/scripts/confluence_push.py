#!/usr/bin/env python3
"""
push_to_confluence_base.py — 推送 PRD docx 到 Confluence Server

两种模式:
    A. 按 space + title upsert(找不到则建新页):
        python3 push_to_confluence_base.py <docx> --title <标题> [--space <key>] [--parent-id <id>]
    B. 按 pageId 直接覆盖(已知页面,跳过 title 查找):
        python3 push_to_confluence_base.py <docx> --page-id <id> [--title <新标题>]

说明:
    - 凭证自动从 .mcp.json / .mcp-disabled.json 读取(MCP 关闭也能跑,脚本走 REST)
    - --space 不传时默认 `jituankejizhongxin`(集团科技中心),与 md_to_confluence.py 一致
    - 项目级配置 projects/{项目}/.confluence.json 自动读写:
        首次推送传 --page-id 或 --title/--parent-id,成功后字段写回 .confluence.json;
        后续直接 `push <docx>` 即可,脚本从配置取 page_id/space/title/parent_id。
    - 图片提取后作为附件上传,再用 ac:image 宏引用
    - push 前 pre-flight 检查 docx 的 Heading 1/2 计数,为 0 时 warning
      (memory #1/#3 踩坑:gen 脚本 BASE 路径错 → h1/h2 函数失效 → docx 全 Normal
       → Confluence 大纲树失效,不先检查事后很难发现)
    - 依赖: pip install mammoth requests

示例:
    # 模式 A:新建/upsert
    python3 .claude/skills/prd/scripts/push_to_confluence_base.py \\
        projects/growth/activity-center/deliverables/prd-htx-v1.docx \\
        --title "HTX 活动中心 PRD v1" --space HTX --parent-id 12345678

    # 模式 B:按 pageId 覆盖
    python3 .claude/skills/prd/scripts/push_to_confluence_base.py \\
        projects/community/base/deliverables/prd-htx-community-v3.4.docx \\
        --page-id 155652375
"""

import os
import sys, re, json, argparse, requests
from pathlib import Path

_ROOT = Path(__file__).resolve().parents[4]
sys.path.insert(0, os.path.join(_ROOT, 'scripts'))

try:
    import mammoth
except ImportError:
    sys.exit("缺依赖: pip install mammoth requests")

from lib.confluence import api_request, base_url, load_creds


DEFAULT_SPACE = "jituankejizhongxin"


def _find_project_config(docx_path: str) -> Path | None:
    """从 docx 路径回溯找 {项目}/.confluence.json。

    Schema v2 后项目路径有两种：
      - projects/{产品线}/{项目}/deliverables/xxx.docx → projects/{产品线}/{项目}/.confluence.json
      - projects/{顶级}/deliverables/xxx.docx → projects/{顶级}/.confluence.json

    用 deliverables/ 作为锚点：项目目录 = 包含 deliverables/ 的那一层。
    找不到返回 None（脚本不强求配置存在）。
    """
    p = Path(docx_path).resolve()
    parts = p.parts
    for i, part in enumerate(parts):
        if part == "deliverables" and i > 0:
            project_dir = Path(*parts[:i])
            return project_dir / ".confluence.json"
    return None


def _load_project_config(cfg_path: Path | None) -> dict:
    if not cfg_path or not cfg_path.exists():
        return {}
    try:
        return json.loads(cfg_path.read_text(encoding="utf-8"))
    except json.JSONDecodeError as e:
        print(f"  ⚠ {cfg_path} 解析失败,忽略: {e}", file=sys.stderr)
        return {}


def _save_project_config(cfg_path: Path | None, data: dict):
    """成功推送后回写,只写有值的字段,缩进 2 便于人读。"""
    if not cfg_path:
        return
    clean = {k: v for k, v in data.items() if v}
    cfg_path.write_text(
        json.dumps(clean, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def _rest(method: str, path: str, **kwargs):
    """requests 版 REST 封装——仅用于 multipart upload（urllib 不支持 multipart）。"""
    _base, _token = load_creds()
    url = f"{_base}/rest/api/{path.lstrip('/')}"
    headers = {"Authorization": f"Bearer {_token}", **kwargs.pop("extra_headers", {})}
    r = requests.request(method, url, headers=headers, **kwargs)
    try:
        r.raise_for_status()
    except requests.HTTPError as e:
        sys.exit(f"Confluence API 错误 [{r.status_code}]: {r.text[:300]}\n{e}")
    return r.json() if r.content else {}


def get_or_create_page(space_key: str, title: str, parent_id: str = None) -> tuple:
    """返回 (page_id, current_version_number)。页面不存在则新建占位页。"""
    resp = _rest("GET", "content", params={
        "spaceKey": space_key, "title": title, "expand": "version"
    })
    results = resp.get("results", [])
    if results:
        p = results[0]
        return p["id"], p["version"]["number"]

    # 新建占位页（先建后填内容，这样附件上传有 pageId）
    payload = {
        "type": "page",
        "title": title,
        "space": {"key": space_key},
        "body": {"storage": {"value": "", "representation": "storage"}},
    }
    if parent_id:
        payload["ancestors"] = [{"id": str(parent_id)}]
    p = _rest("POST", "content",
              json=payload,
              extra_headers={"Content-Type": "application/json"})
    return p["id"], 1


def upload_attachment(page_id: str, filename: str, data: bytes, mime: str = "image/png"):
    """上传图片附件，已存在则更新（upsert）。需要 requests 做 multipart。"""
    _base, _token = load_creds()
    _headers = {"Authorization": f"Bearer {_token}"}
    att_base = f"{_base}/rest/api/content/{page_id}/child/attachment"
    headers = {**_headers, "X-Atlassian-Token": "no-check"}

    existing = requests.get(att_base, headers=_headers,
                            params={"filename": filename}).json().get("results", [])
    if existing:
        att_id = existing[0]["id"]
        url = f"{att_base}/{att_id}/data"
    else:
        url = att_base

    requests.post(url, headers=headers,
                  files={"file": (filename, data, mime)}).raise_for_status()


# ── docx → Confluence storage format ────────────────────────────────────────

def convert_docx(docx_path: str, page_id: str,
                 img_width_web: int = 600, img_width_phone: int = 300) -> str:
    """
    1. mammoth 把 docx 转 HTML，图片提取到内存
    2. 图片上传为页面附件
    3. <img src="..."> 替换为 <ac:image>，按图片宽高比智能选 width：
       - phone 截图（高 > 宽 × 1.3）→ img_width_phone（默认 300）
       - web 截图（横屏或方形）→ img_width_web（默认 600）
    返回 Confluence storage format HTML 字符串
    """
    _images = {}   # filename → (bytes, mime)
    _idx = [0]

    def handle_image(image):
        with image.open() as f:
            data = f.read()
        mime = getattr(image, "content_type", None) or "image/png"
        ext  = mime.split("/")[-1].split("+")[0]   # image/svg+xml → svg
        _idx[0] += 1
        name = f"prd-img-{_idx[0]}.{ext}"
        _images[name] = (data, mime)
        return {"src": f"__ATTACH__{name}"}

    with open(docx_path, "rb") as f:
        result = mammoth.convert_to_html(
            f,
            convert_image=mammoth.images.img_element(handle_image)
        )

    if result.messages:
        for msg in result.messages:
            print(f"  ⚠ mammoth: {msg}")

    html = result.value

    # 算每张图的目标宽度（按宽高比）
    _widths = {}
    try:
        from PIL import Image
        from io import BytesIO
        for name, (data, _mime) in _images.items():
            try:
                w, h = Image.open(BytesIO(data)).size
                _widths[name] = img_width_phone if h > w * 1.3 else img_width_web
            except Exception:
                _widths[name] = img_width_web
    except ImportError:
        for name in _images:
            _widths[name] = img_width_web

    # 上传附件
    for name, (data, mime) in _images.items():
        upload_attachment(page_id, name, data, mime)
        print(f"  ↑ 附件: {name} ({len(data)//1024}KB, w={_widths.get(name, img_width_web)})")

    # <img src="__ATTACH__xxx"> → <ac:image> 宏
    def to_ac(m):
        src = m.group(1)
        if src.startswith("__ATTACH__"):
            fname = src[len("__ATTACH__"):]
            w = _widths.get(fname, img_width_web)
            return f'<ac:image ac:width="{w}"><ri:attachment ri:filename="{fname}"/></ac:image>'
        return m.group(0)   # 外链图片保留原样

    html = re.sub(r'<img[^>]+src="([^"]+)"[^>]*/?>', to_ac, html)

    # 缩进重建：Confluence storage format strip 段落 inline style，padding-left 失效。
    # 唯一可靠方案是把字符前缀的伪 list 段合并成原生 <ol> / <ul>，
    # 让 Confluence 用 list 自带缩进 + 自动编号 + 紧凑行距。
    html = _merge_numbered_into_ol(html)
    html = _merge_bullet_into_ul(html)

    return html


def _merge_bullet_into_ul(html: str) -> str:
    """把 mkp() 产出的 <p>• xxx</p> / <p>· xxx</p> 字符前缀伪 list 合并成 <ul><li>。
    去掉「• 」/「· 」前缀（HTML ul 自带 marker）。

    跟 _merge_numbered_into_ol 互补：那个处理「1./2./3.」编号 → ol，本函数处理无序 → ul。
    必须在 ol 合并之后跑，避免误吞 ol 内嵌套的 ul。
    """
    parts = re.split(r'(<p[^>]*>.*?</p>)', html, flags=re.DOTALL)
    p_re = re.compile(r'^<p[^>]*>(.*?)</p>$', re.DOTALL)
    bullet_re = re.compile(r'^\s*[•·‧]\s+', re.DOTALL)
    bullet_strip = re.compile(r'^((?:\s*<[^>]+>\s*)*)\s*[•·‧]\s+')

    out: list[str] = []
    buf: list[str] = []

    def flush():
        if buf:
            out.append('<ul>' + ''.join(f'<li>{x}</li>' for x in buf) + '</ul>')
            buf.clear()

    for part in parts:
        m = p_re.match(part)
        if not m:
            if part.strip() == '':
                if not buf:
                    out.append(part)
                continue
            flush()
            out.append(part)
            continue
        inner = m.group(1)
        plain = re.sub(r'<[^>]+>', '', inner).lstrip()
        if bullet_re.match(plain):
            content = bullet_strip.sub(r'\1', inner, count=1)
            buf.append(content)
        else:
            flush()
            out.append(part)
    flush()
    return ''.join(out)


def _merge_numbered_into_ol(html: str) -> str:
    """把 fill_cell_blocks 产出的 <p>N. xxx</p> 段合并成 <ol><li>...</li></ol>，
    并把紧随的 <p>- yyy</p> 二级缩进段嵌套为父 <li> 内的 <ul><li>。

    背景：Confluence storage format 会 strip 段落 inline style，padding-left 失效；
    旧实现对二级条目 fallback 写 padding-left 兜底，结果是缩进丢失 + 关掉 <ol>
    导致后续 N. 编号从 1 重新开始。重建为原生嵌套列表，让 Confluence 自带缩进 +
    自动编号同时生效。

    输入示例（mammoth 转出的形态）：
        <p>title</p>
        <p>1. main A</p>
        <p>- sub A1</p>
        <p>- sub A2</p>
        <p>2. main B</p>

    输出：
        <p>title</p>
        <ol>
          <li>main A<ul><li>sub A1</li><li>sub A2</li></ul></li>
          <li>main B</li>
        </ol>
    """
    parts = re.split(r'(<p[^>]*>.*?</p>)', html, flags=re.DOTALL)
    p_re = re.compile(r'^<p[^>]*>(.*?)</p>$', re.DOTALL)
    num_re = re.compile(r'^\s*\d+\.\s+', re.DOTALL)
    sub_re = re.compile(r'^\s*-\s+')
    num_strip = re.compile(r'^((?:\s*<[^>]+>\s*)*)\s*\d+\.\s+')
    sub_strip = re.compile(r'^((?:\s*<[^>]+>\s*)*)\s*-\s+')

    out: list[str] = []
    items: list[dict] = []  # [{'main': str, 'subs': [str, ...]}, ...]

    def flush():
        if not items:
            return
        lis = []
        for it in items:
            if it['subs']:
                ul = '<ul>' + ''.join(f'<li>{s}</li>' for s in it['subs']) + '</ul>'
                lis.append(f"<li>{it['main']}{ul}</li>")
            else:
                lis.append(f"<li>{it['main']}</li>")
        out.append('<ol>' + ''.join(lis) + '</ol>')
        items.clear()

    for part in parts:
        m = p_re.match(part)
        if not m:
            # 纯空白：active <ol> 内吞掉（避免空白挤进列表破坏结构），否则原样保留
            if part.strip() == '':
                if not items:
                    out.append(part)
                continue
            flush()
            out.append(part)
            continue
        inner = m.group(1)
        plain = re.sub(r'<[^>]+>', '', inner).lstrip()
        if num_re.match(plain):
            content = num_strip.sub(r'\1', inner, count=1)
            items.append({'main': content, 'subs': []})
        elif sub_re.match(plain):
            sub_content = sub_strip.sub(r'\1', inner, count=1)
            if items:
                items[-1]['subs'].append(sub_content)
            else:
                # 孤儿子条目：没有前置主条目时单独包成 <ul>
                out.append(f'<ul><li>{sub_content}</li></ul>')
        else:
            flush()
            out.append(part)
    flush()
    return ''.join(out)


def update_page(page_id: str, title: str, body: str, current_version: int):
    payload = {
        "version": {"number": current_version + 1},
        "title": title,
        "type": "page",
        "body": {"storage": {"value": body, "representation": "storage"}},
    }
    _rest("PUT", f"content/{page_id}",
          json=payload,
          extra_headers={"Content-Type": "application/json"})


def get_page_info(page_id: str) -> dict:
    """按 pageId 直接取当前页面 title + version + spaceKey,跳过 title 查找。
    模式 B 用;当 page 已知时比 get_or_create_page 稳(无同名页歧义)。"""
    info = _rest("GET", f"content/{page_id}", params={"expand": "version,space"})
    return {
        "id": info["id"],
        "title": info["title"],
        "version": info["version"]["number"],
        "space_key": info["space"]["key"],
    }


# ── Pre-flight:docx 结构体检 ─────────────────────────────────────────────────

def preflight_check_headings(docx_path: str) -> tuple:
    """扫 docx Heading 1/2 计数。为 0 时 stderr 打印 warning 但不中止。
    memory #1/#3 踩坑:gen 脚本 BASE 路径错时 h1/h2 函数无声失败,
    所有章节都是 Normal style,推 Confluence 后大纲树失效。
    事前 5 秒检查省事后几小时排查。
    返回 (h1_count, h2_count)。"""
    try:
        from docx import Document
    except ImportError:
        print("  ⚠ 跳过 Heading 体检(python-docx 未安装)", file=sys.stderr)
        return (-1, -1)
    doc = Document(docx_path)
    h1 = sum(1 for p in doc.paragraphs if (p.style and p.style.name == "Heading 1"))
    h2 = sum(1 for p in doc.paragraphs if (p.style and p.style.name == "Heading 2"))
    print(f"🔍 pre-flight: Heading 1={h1}, Heading 2={h2}")
    if h1 == 0 and h2 == 0:
        print(
            "⚠️  docx 里没有任何 Heading 1/2!推到 Confluence 后左侧大纲树会失效。\n"
            "   常见原因:① gen 脚本 sys.path 路径错,h1/h2 函数无声失败(memory #1)\n"
            "            ② 章节全部用 add_paragraph 直接写,没调 h1/h2\n"
            "   先跑 normalize_headings(doc) 归一化或修 gen 脚本,再 push。",
            file=sys.stderr,
        )
    elif h1 == 0:
        print("⚠️  docx 里没有 Heading 1,只有 H2 — Confluence 大纲树会缺顶层。", file=sys.stderr)
    return (h1, h2)


def gate_check_quality(docx_path: str) -> int:
    """推 wiki 前硬闸：扫 docx 里的圈数字 + 裸场景编号 + 决策编号。
    任一违规即拒绝推送（非 0 退出码）。--skip-self-check 可绕过。
    与 check_prd.sh 第 2 段 python 检查同源逻辑，独立实现避免依赖外部脚本 / grep -P。"""
    try:
        from docx import Document
    except ImportError:
        print("  ⚠ 跳过质量闸（python-docx 未安装）", file=sys.stderr)
        return 0

    doc = Document(docx_path)
    CIRCLE = re.compile(r'[①-⑳⓫-⓿]')  # ①-⑳ + ⓫-⓿
    SCENE_ID = re.compile(r'[A-G]-\d+(?:\s*/\s*[A-G]-\d+){0,5}')
    DECISION = re.compile(r'决策\s*\d+')

    circle_hits, scene_hits, decision_hits = [], [], []
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                txt = cell.text
                for m in CIRCLE.finditer(txt):
                    circle_hits.append(f'T{ti}R{ri}C{ci}: {m.group()}')
                if ci == 0:
                    continue  # scene_table 第 0 列是 anchor，允许编号
                if ti >= 3 and SCENE_ID.search(txt):
                    scene_hits.append(f'T{ti}R{ri}C{ci}: {SCENE_ID.search(txt).group()}')
                if DECISION.search(txt):
                    decision_hits.append(f'T{ti}R{ri}C{ci}: {DECISION.search(txt).group()}')
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if CIRCLE.search(p.text):
            circle_hits.extend(f'P[{style}]: {m.group()}' for m in CIRCLE.finditer(p.text))
        if not style.startswith('Heading') and DECISION.search(p.text):
            decision_hits.append(f'P[{style}]: {DECISION.search(p.text).group()}')
        # 段落级裸场景编号扫描（与表格级互补，封堵 add_p / bullet 盲区）
        if not style.startswith('Heading') and SCENE_ID.search(p.text):
            scene_hits.append(f'P[{style}]: {SCENE_ID.search(p.text).group()} ...')

    # 1.3 变更范围流水账启发式（限 T2）：左列含迭代过程标记
    # 全文迭代痕迹（PRD 是静态描述，不应含「04-28 决策 / N 月 N 日决策 / 改名为 / 改为 X」之类来历）
    # 词集统一至 humanize/patterns.py，与 check_prd.sh / save_prd 共享
    from humanize import (
        PRD_CHANGELOG_BODY_HISTORY as BODY_HISTORY,
        PRD_CHANGELOG_LANE_HISTORY as LANE_HISTORY,
    )

    # 讲人话扫描（流水账日期 / snake_case 字段 / CSS 实现细节）— 共享模块
    # 完整规则 + 豁免见 references/prd-human-voice.md
    from humanize import scan_human_voice
    voice_result = scan_human_voice(doc)
    date_tag_hits = voice_result['date_tag_hits']
    snake_field_hits = voice_result['snake_field_hits']
    css_impl_hits = voice_result['css_impl_hits']

    history_hits, body_history_hits = [], []
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            if ri == 0:
                continue
            if not row.cells:
                continue
            # 1.3 变更范围左列流水账（限 T2）
            if ti == 2:
                left_col = row.cells[0].text
                if LANE_HISTORY.search(left_col):
                    history_hits.append(f'R{ri}: {left_col[:60]}')
            # 仅扫旧的 BODY_HISTORY（迭代痕迹），其他三类已由 scan_human_voice 处理
            if ti >= 3:
                for ci, cell in enumerate(row.cells):
                    m = BODY_HISTORY.search(cell.text)
                    if m:
                        body_history_hits.append(f'T{ti}R{ri}C{ci}: {m.group()} ...')
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if style.startswith('Heading'):
            continue
        m = BODY_HISTORY.search(p.text)
        if m:
            body_history_hits.append(f'P[{style}]: {m.group()} ...')

    fail = 0
    if circle_hits:
        print(f"❌ 圈数字残留 {len(circle_hits)} 处（CLAUDE.md 禁止 ①②③，统一 1./2./3.）：", file=sys.stderr)
        for h in circle_hits[:5]:
            print(f"   {h}", file=sys.stderr)
        fail = 1
    if scene_hits:
        print(f"❌ 正文裸场景编号 {len(scene_hits)} 处（pm-workflow.md「讲人话」强制，仅第 2 章场景地图允许）：", file=sys.stderr)
        for h in scene_hits[:5]:
            print(f"   {h}", file=sys.stderr)
        fail = 1
    if decision_hits:
        print(f"❌ 正文决策编号 {len(decision_hits)} 处（PRD SKILL.md 禁止「决策 N」）：", file=sys.stderr)
        for h in decision_hits[:5]:
            print(f"   {h}", file=sys.stderr)
        fail = 1
    # 讲人话三类：流水账（FAIL）/ snake_case（WARN）/ CSS（WARN）
    from humanize import report_violations
    if report_violations(voice_result):
        fail = 1
    if history_hits:
        print(f"⚠️  1.3 变更范围疑似流水账 {len(history_hits)} 行（左列含迭代过程标记，应是 vs 线上基线 delta）：", file=sys.stderr)
        for h in history_hits[:5]:
            print(f"   {h}", file=sys.stderr)
        print(f"   1.3 锚点 = vs **线上基线** delta；多轮迭代过程归 context.md 第 7 章。误报可加 --skip-self-check", file=sys.stderr)
        fail = 1
    # 以下两类是软提醒（误报风险高于价值，不阻断推送）：
    if body_history_hits:
        print(f"⚠️  正文疑似迭代痕迹 {len(body_history_hits)} 处（PRD 应描述最终态，不写「N 月 N 日决策 / 从 X 改名」来历）：", file=sys.stderr)
        for h in body_history_hits[:3]:
            print(f"   {h}", file=sys.stderr)
        print(f"   迭代过程归 context.md 第 7 章；本提醒不阻断推送", file=sys.stderr)
    # css_impl_hits 已在 report_violations 内打印

    if fail:
        print("\n💡 修复后重推；如确认刻意保留可加 --skip-self-check 绕过", file=sys.stderr)
    else:
        print(f"✅ 质量闸通过（圈数字 / 裸编号 / 决策编号 三项 0 违规）")
    return fail


# ── 主流程 ───────────────────────────────────────────────────────────────────

def _build_parser():
    p = argparse.ArgumentParser(
        description="推送 PRD docx 到 Confluence Server(两种模式)",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__,
    )
    p.add_argument("docx", help="PRD docx 文件路径")
    p.add_argument("--page-id", dest="page_id", default=None,
                   help="已知 pageId,直接按 ID 覆盖(模式 B);与 --space/--parent-id 互斥")
    p.add_argument("--title", default=None,
                   help="页面标题。模式 A 必填;模式 B 可选(不传则保留 Confluence 现有 title)")
    p.add_argument("--space", dest="space_key", default=None,
                   help="Confluence spaceKey(模式 A 必填)")
    p.add_argument("--parent-id", dest="parent_id", default=None,
                   help="父页面 pageId(模式 A 可选)")
    p.add_argument("--img-width", type=int, default=None,
                   help="所有图片统一宽度(像素)。不传时按宽高比智能选:phone 截图 300 / web 截图 600")
    p.add_argument("--img-width-web", type=int, default=600,
                   help="web 截图宽度(横屏或方形),默认 600")
    p.add_argument("--img-width-phone", type=int, default=300,
                   help="phone 截图宽度(高瘦比例 h>w*1.3),默认 300")
    p.add_argument("--skip-preflight", action="store_true",
                   help="跳过 Heading 体检(不推荐,除非明知 docx 故意无 H1/H2)")
    p.add_argument("--skip-self-check", action="store_true",
                   help="跳过质量闸（圈数字 / 裸场景编号 / 决策编号 三项）。不推荐，违规会污染 wiki")
    return p


def _resolve_from_pages_map(docx_path: str, cfg: dict) -> dict:
    """多 PRD 项目:cfg.pages 按 docx basename 精确匹配,返回 {'page_id', 'title'} 或空 dict。
    没命中时返回空,让上层走 prd_page_id 单 PRD 兜底。"""
    pages = cfg.get("pages") or {}
    if not pages:
        return {}
    basename = Path(docx_path).name
    if basename in pages:
        return pages[basename]
    # 模糊匹配:文件名含 key 子串(如 key="V2.8" 命中 "...V2.8-Phase1.docx")
    for key, val in pages.items():
        if key in basename:
            return val
    return {}


def _resolve_mode(args, cfg: dict, docx_path: str):
    """返回 (page_id, title, version, space_key) —— 无论模式 A/B 都归一化为这 4 个字段。
    CLI 参数优先于 cfg.pages[docx 匹配项],再 fallback 到 cfg 顶层 prd_page_id/title。"""
    matched = _resolve_from_pages_map(docx_path, cfg)
    page_id = args.page_id or matched.get("page_id") or cfg.get("prd_page_id")
    title = args.title or matched.get("title") or cfg.get("title")
    space_key = args.space_key or cfg.get("space") or DEFAULT_SPACE
    parent_id = args.parent_id or cfg.get("parent_id")

    if page_id:
        info = get_page_info(page_id)
        return info["id"], title or info["title"], info["version"], info["space_key"]
    # 模式 A:按 title upsert
    if not title:
        sys.exit("❌ 模式 A 必须传 --title (或 .confluence.json 含 title);也可用 --page-id 走模式 B")
    pid, version = get_or_create_page(space_key, title, parent_id)
    return pid, title, version, space_key


def main(argv=None):
    # 向后兼容:旧位置参数 `docx title space [parent]` 仍能工作
    if argv is None:
        argv = sys.argv[1:]
    if argv and not argv[0].startswith("-") and len(argv) >= 3 and not any(
        a.startswith("--") for a in argv[:4]
    ):
        # 位置参数模式:[docx, title, space, parent_id?]
        legacy = ["--title", argv[1], "--space", argv[2]]
        if len(argv) >= 4:
            legacy += ["--parent-id", argv[3]]
        argv = [argv[0]] + legacy

    args = _build_parser().parse_args(argv)

    if not Path(args.docx).exists():
        sys.exit(f"❌ 文件不存在: {args.docx}")

    if not args.skip_preflight:
        preflight_check_headings(args.docx)

    if not args.skip_self_check:
        if gate_check_quality(args.docx) != 0:
            sys.exit("❌ 推送已拒绝（质量闸未过）。修复后重推，或加 --skip-self-check 强制推。")

    cfg_path = _find_project_config(args.docx)
    cfg = _load_project_config(cfg_path)
    if cfg:
        print(f"📦 读取项目配置: {cfg_path}")

    print(f"📄 转换: {args.docx}")
    page_id, title, version, space_key = _resolve_mode(args, cfg, args.docx)
    print(f"📑 页面 ID: {page_id}  标题: {title}  space: {space_key}  当前版本: {version}")

    # --img-width 兼容：传单值 → 所有图统一；不传 → phone/web 各自走 default
    if args.img_width is not None:
        web_w = phone_w = args.img_width
    else:
        web_w = args.img_width_web
        phone_w = args.img_width_phone
    body = convert_docx(args.docx, page_id,
                        img_width_web=web_w, img_width_phone=phone_w)

    print("📤 更新页面内容...")
    update_page(page_id, title, body, version)

    # 推送成功后写回项目配置,下次 session 不用再问 pageId
    new_cfg = {**cfg, "space": space_key}
    if args.parent_id or cfg.get("parent_id"):
        new_cfg["parent_id"] = args.parent_id or cfg.get("parent_id")
    # 多 PRD 项目(已有 pages 字段)写到 pages[basename];单 PRD 项目写顶层 prd_page_id
    if cfg.get("pages"):
        basename = Path(args.docx).name
        pages = dict(new_cfg.get("pages") or {})
        # 找已有 key (可能是模糊匹配命中);找不到就用 basename 新增一项
        match_key = basename if basename in pages else next(
            (k for k in pages if k in basename), basename
        )
        pages[match_key] = {"page_id": page_id, "title": title}
        new_cfg["pages"] = pages
    else:
        new_cfg["prd_page_id"] = page_id
        new_cfg["title"] = title
    if new_cfg != cfg:
        _save_project_config(cfg_path, new_cfg)
        if cfg_path:
            print(f"💾 已写回 {cfg_path}")

    page_url = f"{base_url()}/pages/viewpage.action?pageId={page_id}"
    print(f"✅ 完成: {page_url}")


if __name__ == "__main__":
    main()

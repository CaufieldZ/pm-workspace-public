"""docx 文档级 helper：初始化 + 封面 + save_prd 一键保存（自动跑全套规范化）。

依赖：
- core.styles  → C, para_run, cell_text
- core.normalize → normalize_punctuation, normalize_fonts, ensure_theme
- humanize → humanize_prd_voice, scan_human_voice, report_violations
- sections（save_prd 内 lazy import）→ scan_completeness, report_completeness

save_prd 是 PRD 的一键保存入口：humanize → punctuation → fonts → save → ensure_theme
→ scan_human_voice → scan_completeness（按档位）。各步可独立关闭。
"""
import json
import os
import pathlib

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

from core.styles import C, para_run, cell_text


# ── 元信息 ────────────────────────────────────────────────────────────────

def get_author() -> str:
    cfg = pathlib.Path(__file__).resolve().parents[3] / ".claude" / "skills" / "_shared" / "workspace.json"
    try:
        return json.loads(cfg.read_text(encoding='utf-8')).get("author", "")
    except Exception:
        return ""


# ── 文档初始化 ────────────────────────────────────────────────────────────

def init_doc(landscape=True):
    """创建并返回一个配置好页面的 Document 对象"""
    doc = Document()
    section = doc.sections[0]
    if landscape:
        section.page_width  = Cm(27.94)
        section.page_height = Cm(21.59)
        section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)
    section.top_margin    = Cm(1.905)
    section.bottom_margin = Cm(1.905)
    return doc


# ── 封面 ──────────────────────────────────────────────────────────────────

def cover_page(doc, title, subtitle="产品需求文档（PRD）", scope="", meta_rows=None):
    """生成封面页
    meta_rows: list of [label, value] 用于文档属性表
    """
    if meta_rows is None:
        meta_rows = []
    _author = get_author()
    if _author and not any(r[0] == "作者" for r in meta_rows):
        meta_rows = [["作者", _author]] + list(meta_rows)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    p.paragraph_format.space_after = Pt(4)
    para_run(p, title, font='title', size_pt=24, bold=True, color=C["textHeading"])

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    para_run(p, subtitle, font='title', size_pt=14, color=C["textSecondary"])

    if scope:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(30)
        para_run(p, scope, size_pt=10, color=C["textMuted"], italic=True)

    if meta_rows:
        meta_tbl = doc.add_table(rows=len(meta_rows), cols=2)
        meta_tbl.style = 'Table Grid'
        meta_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, value) in enumerate(meta_rows):
            cell_text(meta_tbl.rows[i].cells[0], label, size_pt=9, bold=True,
                      color=C["textSecondary"], align=WD_ALIGN_PARAGRAPH.RIGHT)
            meta_tbl.rows[i].cells[0].width = Cm(4)
            cell_text(meta_tbl.rows[i].cells[1], value, size_pt=9,
                      color=C["textPrimary"])
            meta_tbl.rows[i].cells[1].width = Cm(14)


# ── save_prd: 一键保存 + 全套规范化（gen / update 通用）──────────────────

def save_prd(doc, path, *,
             tier='standard',
             extra_jargon=None,
             scene_table_indices=None,
             humanize=True,
             punctuation=True,
             fonts=True,
             theme=True,
             scan=True,
             structural=True,
             completeness=True,
             scene_count=None,
             screenshots_fresh=False,
             prototype_html=None,
             shot_dir=None):
    """保存 PRD docx，自动应用「讲人话 + 标点 + 字体 + theme + 完整性 + 结构性」全套规范化。

    Args:
        doc: python-docx Document
        path: 输出路径（str 或 Path）
        tier: 档位，'mini' / 'standard' / 'full'。极简档跳过完整性 + 结构性体检。
              详见 references/prd-template.md「何时用哪档」
        extra_jargon: 项目特定字段 → 中文映射，list of (re.Pattern, str)
        scene_table_indices: 仅对这些表索引应用 humanize（None = range(5, len(tables))）
        humanize / punctuation / fonts / theme / scan / structural / completeness:
            各步开关，默认全开（mini 档自动关 structural / completeness）
        scene_count: 显式场景数（升版路径已知时传入，缺省自动从 doc 数 H2）
        screenshots_fresh: True 时调 assert_screenshots_fresh（升版路径推荐传 True，
            截图过期 raise；新建路径默认 False）
        prototype_html / shot_dir: screenshots_fresh=True 时必传

    体检（stderr 输出，不抛异常 — 自检 / push gate 兜底阻断）：
        - humanize.scan_human_voice：流水账 / snake_case / CSS 实现细节
        - humanize.scan_prd_structural：圈数字 / 决策编号 / 占位符 / 1.3 流水账 /
          章节用户故事引言 / Scene 右列层次 / 段落表格数 / DPI / theme / 老字体
        - sections.scan_completeness：按档位扫 1.2 反向指标 / 1.5 Non-goals /
          6.x Open Questions / 1.6 方案选型（仅 full）/ 7.x Risks（仅 full）
    """
    from humanize import humanize_prd_voice as _humanize_prd_voice
    from core.normalize import normalize_punctuation, normalize_fonts, ensure_theme

    if humanize:
        _humanize_prd_voice(doc, scene_table_indices=scene_table_indices,
                            extra_jargon=extra_jargon)
    if punctuation:
        normalize_punctuation(doc)
    if fonts:
        normalize_fonts(doc)
    doc.save(str(path))
    if theme:
        ensure_theme(str(path))

    if screenshots_fresh:
        if not (prototype_html and shot_dir):
            raise ValueError(
                'save_prd(screenshots_fresh=True) 必须传 prototype_html + shot_dir'
            )
        from screenshots import assert_screenshots_fresh
        assert_screenshots_fresh(str(path), prototype_html, shot_dir)

    if scan:
        from humanize import scan_human_voice, report_violations
        result = scan_human_voice(doc)
        report_violations(result)  # stderr，不抛异常（自检/push gate 兜底阻断）

    if structural and tier != 'mini':
        from humanize import scan_prd_structural, report_structural_violations
        struct_result = scan_prd_structural(doc, scene_count=scene_count, docx_path=str(path))
        report_structural_violations(struct_result)  # stderr，不抛异常

    if completeness and tier != 'mini':
        from sections import scan_completeness, report_completeness
        completeness_result = scan_completeness(doc, tier=tier)
        report_completeness(completeness_result)  # stderr warn，不阻断

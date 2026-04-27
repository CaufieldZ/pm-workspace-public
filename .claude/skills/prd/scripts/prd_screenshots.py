#!/usr/bin/env python3
# PM-Workspace | (c) 2026 CaufieldZ | Apache 2.0 + AI Training Restriction
"""PRD 截图回填 — 通用脚本。

从 imap HTML 截取各 Scene 设备框，圆角处理后插入 PRD docx 对应表格。

用法：
    python3 prd_screenshots.py --project demo-private-fund
    python3 prd_screenshots.py --project htx-community --imap deliverables/imap-foo-v2.html
    python3 prd_screenshots.py --project foo --docx deliverables/PRD_v2.docx
    python3 prd_screenshots.py --project foo --scenes A-1,B-1,C-1
    python3 prd_screenshots.py --project foo --shot-only    # 只截图不插入
    python3 prd_screenshots.py --project foo --insert-only  # 只插入已有截图

覆盖范围：
    - imap HTML 标准结构：fade-section[id="scene-*"] 内含 .phone / .webframe
    - prototype HTML：同上结构（若用了标准骨架也适用）
    - 不覆盖：使用自定义 JS API 的 prototype（htx-activity-center 类型），
      那类留在各项目 scripts/ 目录维护

Scene → docx 表格匹配策略：
    遍历 doc.tables，找首行首列 cell.text 包含场景编号的表格（如 "A-1"、"M-2"），
    按 scene_id 正则匹配，无需 hardcode table index。
"""

import argparse, os, re, sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent.parent.parent.parent  # pm-workspace


# ── 工具函数 ──────────────────────────────────────────────────────────────────

def round_phone_corners(png_path: str, radius_px: int = 72):
    """给 phone 截图加透明圆角蒙版。
    CSS border-radius 36px × deviceScaleFactor 2 = 72px。"""
    from PIL import Image, ImageDraw
    img = Image.open(png_path).convert("RGBA")
    mask = Image.new("L", img.size, 0)
    ImageDraw.Draw(mask).rounded_rectangle([(0, 0), img.size], radius=radius_px, fill=255)
    img.putalpha(mask)
    img.save(png_path, "PNG")


def find_latest(directory: Path, glob: str) -> Path | None:
    files = sorted(directory.glob(glob), key=lambda p: p.stat().st_mtime, reverse=True)
    return files[0] if files else None


def discover_scenes_from_html(html_path: Path) -> list[tuple[str, str]]:
    """从 imap/proto HTML 解析 scene 列表，返回 [(scene_id, device_type), ...]。
    device_type: 'phone' | 'web'
    通过检查 scene 容器内是否含 .phone 或 .webframe 判断。"""
    text = html_path.read_text(encoding="utf-8")
    # 找所有 fade-section id
    scene_ids = re.findall(r'<div[^>]+class="[^"]*fade-section[^"]*"[^>]+id="(scene-[^"]+)"', text)
    if not scene_ids:
        # fallback: 任意带 id="scene-*" 的块
        scene_ids = re.findall(r'id="(scene-[a-z0-9-]+)"', text)

    results = []
    for sid in scene_ids:
        # 粗略判断：找到 scene 容器后看里面第一个 .phone / .webframe
        pattern = rf'id="{re.escape(sid)}".*?(?=id="scene-|</body>)'
        block = re.search(pattern, text, re.DOTALL)
        device = "phone"
        if block:
            block_text = block.group(0)
            if 'class="webframe' in block_text or "class='webframe" in block_text:
                device = "web"
            elif 'class="phone' not in block_text and "class='phone" not in block_text:
                device = "web"
        results.append((sid, device))
    return results


# ── 截图 ──────────────────────────────────────────────────────────────────────

def take_screenshots(html_path: Path, out_dir: Path,
                     scenes: list[tuple[str, str]] | None = None) -> dict[str, Path]:
    """打开 HTML，对每个 scene 截设备框，返回 {scene_id: png_path}。"""
    from playwright.sync_api import sync_playwright

    out_dir.mkdir(parents=True, exist_ok=True)
    url = html_path.as_uri()

    if scenes is None:
        scenes = discover_scenes_from_html(html_path)
    if not scenes:
        sys.exit(f"错误：在 {html_path.name} 里找不到 scene-* 容器")

    shots: dict[str, Path] = {}

    with sync_playwright() as pw:
        browser = pw.chromium.launch(headless=True)
        page = browser.new_page(
            viewport={"width": 1440, "height": 900},
            device_scale_factor=2,
        )
        page.goto(url)
        page.wait_for_load_state("networkidle")

        # 让所有 fade-section 可见，隐藏标注层（避免被截进图里）
        page.evaluate("""
            document.querySelectorAll('.fade-section').forEach(el => el.classList.add('visible'));
            document.querySelectorAll(
              '.anno, .anno-n, .ann-card, .ann-tag, .aw, .flow-note, .side-nav, .scene-nav'
            ).forEach(el => { el.style.display = 'none'; });
        """)
        page.wait_for_timeout(600)

        for scene_id, device_type in scenes:
            scene_loc = page.locator(f"#{scene_id}").first
            if not scene_loc.is_visible():
                print(f"  ⊘ {scene_id}: 容器不可见，跳过")
                continue
            scene_loc.scroll_into_view_if_needed()
            page.wait_for_timeout(200)

            selector = ".phone" if device_type == "phone" else ".webframe"
            device_loc = scene_loc.locator(selector).first
            if not device_loc.is_visible():
                # 再试另一种
                alt = ".webframe" if device_type == "phone" else ".phone"
                device_loc = scene_loc.locator(alt).first
                if device_loc.is_visible():
                    device_type = "web" if device_type == "phone" else "phone"
                else:
                    print(f"  ⊘ {scene_id}: 找不到 .phone / .webframe")
                    continue

            # scene_id 格式 "scene-a-1" → 文件名 "a-1.png"
            fname = scene_id.removeprefix("scene-") + ".png"
            out_path = out_dir / fname
            device_loc.screenshot(path=str(out_path), omit_background=True)

            if device_type == "phone":
                round_phone_corners(str(out_path))

            shots[scene_id] = out_path
            print(f"  截图 {scene_id} ({device_type}) → {fname}")

        browser.close()

    return shots


# ── 插入 docx ─────────────────────────────────────────────────────────────────

_SCENE_ID_RE = re.compile(r'\b([A-Z])-(\d+[a-z]?)\b')
_SCENE_TITLE_RE = re.compile(r'(📱\s*[A-Z]-\d+[a-z]?\s*·\s*[^\n]+)')


def _normalize_scene_id(raw: str) -> str:
    """把 scene-a-1 / A-1 / a-1 统一成 A-1 大写形式。"""
    raw = raw.removeprefix("scene-")
    parts = raw.split("-", 1)
    if len(parts) == 2:
        return f"{parts[0].upper()}-{parts[1]}"
    return raw.upper()


def _cell_scene_ids(cell) -> set[str]:
    text = cell.text.strip()
    return {f"{m.group(1)}-{m.group(2)}" for m in _SCENE_ID_RE.finditer(text)}


def insert_into_docx(docx_path: Path, shots: dict[str, Path],
                     width_phone_cm: float = 5.0, width_web_cm: float = 7.0):
    """把截图插入 PRD docx，按场景编号匹配表格首行首列。

    自动检测 cell 首段是否含 scene_table 标识（📱 + 编号），有则保留场景标题段
    （memory feedback_prd_iter_screenshot_pitfalls.md 坑 2：默认 replace_cell_image
    会清掉 scene_table 写在 cell 首段的「📱 X-N · 名称」标识，重推后看不出图属于哪个场景）。
    """
    import sys as _sys
    _sys.path.insert(0, str(ROOT / ".claude/skills/prd/scripts"))
    from update_prd_base import replace_cell_image, replace_cell_image_keep_title

    from docx import Document
    doc = Document(str(docx_path))

    # 建立 scene_id → png_path 映射（大写归一化）
    shots_norm = {_normalize_scene_id(k): v for k, v in shots.items()}
    inserted: set[str] = set()

    for table in doc.tables:
        if not table.rows:
            continue
        cell = table.rows[0].cells[0]
        cell_ids = _cell_scene_ids(cell)
        for sid in cell_ids:
            if sid in shots_norm and sid not in inserted:
                img_path = shots_norm[sid]
                # 判断宽度：phone 截图文件名不含 web，且文件尺寸比较窄
                is_phone = "phone" in str(img_path) or _is_portrait(img_path)
                width = width_phone_cm if is_phone else width_web_cm
                # scene_table 写在 cell 首段的「📱 X-N · ...」标题必须保留
                first_para_text = cell.paragraphs[0].text.strip() if cell.paragraphs else ""
                title_match = _SCENE_TITLE_RE.match(first_para_text)
                if title_match:
                    replace_cell_image_keep_title(cell, title_match.group(1), str(img_path), width_cm=width)
                else:
                    replace_cell_image(cell, str(img_path), width_cm=width)
                print(f"  插入 {sid} → Table ({table.rows[0].cells[0].text[:20].strip()!r})")
                inserted.add(sid)
                break

    doc.save(str(docx_path))
    print(f"\n  已保存: {docx_path.name}")
    missing = set(shots_norm) - inserted
    if missing:
        print(f"  未匹配场景（docx 里没找到对应表格）: {', '.join(sorted(missing))}")
    return inserted


def _is_portrait(png_path: Path) -> bool:
    """粗判断是否竖屏（phone），高 > 宽时认为是。"""
    try:
        from PIL import Image
        w, h = Image.open(str(png_path)).size
        return h > w
    except Exception:
        return False


# ── 主流程 ────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="PRD 截图回填通用脚本")
    parser.add_argument("--project", "-p", required=True, help="项目名（projects/ 下目录名）")
    parser.add_argument("--imap", help="imap/proto HTML 路径（默认：自动找 deliverables/imap-*.html）")
    parser.add_argument("--docx", help="PRD docx 路径（默认：自动找 deliverables/*PRD*.docx）")
    parser.add_argument("--scenes", help="限定场景列表，逗号分隔，如 A-1,B-2（默认：全部）")
    parser.add_argument("--shot-only", action="store_true", help="只截图，不插入 docx")
    parser.add_argument("--insert-only", action="store_true", help="只插入已有截图，跳过截图步骤")
    parser.add_argument("--width-phone", type=float, default=5.0, help="phone 截图宽度 cm（默认 5.0）")
    parser.add_argument("--width-web", type=float, default=7.0, help="web 截图宽度 cm（默认 7.0）")
    args = parser.parse_args()

    proj_dir = ROOT / "projects" / args.project
    if not proj_dir.is_dir():
        sys.exit(f"错误：项目目录不存在 {proj_dir}")

    deliverables = proj_dir / "deliverables"
    shot_dir = proj_dir / "screenshots" / "prd"

    # ── 确定 HTML 源 ──
    html_path = None
    if not args.insert_only:
        if args.imap:
            html_path = Path(args.imap) if Path(args.imap).is_absolute() else proj_dir / args.imap
        else:
            html_path = find_latest(deliverables, "imap-*.html")
            if not html_path:
                html_path = find_latest(deliverables, "proto-*.html")
            if not html_path:
                html_path = find_latest(deliverables, "*.html")
        if not html_path or not html_path.exists():
            sys.exit(f"错误：找不到 imap/proto HTML，请用 --imap 指定路径")
        print(f"HTML 源: {html_path.relative_to(ROOT)}")

    # ── 确定 docx ──
    docx_path = None
    if not args.shot_only:
        if args.docx:
            docx_path = Path(args.docx) if Path(args.docx).is_absolute() else proj_dir / args.docx
        else:
            docx_path = find_latest(deliverables, "*PRD*.docx")
            if not docx_path:
                docx_path = find_latest(deliverables, "*prd*.docx")
        if not docx_path or not docx_path.exists():
            sys.exit(f"错误：找不到 PRD docx，请用 --docx 指定路径")
        print(f"PRD:   {docx_path.relative_to(ROOT)}")

    # ── 限定 scenes ──
    scene_filter: set[str] | None = None
    if args.scenes:
        scene_filter = {_normalize_scene_id(s.strip()) for s in args.scenes.split(",")}

    # ── 截图 ──
    shots: dict[str, Path] = {}
    if not args.insert_only:
        all_scenes = discover_scenes_from_html(html_path)
        if scene_filter:
            all_scenes = [(sid, dt) for sid, dt in all_scenes
                          if _normalize_scene_id(sid) in scene_filter]
        print(f"\n截图 {len(all_scenes)} 个 Scene...")
        shots = take_screenshots(html_path, shot_dir, scenes=all_scenes)
    else:
        # insert-only：从 shot_dir 加载已有截图
        if not shot_dir.exists():
            sys.exit(f"错误：截图目录不存在 {shot_dir}，请先运行截图步骤")
        for png in shot_dir.glob("*.png"):
            # 文件名 a-1.png → scene-a-1
            sid = "scene-" + png.stem
            shots[sid] = png
        if scene_filter:
            shots = {k: v for k, v in shots.items()
                     if _normalize_scene_id(k) in scene_filter}
        print(f"加载已有截图 {len(shots)} 张...")

    if not shots:
        sys.exit("没有可用截图，退出")

    if args.shot_only:
        print(f"\n截图完成，共 {len(shots)} 张，存于 {shot_dir.relative_to(ROOT)}")
        return

    # ── 插入 docx ──
    print(f"\n插入 docx...")
    inserted = insert_into_docx(docx_path, shots,
                                width_phone_cm=args.width_phone,
                                width_web_cm=args.width_web)
    print(f"\n完成：{len(inserted)}/{len(shots)} 个场景已插入")


if __name__ == "__main__":
    main()

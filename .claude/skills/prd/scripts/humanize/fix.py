"""自动修复（write 端，gen / update 时调用）。

PRD 讲人话抛光：流水账清理 + 代码字段 → 白话 + UI 设计参数砍掉 + 视觉问题修复。
详细规则与例外见 `references/prd-human-voice.md`。
"""
import re

from core.patch import replace_para_text
from core.styles import C, set_cell_bg
from docx.oxml.ns import qn

from .patterns import (
    DARK_FILLS_KEEP_WHITE,
    EMPTY_BULLET_RE,
    H2_V_TAG,
    LEGACY_BLUE_FILLS,
    NUMBERED_LINE_RE,
    PRD_CHANGELOG_PATTERNS,
    PRD_JARGON_REPLACEMENTS,
    PRD_KILL_BULLET_KEYWORDS,
    PRD_TRAILING_JUNK_PATTERNS,
    PRD_UI_STRIP_PATTERNS,
    W_NS,
    ZOMBIE_HEADING,
)


def _is_kill_bullet(text: str) -> bool:
    t = text.strip()
    if not re.match(r'^\d+\.\s*', t):
        return False
    return any(k in t for k in PRD_KILL_BULLET_KEYWORDS)


def _kill_dead_bullets(cell):
    """删 cell 内：① 含 字体/字号/砍掉/反转说明 关键词的整段；② 删后空 N. 编号段。"""
    for p in list(cell.paragraphs):
        t = p.text
        if _is_kill_bullet(t) or EMPTY_BULLET_RE.match(t.strip()):
            replace_para_text(p, '')


def _renumber_cell_bullets(cell):
    """同一 section（H 标题段之间）的 N. bullet 重排为 1./2./3./...。"""
    counter = 0
    for p in cell.paragraphs:
        t = p.text.strip()
        m = NUMBERED_LINE_RE.match(t)
        if m and m.group(2):
            counter += 1
            if int(m.group(1)) != counter:
                replace_para_text(p, f'{counter}. {m.group(2)}')
        else:
            counter = 0


def _cleanup_blank_paragraphs(cell):
    """删 cell 内连续空段（保留单空段做 section 分隔）。"""
    paras = list(cell.paragraphs)
    prev_empty = False
    for p in paras:
        if not p.text.strip():
            if prev_empty:
                p._element.getparent().remove(p._element)
            prev_empty = True
        else:
            prev_empty = False


def _clean_text(text: str, *pattern_groups) -> str:
    out = text
    for group in pattern_groups:
        for pat, repl in group:
            out = pat.sub(repl, out)
    return out.strip()


# scene_table 右列 list 项 bold 剥离规则（V2.7-era 残留：title + list 项全 bold；
# 新规范「title 粗体 + list 项非粗体」，仅识别 list 段去 bold，title 段保留）
# list 判定：以 「N.」/「N、」/ emoji / 「- 」/「· 」/「• 」开头
_LIST_PREFIX = re.compile(
    r'^\s*(?:\d+[\.、)]|[\-•·▪◦]\s|'
    r'[\U0001F300-\U0001FAFF\U00002600-\U000027BF])'
)


def humanize_prd_voice(doc, scene_table_indices=None, extra_jargon=None,
                       skip_chapters=None):
    """PRD 讲人话抛光：流水账清理 + 代码字段 → 白话 + UI 设计参数砍掉。

    详细规则与例外见 `references/prd-human-voice.md`。

    Args:
        doc: python-docx Document
        scene_table_indices: 仅对这些表索引应用「讲人话 + 砍 UI」，None 时 = range(5, len(doc.tables))。
            前 5 张通常是封面/角色/场景地图/CMS 概览，不需要 jargon 替换。
        extra_jargon: list of (re.Pattern, str) — 项目特定的字段映射（如 futures_holding → 合约持仓）。
            会在通用 PRD_JARGON_REPLACEMENTS 之后追加应用。
        skip_chapters: list of str — H2 标题包含这些关键词的章节豁免（如「枚举值」「字段」「埋点」）。
            未传时默认豁免：枚举值 / 字段对照 / 核心 ID / 埋点 / 事件。

    使用示例（项目 patch 脚本）：

        from humanize import humanize_prd_voice
        import re

        PROJECT_JARGON = [
            (re.compile(r'`futures_holding`|futures_holding'), '合约持仓'),
            (re.compile(r'`spot_holding`|spot_holding'), '现货持仓'),
        ]
        humanize_prd_voice(doc, scene_table_indices=range(5, 19),
                           extra_jargon=PROJECT_JARGON)
    """
    if scene_table_indices is None:
        scene_table_indices = range(5, len(doc.tables))
    jargon = list(PRD_JARGON_REPLACEMENTS)
    if extra_jargon:
        # 项目映射放前面，避免被通用规则提前消耗（如 futures_holding 含 holding 子串）
        jargon = list(extra_jargon) + jargon

    # 1a. 物理删除僵尸 Scene（H2/H3 含「砍掉/已删除/废弃」→ 删段 + 紧跟 scene_table）
    # 必须先于流水清理：V tag 剥离后字面「砍掉」会消失，先扫再删才能命中
    zombie_h2_paras = [p for p in doc.paragraphs
                       if p.style and p.style.name.startswith('Heading')
                       and ZOMBIE_HEADING.search(p.text)]
    for h2_para in zombie_h2_paras:
        body = h2_para._element.getparent()
        # 找下一个 Heading 段位置（边界）
        cur = h2_para._element
        siblings = []
        while cur is not None:
            siblings.append(cur)
            nxt = cur.getnext()
            if nxt is None: break
            tag = nxt.tag.split('}')[-1]
            if tag == 'p':
                pStyle = nxt.find(f'{{{W_NS}}}pPr/{{{W_NS}}}pStyle')
                if pStyle is not None:
                    val = pStyle.get(f'{{{W_NS}}}val') or ''
                    if val.startswith('Heading') or val in ('1', '2', '3'):
                        break
            cur = nxt
        for el in siblings:
            body.remove(el)

    # 1. 全文清流水账：body paragraph + 所有 cell（含场景地图表 T4 等非 scene_table 范围）
    for p in doc.paragraphs:
        new = _clean_text(p.text, PRD_CHANGELOG_PATTERNS)
        # H2/H3 加激进 V 流水清理（标题应描述当前态，不论动作词）
        if p.style and p.style.name.startswith('Heading'):
            new = H2_V_TAG.sub('', new).rstrip()
        if new != p.text:
            replace_para_text(p, new)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for cp in cell.paragraphs:
                    new = _clean_text(cp.text, PRD_CHANGELOG_PATTERNS)
                    if new != cp.text:
                        replace_para_text(cp, new)

    # 1b. 修视觉问题：旧蓝色 cell shading + 白字+浅底
    for table in doc.tables:
        rows = list(table.rows)
        is_scene_table = (len(rows) == 1 and len(rows[0].cells) == 2)
        for ri, row in enumerate(rows):
            for ci, cell in enumerate(row.cells):
                tcPr = cell._element.find(qn('w:tcPr'))
                fill = ''
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = (shd.get(qn('w:fill')) or '').upper()
                # 旧蓝色 → 按位置归一为新规范配色
                if fill in LEGACY_BLUE_FILLS:
                    if is_scene_table and ci == 0:
                        set_cell_bg(cell, C['tableAltRow'])
                    elif is_scene_table and ci == 1:
                        tcPr.remove(shd)
                    elif ri == 0:
                        set_cell_bg(cell, C['tableHeaderBg'])
                    else:
                        set_cell_bg(cell, C['tableAltRow'])
                    fill = ''  # 重新评估：清完后 cell 不再算深色
                # 白字 + 非深色填色 → 白字看不见，改 textPrimary
                if fill not in DARK_FILLS_KEEP_WHITE:
                    for run in cell._element.iter(qn('w:r')):
                        rPr = run.find(qn('w:rPr'))
                        if rPr is None: continue
                        color = rPr.find(qn('w:color'))
                        if color is None: continue
                        if (color.get(qn('w:val')) or '').upper() == 'FFFFFF':
                            color.set(qn('w:val'), C['textPrimary'])

    # 2. scene 表右列：讲人话 + 砍 UI + 清死 bullet + 重排编号
    for ti in scene_table_indices:
        if ti >= len(doc.tables):
            continue
        t = doc.tables[ti]
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                if ci == 0:
                    # 左 anchor 列只清流水账
                    for p in cell.paragraphs:
                        new = _clean_text(p.text, PRD_CHANGELOG_PATTERNS)
                        if new != p.text:
                            replace_para_text(p, new)
                    continue
                _kill_dead_bullets(cell)
                for p in cell.paragraphs:
                    new = _clean_text(
                        p.text,
                        PRD_CHANGELOG_PATTERNS,
                        jargon,
                        PRD_UI_STRIP_PATTERNS,
                        PRD_TRAILING_JUNK_PATTERNS,
                    )
                    if new != p.text:
                        replace_para_text(p, new)
                _kill_dead_bullets(cell)
                _renumber_cell_bullets(cell)
                _cleanup_blank_paragraphs(cell)
                # scene_table 右列去 list 项的 bold（仅识别 list 段去 bold，title 段保留）
                for para in cell._element.findall(qn('w:p')):
                    text = ''.join(t.text or '' for t in para.iter(qn('w:t')))
                    if not _LIST_PREFIX.match(text):
                        continue  # title / 普通段，保留 bold
                    for run in para.findall(qn('w:r')):
                        rPr = run.find(qn('w:rPr'))
                        if rPr is None: continue
                        b = rPr.find(qn('w:b'))
                        if b is not None:
                            rPr.remove(b)

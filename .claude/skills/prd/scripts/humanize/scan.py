"""扫描器（read 端，被 check_prd.sh / push gate 调用）。"""
from docx.oxml.ns import qn

from .patterns import (
    CSS_IMPL,
    DARK_FILLS_KEEP_WHITE,
    DATE_TAG,
    DIRTY_CELL_MIN_CHARS,
    EMOJI_RE,
    EXEMPT_H2_KW,
    EXEMPT_HEADER_KW,
    H2_V_TAG,
    LEGACY_BLUE_FILLS,
    SNAKE_FIELD,
    W_NS,
    ZOMBIE_HEADING,
)


def _build_h2_index(doc):
    """返回 {body 子元素: 当前 H2 标题} 映射。直接读 pStyle val，避免 Paragraph wrapper 出错。"""
    body = doc.element.body
    elem_to_h2 = {}
    current_h2 = ''
    for elem in body.iterchildren():
        tag = elem.tag.split('}')[-1]
        if tag == 'p':
            pStyle = elem.find(f'{{{W_NS}}}pPr/{{{W_NS}}}pStyle')
            style_val = pStyle.get(f'{{{W_NS}}}val') if pStyle is not None else ''
            if style_val in ('Heading2', '2'):
                current_h2 = ''.join(t.text or '' for t in elem.iter(f'{{{W_NS}}}t')).strip()
            elem_to_h2[elem] = current_h2
        elif tag == 'tbl':
            elem_to_h2[elem] = current_h2
    return elem_to_h2


def _is_exempt_h2(h2_text: str) -> bool:
    return any(kw in h2_text for kw in EXEMPT_H2_KW)


def _is_field_table(table) -> bool:
    """表头命中字段表关键词 → 豁免（兜底 H2 章节定位失效的场景）。"""
    if not table.rows:
        return False
    headers = [c.text.strip() for c in table.rows[0].cells]
    return any(any(kw in h for kw in EXEMPT_HEADER_KW) for h in headers)


def scan_human_voice(doc) -> dict:
    """扫描 PRD docx 的「讲人话」违规。返回多类 hit 列表。

    Args:
        doc: python-docx Document

    Returns:
        {
            'date_tag_hits':       list[str],  # 流水账日期 / 版本标记（FAIL）
            'snake_field_hits':    list[str],  # snake_case 字段名（WARN，字段表已豁免）
            'css_impl_hits':       list[str],  # CSS 实现细节（WARN）
            'zombie_heading_hits': list[str],  # 僵尸 H2/H3（FAIL，应物理删除）
            'v_tag_heading_hits':  list[str],  # H2/H3 V 版本流水（FAIL）
            'legacy_blue_hits':    list[str],  # 旧蓝色 cell shading（FAIL）
            'invisible_text_hits': list[str],  # 白字+浅底视觉死字（FAIL）
        }
    """
    elem_to_h2 = _build_h2_index(doc)
    date_tag_hits, snake_field_hits, css_impl_hits = [], [], []
    zombie_heading_hits, v_tag_heading_hits = [], []
    legacy_blue_hits, invisible_text_hits = [], []
    dirty_cell_hits = []

    for ti, t in enumerate(doc.tables):
        # 旧蓝色 + 视觉死字扫所有表（含封面）— 视觉问题不应跳过
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                tcPr = cell._element.find(qn('w:tcPr'))
                fill = ''
                if tcPr is not None:
                    shd = tcPr.find(qn('w:shd'))
                    if shd is not None:
                        fill = (shd.get(qn('w:fill')) or '').upper()
                if fill in LEGACY_BLUE_FILLS:
                    legacy_blue_hits.append(f'T{ti}R{ri}C{ci}: fill=#{fill}')
                # 白字 + 非深色填色 → 字看不见
                if fill not in DARK_FILLS_KEEP_WHITE:
                    for run in cell._element.iter(qn('w:r')):
                        rPr = run.find(qn('w:rPr'))
                        if rPr is None: continue
                        color = rPr.find(qn('w:color'))
                        if color is None: continue
                        if (color.get(qn('w:val')) or '').upper() == 'FFFFFF':
                            txt = ''.join(t.text or '' for t in run.iter(qn('w:t')))[:30]
                            invisible_text_hits.append(f'T{ti}R{ri}C{ci}: {txt!r}')
                            break
                # Dirty cell：cell 任一段超长 + 多 list 标记 = set_cell_text 误塞 \n 串
                if ti >= 3 and ci != 0:  # 跳封面 / 元信息 + scene_table 左 anchor
                    for pi, p in enumerate(cell.paragraphs):
                        text = p.text
                        if len(text) < DIRTY_CELL_MIN_CHARS: continue
                        emoji_n = len(EMOJI_RE.findall(text))
                        arrow_n = text.count('→')
                        # ≥ 2 emoji 或 ≥ 2 个 → 分隔 = 多 list 项被塞一段
                        if emoji_n >= 2 or arrow_n >= 2:
                            dirty_cell_hits.append(
                                f'T{ti}R{ri}C{ci} P{pi} ({len(text)}字 emoji={emoji_n} →={arrow_n}): '
                                f'{text[:60]}...'
                            )

        if ti < 3:
            continue  # 文本扫描跳过封面 / 元信息 / 1.3 变更范围
        t_h2 = elem_to_h2.get(t._element, '')
        t_exempt_snake = _is_exempt_h2(t_h2) or _is_field_table(t)
        for ri, row in enumerate(t.rows):
            if ri == 0:
                continue
            for ci, cell in enumerate(row.cells):
                txt = cell.text
                m = DATE_TAG.search(txt)
                if m:
                    date_tag_hits.append(f'T{ti}R{ri}C{ci}: {m.group()}')
                m2 = CSS_IMPL.search(txt)
                if m2:
                    css_impl_hits.append(f'T{ti}R{ri}C{ci}: {m2.group()}')
                if not t_exempt_snake:
                    sm = SNAKE_FIELD.search(txt)
                    if sm:
                        snake_field_hits.append(f'T{ti}R{ri}C{ci}: {sm.group()}')

    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        # Heading 段：扫僵尸标题 + V 流水
        if style.startswith('Heading'):
            if ZOMBIE_HEADING.search(p.text):
                zombie_heading_hits.append(f'[{style}] {p.text.strip()[:80]}')
            if H2_V_TAG.search(p.text):
                v_tag_heading_hits.append(f'[{style}] {p.text.strip()[:80]}')
            continue
        m = DATE_TAG.search(p.text)
        if m:
            date_tag_hits.append(f'P[{style}]: {m.group()}')
        m2 = CSS_IMPL.search(p.text)
        if m2:
            css_impl_hits.append(f'P[{style}]: {m2.group()}')
        p_h2 = elem_to_h2.get(p._element, '')
        if not _is_exempt_h2(p_h2):
            sm = SNAKE_FIELD.search(p.text)
            if sm:
                snake_field_hits.append(f'P[{style}]: {sm.group()}')

    return {
        'date_tag_hits': date_tag_hits,
        'snake_field_hits': snake_field_hits,
        'css_impl_hits': css_impl_hits,
        'zombie_heading_hits': zombie_heading_hits,
        'v_tag_heading_hits': v_tag_heading_hits,
        'legacy_blue_hits': legacy_blue_hits,
        'invisible_text_hits': invisible_text_hits,
        'dirty_cell_hits': dirty_cell_hits,
    }

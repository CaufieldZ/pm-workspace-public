"""PRD 结构性 / 内容性扫描器。

与 scan_human_voice (scan.py) 互补:scan_human_voice 扫「讲人话」三类
(流水账日期 / snake_case / CSS 实现细节),本模块扫 PRD 结构 / 内容硬错误:
圈数字 / 决策编号 / 占位符 / 1.3 流水账 / 章节用户故事引言 / Scene 右列层次
/ 截图 DPI / 段落表格数 / CJK 半角 / theme / 老字体。

被 save_prd 写时调用 + check_prd.sh 推前调用,共享同一份规则。
"""
import io
import re
import sys
import zipfile
from docx.oxml.ns import qn

from .patterns import (
    CIRCLE_NUM_RE,
    CJK_HALF_PUNCT_RE,
    DECISION_NUM_RE,
    LEGACY_FONTS,
    LOW_DPI_THRESHOLD,
    PLACEHOLDER_TOKENS,
    PRD_CHANGELOG_ITERATION_WORDS,
    SCENE_FLAT_MIN_PARAS,
    SCENE_FLAT_SINGLE_PARA_CHARS,
    TECH_CHAPTER_KW,
    docx_min_paragraphs,
    docx_min_tables,
)


_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_BLIP = '{http://schemas.openxmlformats.org/drawingml/2006/main}blip'

# 第 2 章 H2 「2.X View N · ...」/「2.X Part N · ...」识别多端项目
_VIEW_HEADING_RE = re.compile(r'^\s*\d+\.\d+\s+(?:View|Part)\s+\d+', re.IGNORECASE)
# numbered list 段首
_NUMBERED_LINE_RE = re.compile(r'^\d+\.\s')


def _count_scenes(doc) -> int:
    """从 doc 数 H2 「N.M ...」段落作为场景数(与 check_prd.sh 第 2 章 view_count 同源)。

    save_prd 写时拿不到 scene-list.md,自动数。update / check_prd 路径可显式传入。
    """
    n = 0
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if style != 'Heading 2':
            continue
        if re.match(r'^\s*\d+\.\d+\s+', p.text):
            n += 1
    return n


def _gather_full_text(doc) -> str:
    chunks = [p.text for p in doc.paragraphs]
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                chunks.append(c.text)
    return '\n'.join(chunks)


def _scan_circle_nums(doc, full_text: str) -> list:
    if not CIRCLE_NUM_RE.search(full_text):
        return []
    return sorted(set(CIRCLE_NUM_RE.findall(full_text)))


def _scan_placeholders(full_text: str) -> list:
    return [kw for kw in PLACEHOLDER_TOKENS if kw in full_text]


def _scan_decision_nums(doc) -> list:
    """正文「决策 N」内部 ID:第 0 列 anchor + Heading 段豁免。"""
    hits = []
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                if ci == 0:
                    continue  # scene_table 第 0 列是 anchor,允许编号
                m = DECISION_NUM_RE.search(cell.text)
                if m:
                    hits.append(f'T{ti}R{ri}C{ci}: {cell.text[:40]!r}')
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        if style.startswith('Heading'):
            continue
        if DECISION_NUM_RE.search(p.text):
            hits.append(f'P[{style}]: {p.text[:40]!r}')
    return hits


def _scan_iteration_trace(doc) -> list:
    """1.3 节内 PM 内部多轮决策流水(覆盖条目 / 反转回 / 中间稿 / 上一稿 / 前一版)。

    1.3 锚点 = vs 线上基线 delta,迭代过程归 context.md 第 7 章。
    """
    in_13 = False
    section_paras = []
    for p in doc.paragraphs:
        style = p.style.name if p.style else ''
        text = p.text.strip()
        if style == 'Heading 2':
            if text.startswith('1.3'):
                in_13 = True
                continue
            elif in_13:
                in_13 = False
        elif style == 'Heading 1':
            in_13 = False
        if in_13:
            section_paras.append(text)
    section_text = '\n'.join(section_paras)
    return [w for w in PRD_CHANGELOG_ITERATION_WORDS if w in section_text]


def _scan_missing_story_intro(doc) -> list:
    """第 3/4/5 章 H1 后必须有一段用户故事引言(chapter_story),技术骨架章豁免。"""
    heading1 = []
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ''
        if style == 'Heading 1':
            heading1.append((i, p.text.strip()))

    missing = []
    for i, h1 in heading1:
        if any(kw in h1 for kw in TECH_CHAPTER_KW):
            continue
        next_p = None
        for j in range(i + 1, min(i + 5, len(doc.paragraphs))):
            if doc.paragraphs[j].text.strip():
                next_p = doc.paragraphs[j]
                break
        if next_p is None:
            missing.append(f'{h1!r}(章后无内容)')
            continue
        next_style = next_p.style.name if next_p.style else ''
        if next_style.startswith('Heading'):
            missing.append(f'{h1!r}(直接接 {next_style}, 缺 chapter_story 引言)')
    return missing


def _scan_flat_right_col(doc) -> dict:
    """Scene 右列扁平化:单段 > 100 字 OR ≥ 4 段无 numbered 编号。"""
    flat, no_numbered = [], []
    for ti, t in enumerate(doc.tables):
        if len(t.columns) != 2 or len(t.rows) < 1:
            continue
        right = t.rows[0].cells[1]
        ps = [p for p in right.paragraphs if p.text.strip()]
        if len(ps) == 1 and len(ps[0].text) > SCENE_FLAT_SINGLE_PARA_CHARS:
            flat.append(f'T{ti}({len(ps[0].text)}字)')
            continue
        if len(ps) >= SCENE_FLAT_MIN_PARAS:
            has_num = any(_NUMBERED_LINE_RE.match(p.text) for p in ps)
            if not has_num:
                no_numbered.append(f'T{ti}({len(ps)}段)')
    return {'flat': flat, 'no_numbered': no_numbered}


def _scan_low_dpi_shots(doc) -> int:
    """截图 DPI < 130 计数(Playwright 默认 72,需 fix_dpi 抬到 144)。"""
    try:
        from PIL import Image
    except ImportError:
        return 0
    bad = 0
    for rel in doc.part.rels.values():
        if 'image' not in rel.target_ref:
            continue
        try:
            img = Image.open(io.BytesIO(rel.target_part.blob))
            dpi = img.info.get('dpi', (None, None))
            if not dpi[0] or dpi[0] < LOW_DPI_THRESHOLD:
                bad += 1
        except Exception:
            pass
    return bad


def _scan_count_anomaly(doc, scene_count: int) -> dict:
    """段落数 / 表格数异常:小于期望下限。"""
    paras = len(doc.paragraphs)
    tables = len(doc.tables)
    min_p = docx_min_paragraphs(scene_count)
    min_t = docx_min_tables(scene_count)
    out = {}
    if paras < min_p:
        out['paragraphs'] = (paras, min_p)
    if tables < min_t:
        out['tables'] = (tables, min_t)
    return out


def _scan_cjk_half_punct(full_text: str) -> list:
    return list(set(CJK_HALF_PUNCT_RE.findall(full_text)))[:5]


def _scan_missing_theme(docx_path) -> bool:
    """缺 word/theme/theme1.xml → docDefaults themed 失效 → Word fallback Arial。"""
    if docx_path is None:
        return False
    try:
        with zipfile.ZipFile(str(docx_path)) as z:
            return 'word/theme/theme1.xml' not in z.namelist()
    except Exception:
        return False


def _scan_legacy_fonts(doc) -> int:
    """正文残留老字体 run 计数(Arial / Calibri / Microsoft YaHei 等)。"""
    n = 0
    for r in doc.element.body.iter(f'{{{_W}}}r'):
        rPr = r.find(qn('w:rPr'))
        if rPr is None:
            continue
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            continue
        for slot in ('ascii', 'hAnsi', 'eastAsia'):
            if rFonts.get(qn(f'w:{slot}')) in LEGACY_FONTS:
                n += 1
                break
    return n


def _scan_chapter2_multi_end(doc) -> dict:
    """≥ 2 端项目第 2 章 2.1 用户旅程主线 / 2.3 跨端时序 缺失或缺嵌图(warn 不阻断)。"""
    view_count = sum(
        1 for p in doc.paragraphs
        if p.style and p.style.name == 'Heading 2' and _VIEW_HEADING_RE.match(p.text)
    )
    if view_count < 2:
        return {'view_count': view_count, 'warns': []}

    para_idx_21 = para_idx_22 = para_idx_23 = para_idx_3 = -1
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ''
        if style != 'Heading 2':
            continue
        text = p.text.strip()
        if '2.1 用户旅程主线' in text or text.startswith('2.1 用户旅程'):
            para_idx_21 = i
        elif '2.2 端' in text or '2.2 角色' in text:
            para_idx_22 = i
        elif '2.3 跨端时序' in text or text.startswith('2.3 跨端'):
            para_idx_23 = i
    for i, p in enumerate(doc.paragraphs):
        style = p.style.name if p.style else ''
        if style == 'Heading 1' and i > para_idx_23 >= 0:
            para_idx_3 = i
            break

    def _has_image_in_range(start, end):
        for p in doc.paragraphs[start:end]:
            if p._element.findall(f'.//{_BLIP}'):
                return True
        return False

    warns = []
    if para_idx_21 < 0:
        warns.append('缺 H2 "2.1 用户旅程主线"')
    elif para_idx_22 > 0 and not _has_image_in_range(para_idx_21, para_idx_22):
        warns.append('2.1 用户旅程主线无嵌入图')
    if para_idx_23 < 0:
        warns.append('缺 H2 "2.3 跨端时序"')
    elif para_idx_3 > 0 and not _has_image_in_range(para_idx_23, para_idx_3):
        warns.append('2.3 跨端时序无嵌入图')
    return {'view_count': view_count, 'warns': warns}


def scan_prd_structural(doc, scene_count: int = None, docx_path=None) -> dict:
    """扫描 PRD docx 的结构性 / 内容性硬错误。

    Args:
        doc: python-docx Document
        scene_count: 场景数。缺省自动从 doc 数 H2 「N.M」段落
        docx_path: docx 文件路径(仅用于扫 zip 内 theme1.xml,可选)

    Returns:
        dict 形态对齐 scan_human_voice,各字段非空即视为违规
    """
    if scene_count is None:
        scene_count = _count_scenes(doc)
    full_text = _gather_full_text(doc)

    flat_result = _scan_flat_right_col(doc)
    chapter2 = _scan_chapter2_multi_end(doc)

    return {
        'scene_count': scene_count,
        'circle_nums': _scan_circle_nums(doc, full_text),
        'decision_nums': _scan_decision_nums(doc),
        'placeholders': _scan_placeholders(full_text),
        'iteration_trace': _scan_iteration_trace(doc),
        'missing_story_intro': _scan_missing_story_intro(doc),
        'flat_right_col': flat_result['flat'],
        'no_numbered_right_col': flat_result['no_numbered'],
        'low_dpi_shot_count': _scan_low_dpi_shots(doc),
        'count_anomaly': _scan_count_anomaly(doc, scene_count),
        'cjk_half_punct': _scan_cjk_half_punct(full_text),
        'missing_theme': _scan_missing_theme(docx_path),
        'legacy_font_count': _scan_legacy_fonts(doc),
        'chapter2': chapter2,
    }


def report_structural_violations(result: dict, file=sys.stderr) -> int:
    """打印 scan_prd_structural 结果到 stderr,返回 fail 标志(1 = 有 FAIL 级违规)。

    形态对齐 report_violations:WARN 仅打印,不影响 fail。
    """
    fail = 0

    if result['circle_nums']:
        print(f"❌ 圈数字残留(CLAUDE.md 禁止 ①②③, 统一 1./2./3.): {''.join(result['circle_nums'])}", file=file)
        fail = 1

    if result['placeholders']:
        for kw in result['placeholders']:
            print(f"❌ 残留占位符: {kw}", file=file)
        fail = 1

    if result['decision_nums']:
        n = len(result['decision_nums'])
        print(f"❌ 正文残留决策编号 {n} 处(SKILL.md 硬规则 #11): 样例 {result['decision_nums'][:5]}", file=file)
        fail = 1

    if result['iteration_trace']:
        print(f"❌ PRD 1.3 节残留迭代痕迹: {result['iteration_trace']}", file=file)
        print('   规则源: soul.md 4-28「1.3 必须收敛到最终态一行,迭代过程归 context.md 第 7 章」', file=file)
        print('   修法: 把多轮决策流水删掉, 只保留 vs wiki 线上基线的 delta', file=file)
        fail = 1

    if result['missing_story_intro']:
        n = len(result['missing_story_intro'])
        print(f"❌ 功能章缺用户故事引言 {n} 处(pm-workflow.md「PART/章节用户故事陈述」强制): {result['missing_story_intro'][:5]}", file=file)
        print('   修复: h1(doc, "3. ...") 后紧跟 chapter_story(doc, "用户...一句话")', file=file)
        fail = 1

    if result['flat_right_col']:
        print(f"❌ Scene 右列扁平化(单段落>100字符, 用 set_cell_blocks 重建): {', '.join(result['flat_right_col'])}", file=file)
        fail = 1

    if result['no_numbered_right_col']:
        print(f"❌ Scene 右列缺 numbered list(≥4 段且无 \"N.\" 编号, prd-template 强制): {', '.join(result['no_numbered_right_col'])}", file=file)
        fail = 1

    ca = result['count_anomaly']
    if ca:
        if 'paragraphs' in ca:
            paras, mn = ca['paragraphs']
            print(f'❌ 段落数异常: {paras}, 期望 > {mn}', file=file)
            fail = 1
        if 'tables' in ca:
            tables, mn = ca['tables']
            print(f'❌ 表格数异常: {tables}, 期望 > {mn}', file=file)
            fail = 1

    if result['cjk_half_punct']:
        n_total = '?'  # 只保了样例,没数总数
        print(f"❌ 中文相邻半角标点(调 normalize_punctuation): 样例 {result['cjk_half_punct']}", file=file)
        fail = 1

    if result['low_dpi_shot_count']:
        n = result['low_dpi_shot_count']
        print(f'❌ 截图 DPI<130 有 {n} 张(Playwright 默认 72, 需 fix_dpi 到 144)', file=file)
        fail = 1

    if result['missing_theme']:
        print('❌ 缺 word/theme/theme1.xml(老 python-docx 模板漏注入,导致 docDefaults themed 失效→ Word fallback Arial)。', file=file)
        print('   修复: from update_prd_base import ensure_theme; ensure_theme(docx_path)', file=file)
        fail = 1

    if result['legacy_font_count']:
        n = result['legacy_font_count']
        print(f'❌ 正文残留老字体 run {n} 个(Arial/Calibri/Microsoft YaHei 等,会渲染成 Arial 风)。', file=file)
        print('   修复: from update_prd_base import normalize_fonts; normalize_fonts(doc)', file=file)
        fail = 1

    # WARN(不阻断)
    ch2 = result['chapter2']
    if ch2['warns']:
        for w in ch2['warns']:
            print(f"⚠️  PRD 涉及 {ch2['view_count']} 端: {w}", file=file)
        print('   规范: prd-template.md 第 2 章; sections.journey_main_section / cross_end_sequence_section', file=file)

    return fail

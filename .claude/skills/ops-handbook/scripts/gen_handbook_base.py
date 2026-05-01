#!/usr/bin/env python3
"""ops-handbook 框架函数 —— 薄封装层，复用 PRD core/ 美学栈。

调用方式（项目脚本 projects/{项目}/scripts/gen_handbook_v{N}.py）：

    import sys
    from pathlib import Path
    BASE = Path(__file__).resolve().parents[3]
    sys.path.insert(0, str(BASE / '.claude/skills/ops-handbook/scripts'))
    from gen_handbook_base import *

提供的运营手册专属 helper：
- cover_page          封面 + 元信息表（产品/版本/受众/维护人/范围）
- step_table          三列步骤表（动作 / 校验 / 异常态）
- field_table         字段定义表（允许 snake_case）
- faq_block           FAQ 单条（Q 粗体 + A 缩进）
- permission_matrix   角色权限矩阵
- save_handbook       保存 + 字体三件套（不调 humanize，运营手册允许 snake_case + 像素）

复用 PRD core 的（直接 re-export）：
- styles:    C / FONT / para_run / cell_text / set_cell_bg / set_cell_border
- headings:  h1 / h2 / h3 / add_p / bullet
- tables:    make_table / set_cell_text / set_cell_blocks
- normalize: normalize_punctuation / normalize_fonts / ensure_theme
- patch:     replace_para_text / insert_heading_before / insert_paragraph_before
- images:    fix_dpi / replace_cell_image
"""
from __future__ import annotations

import sys
from pathlib import Path

# ── 接 PRD core 美学栈 ────────────────────────────────────────────────
_PRD_SCRIPTS = Path(__file__).resolve().parent.parent.parent / 'prd' / 'scripts'
if str(_PRD_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(_PRD_SCRIPTS))

from docx import Document  # noqa: E402
from docx.shared import Pt, Cm  # noqa: E402
from docx.enum.text import WD_ALIGN_PARAGRAPH  # noqa: E402
from docx.enum.table import WD_TABLE_ALIGNMENT  # noqa: E402

from core.styles import (  # noqa: E402
    C, FONT, set_cell_bg, set_cell_border,
    para_run, para_run_md, cell_text,
)
from core.headings import (  # noqa: E402
    add_p, bullet,
    h1, h2, h3,
)
from core.tables import (  # noqa: E402
    make_table, set_cell_text, set_cell_blocks,
)
from core.images import (  # noqa: E402
    fix_dpi, replace_cell_image,
)
from core.normalize import (  # noqa: E402
    normalize_punctuation, normalize_fonts, ensure_theme,
)
from core.patch import (  # noqa: E402
    replace_para_text, search_replace_para,
    insert_heading_before, insert_paragraph_before, insert_description_after,
    remove_table,
)
from core.headings import normalize_headings  # noqa: E402


# ── 文档初始化 ────────────────────────────────────────────────────────
def init_doc(landscape: bool = False) -> Document:
    """运营手册默认竖版（A4），与 PRD landscape 横版不同。"""
    doc = Document()
    section = doc.sections[0]
    if landscape:
        from docx.enum.section import WD_ORIENT
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    return doc


def cover_page(doc, *, title: str, version: str, owner: str,
               audience: str, scope: str = '', extra_meta: list[tuple[str, str]] | None = None):
    """运营手册封面：大标题 + 副标题 + 元信息属性表。

    Args:
        title: 主标题，如「活动中心运营操作手册」
        version: 版本号，如「V3.0」
        owner: 维护人
        audience: 受众，如「业务运营 / 总运营 / 客服」
        scope: 适用范围，如「活动中心 MGT 后台 9 个页面」
        extra_meta: 额外元信息行 [(label, value), ...]
    """
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_run(title_p, title, font='title', size_pt=28, bold=True,
             color=C['textHeading'])

    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para_run(subtitle_p, '运营操作手册', font='title', size_pt=14,
             color=C['textSecondary'])

    doc.add_paragraph()

    from datetime import date
    meta = [
        ('文档版本', version),
        ('适用范围', scope or '—'),
        ('受众', audience),
        ('维护人', owner),
        ('上次更新', date.today().isoformat()),
    ]
    if extra_meta:
        meta.extend(extra_meta)

    table = doc.add_table(rows=len(meta), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for r, (label, value) in enumerate(meta):
        cell_text(table.rows[r].cells[0], label, size_pt=10, bold=True)
        cell_text(table.rows[r].cells[1], value, size_pt=10)
        table.rows[r].cells[0].width = Cm(4.5)
        table.rows[r].cells[1].width = Cm(11)
    doc.add_page_break()


# ── 三列步骤表（动作 / 校验 / 异常） ────────────────────────────────
def step_table(doc, *, step_id: str, name: str,
               rows: list[tuple[str, str, str]]):
    """运营手册核心：每个操作步骤的三列展开。

    Args:
        step_id: 步骤编号，如「3.1」
        name: 步骤名称，如「同步活动从哪来」
        rows: [(action, check, exception), ...]
            action:    第几步具体怎么操作（动词开头）
            check:     看到什么 / 校验规则
            exception: 异常态怎么办（XX 不显示 / 报错 等）
    """
    h2(doc, f'{step_id} {name}')

    table = make_table(
        doc,
        headers=['#', '动作', '看到 / 校验', '异常态'],
        rows_data=[
            [str(i + 1), action, check, exception]
            for i, (action, check, exception) in enumerate(rows)
        ],
        col_widths_cm=[1.0, 5.5, 5.0, 4.5],
    )
    return table


# ── 字段定义表（允许 snake_case） ───────────────────────────────────
def field_table(doc, fields: list[tuple]):
    """字段速查表，运营要照着 CMS 字段名搜，允许 snake_case。

    Args:
        fields: 每行 5 元组 (field_name, cn_name, field_type, range_, c_side_show)
            field_name:  CMS 内部字段名（snake_case 允许）
            cn_name:     运营口头叫法
            field_type:  类型 / 枚举 / 多选等
            range_:      取值范围 / 枚举数量
            c_side_show: C 端展示效果
    """
    return make_table(
        doc,
        headers=['字段名', '中文名', '类型', '取值范围', 'C 端展示'],
        rows_data=[list(row) for row in fields],
        col_widths_cm=[3.5, 3.0, 2.0, 3.5, 4.0],
    )


# ── FAQ 单条块 ───────────────────────────────────────────────────────
def faq_block(doc, *, question: str, answer: str):
    """单条 FAQ：Q 粗体红色，A 正文缩进。"""
    q_p = doc.add_paragraph()
    para_run(q_p, f'Q · {question}', font='title', size_pt=11, bold=True,
             color=C['accentBlue'])

    a_p = doc.add_paragraph()
    a_p.paragraph_format.left_indent = Cm(0.6)
    para_run(a_p, answer, size_pt=10, color=C['textSecondary'])
    doc.add_paragraph()


# ── 角色权限矩阵 ────────────────────────────────────────────────────
def permission_matrix(doc, *, roles: list[str], capabilities: list[tuple]):
    """权限矩阵：左列能力，每个角色一列。

    Args:
        roles: 角色列表，如 ['业务运营', '总运营', '客服']
        capabilities: 每行 (capability, *role_values)
            capability:   能力描述
            role_values:  对应每个角色的权限（'✓' / '✗' / '业务线内' / '全部' 等）
    """
    headers = ['操作能力'] + roles
    rows = [list(cap) for cap in capabilities]
    return make_table(
        doc,
        headers=headers,
        rows_data=rows,
        col_widths_cm=[6.0] + [3.5] * len(roles),
    )


# ── 一键保存（字体三件套，不调 humanize） ───────────────────────────
def save_handbook(doc, path: str | Path) -> str:
    """保存 + 标点 / 标题 / 字体规范化 + 注入 theme1.xml。

    与 prd.save_prd 区别：
    - 不调 humanize（运营手册接受 snake_case + 像素）
    - 不做档位体检（手册无三档分层）
    """
    path = str(path)
    normalize_punctuation(doc)
    normalize_headings(doc)
    normalize_fonts(doc)
    doc.save(path)
    ensure_theme(path)
    return path


__all__ = [
    # docx 类型 / 标准库（保留 from gen_handbook_base import * 拿到）
    'Document', 'Pt', 'Cm', 'WD_ALIGN_PARAGRAPH', 'WD_TABLE_ALIGNMENT',
    # core re-export
    'C', 'FONT',
    'set_cell_bg', 'set_cell_border', 'para_run', 'para_run_md', 'cell_text',
    'add_p', 'bullet', 'h1', 'h2', 'h3',
    'make_table', 'set_cell_text', 'set_cell_blocks',
    'fix_dpi', 'replace_cell_image',
    'normalize_punctuation', 'normalize_fonts', 'normalize_headings', 'ensure_theme',
    'replace_para_text', 'search_replace_para',
    'insert_heading_before', 'insert_paragraph_before', 'insert_description_after',
    'remove_table',
    # ops-handbook 专属
    'init_doc', 'cover_page',
    'step_table', 'field_table', 'faq_block', 'permission_matrix',
    'save_handbook',
]

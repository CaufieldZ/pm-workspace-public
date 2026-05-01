#!/usr/bin/env python3
"""ops-handbook · CMS 后台截图回填工具

复用 PRD core/images 的 fix_dpi。骨架版本，按 steps.json 配置驱动；项目落地时
按需扩展（多步翻页 / hover state / drawer 截图等），写在
projects/{项目}/scripts/handbook_screenshots_v{N}.py 里继承本脚本。

steps.json 格式：

    {
      "base_url": "https://mgt-staging.example.com",
      "cookie_file": "~/.config/ops-cookie.json",
      "steps": [
        {
          "id": "step-1",
          "url": "/activity/list",
          "selector": "#main-content",
          "anchor_keyword": "进入「活动列表」"
        },
        {
          "id": "step-2",
          "url": "/activity/edit/123",
          "selector": ".form-container",
          "anchor_keyword": "填活动名称"
        }
      ]
    }

CLI：

    python3 handbook_screenshots.py --project growth/activity-center  # 拍 + 回填
    python3 handbook_screenshots.py --project XX --shot-only       # 只拍
    python3 handbook_screenshots.py --project XX --insert-only     # 只回填
    python3 handbook_screenshots.py --project XX --steps step-1    # 只拍指定 step
"""
from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

# ── 接 PRD core 复用 fix_dpi ──────────────────────────────────────────
_PRD_SCRIPTS = Path(__file__).resolve().parent.parent.parent / 'prd' / 'scripts'
if str(_PRD_SCRIPTS) not in sys.path:
    sys.path.insert(0, str(_PRD_SCRIPTS))

from core.images import fix_dpi  # noqa: E402


def take_screenshots(project_dir: Path, *, only_steps: list[str] | None = None) -> list[Path]:
    """按 steps.json 用 Playwright 截 CMS 后台截图，存 screenshots/handbook/。"""
    config_path = project_dir / 'scripts' / 'handbook_steps.json'
    if not config_path.exists():
        raise FileNotFoundError(
            f'缺 {config_path}。请创建 handbook_steps.json，参考 ops-handbook SKILL.md 截图章节。'
        )
    config = json.loads(config_path.read_text(encoding='utf-8'))

    out_dir = project_dir / 'screenshots' / 'handbook'
    out_dir.mkdir(parents=True, exist_ok=True)

    try:
        from playwright.sync_api import sync_playwright
    except ImportError:
        raise SystemExit(
            '缺 Playwright：cd projects && npm init playwright@latest && npx playwright install chromium'
        )

    base_url = config['base_url'].rstrip('/')
    cookie_file = config.get('cookie_file')

    written: list[Path] = []
    with sync_playwright() as p:
        browser = p.chromium.launch()
        context = browser.new_context(
            viewport={'width': 1440, 'height': 900},
            device_scale_factor=2,
        )
        if cookie_file:
            cookies_path = Path(cookie_file).expanduser()
            if cookies_path.exists():
                context.add_cookies(json.loads(cookies_path.read_text(encoding='utf-8')))
            else:
                print(f'⚠️  cookie_file 不存在: {cookies_path}（截图可能因未登录失败）', file=sys.stderr)
        page = context.new_page()

        for step in config['steps']:
            sid = step['id']
            if only_steps and sid not in only_steps:
                continue
            url = base_url + step['url']
            selector = step.get('selector')
            print(f'  [{sid}] → {url}')
            page.goto(url, wait_until='networkidle')
            page.evaluate("""
                document.querySelectorAll(
                  '.modal-mask, .ant-modal-mask, .drawer-mask, [data-overlay]'
                ).forEach(el => el.style.display = 'none');
            """)
            out_path = out_dir / f'{sid}.png'
            target = page.locator(selector) if selector else page
            target.screenshot(path=str(out_path))
            fix_dpi(str(out_path))
            written.append(out_path)
            print(f'    ✓ {out_path.name}')

        browser.close()
    return written


def insert_into_docx(docx_path: Path, screenshots_dir: Path,
                     *, anchor_map: dict[str, str] | None = None,
                     img_width_cm: float = 13.0):
    """在 docx 里按 anchor_keyword 找段落，紧跟其后插入对应截图。

    anchor_map: {step_id: anchor_keyword}。未传则从同目录 handbook_steps.json 读。
    img_width_cm: 截图宽度（默认 13cm，A4 竖版可用宽度 ≈ 16cm，留一点边距）。
    """
    from docx import Document
    from docx.shared import Cm

    doc = Document(str(docx_path))

    if anchor_map is None:
        config_path = screenshots_dir.parent.parent / 'scripts' / 'handbook_steps.json'
        config = json.loads(config_path.read_text(encoding='utf-8'))
        anchor_map = {s['id']: s['anchor_keyword'] for s in config['steps']
                      if 'anchor_keyword' in s}

    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    inserted = 0
    for step_id, keyword in anchor_map.items():
        png = screenshots_dir / f'{step_id}.png'
        if not png.exists():
            print(f'  ⚠️  截图缺失: {png.name}', file=sys.stderr)
            continue

        target_para = None
        for p in doc.paragraphs:
            if keyword in p.text:
                target_para = p
                break
        if target_para is None:
            print(f'  ⚠️  找不到 anchor「{keyword}」，跳过 {step_id}', file=sys.stderr)
            continue

        new_p_xml = OxmlElement('w:p')
        target_para._p.addnext(new_p_xml)
        new_para = Paragraph(new_p_xml, target_para._parent)
        new_para.add_run().add_picture(str(png), width=Cm(img_width_cm))
        inserted += 1
        print(f'  ✓ {step_id} 插入到「{keyword[:20]}...」后')

    doc.save(str(docx_path))
    print(f'插入 {inserted} 张截图 → {docx_path}')


def find_handbook_docx(project_dir: Path) -> Path:
    """找 deliverables/ 下最新的 handbook-*.docx。"""
    candidates = sorted((project_dir / 'deliverables').glob('handbook-*.docx'),
                        key=lambda p: p.stat().st_mtime, reverse=True)
    if not candidates:
        raise FileNotFoundError(f'{project_dir}/deliverables/ 下无 handbook-*.docx')
    return candidates[0]


def main():
    ap = argparse.ArgumentParser(description='ops-handbook CMS 截图 + 回填')
    ap.add_argument('--project', required=True, help='项目名（projects/{项目}/）')
    ap.add_argument('--shot-only', action='store_true', help='只拍截图不回填')
    ap.add_argument('--insert-only', action='store_true', help='只回填不重拍')
    ap.add_argument('--steps', help='只处理这些 step（逗号分隔）')
    args = ap.parse_args()

    BASE = Path(__file__).resolve().parents[3]
    project_dir = BASE / 'projects' / args.project
    if not project_dir.exists():
        raise SystemExit(f'项目不存在: {project_dir}')

    only_steps = args.steps.split(',') if args.steps else None

    if not args.insert_only:
        print('=== 拍 CMS 截图 ===')
        take_screenshots(project_dir, only_steps=only_steps)

    if not args.shot_only:
        print('=== 回填到 docx ===')
        docx_path = find_handbook_docx(project_dir)
        insert_into_docx(docx_path, project_dir / 'screenshots' / 'handbook')


if __name__ == '__main__':
    main()
